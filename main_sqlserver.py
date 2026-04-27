#!/usr/bin/env python3
# -*- coding:utf-8 -*-
#
# Copyright (c) 2024 DBCheck Contributors
# sdfiyon@gmail.com
#
# This file is part of DBCheck, an open-source database health inspection tool.
# DBCheck is released under the MIT License.
# See LICENSE or visit https://opensource.org/licenses/MIT for full license text.
#
from version import __version__ as VER

# 磁盘采集时忽略的外接 ISO / Media 挂载点前缀
IGNORE_MOUNTS = {'/mnt/iso', '/media', '/run/media', '/iso', '/cdrom'}

import warnings
warnings.filterwarnings("ignore")
import itertools
import math
import sys
import datetime
import argparse
import subprocess
import logging
import logging.handlers
import socket
import re
import time
from pathlib import Path
import sys, getopt, os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate
import configparser
import importlib
import subprocess
import json
import hashlib
import base64
from datetime import datetime, timedelta
import platform
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import tempfile
import io
import psutil
import shutil
import getpass

importlib.reload(sys)

# ── i18n setup for CLI ─────────────────────────────────────────────
try:
    from i18n import get_lang
    _SQLSERVER_LANG = get_lang()
except Exception:
    _SQLSERVER_LANG = 'zh'

def _t(key):
    try:
        from i18n import t as _tt
        return _tt(key, _SQLSERVER_LANG)
    except Exception:
        return key


def print_banner():
    """打印横幅"""
    print("\n" + "=" * 60)
    print("       " + _t("sqlserver_cli_banner_title"))
    print("       " + _t("sqlserver_cli_banner_subtitle"))
    print("=" * 60)


class ExcelTemplateManager:
    """Excel配置模板管理器（参考main_pg.py实现）"""

    def __init__(self):
        self.template_file = os.path.join(os.path.dirname(__file__), 'configs', 'sqlserver_batch_template.xlsx')

    def create_template(self):
        """创建批量巡检Excel模板"""
        os.makedirs(os.path.dirname(self.template_file), exist_ok=True)

        wb = Workbook()
        ws = wb.active
        ws.title = "SQL Server巡检配置"

        # 表头样式
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")

        # 写入表头
        headers = ["名称", "主机地址", "端口", "用户名", "密码", "数据库", "巡检人", "SSH主机", "SSH端口", "SSH用户", "SSH密码", "SSH密钥文件"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # 写入示例数据
        example_data = [
            ["SQLServer_Prod", "192.168.1.100", "1433", "sa", "YourPassword", "master", "DBA", "", "22", "root", "", ""],
            ["SQLServer_Dev", "192.168.1.101", "1433", "sa", "YourPassword", "master", "DBA", "", "22", "root", "", ""],
        ]
        for row_idx, row_data in enumerate(example_data, 2):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # 设置列宽
        column_widths = [18, 15, 8, 12, 15, 12, 8, 15, 8, 12, 15, 20]
        for idx, width in enumerate(column_widths, 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = width

        try:
            wb.save(self.template_file)
            print("\u2705 " + _t("sqlserver_cli_excel_created").format(path=self.template_file))
            print(_t("sqlserver_cli_excel_fill_note"))
        except Exception as e:
            print("\u274c " + _t("sqlserver_cli_excel_create_fail").format(e=e))

    def read_template(self):
        """读取批量巡检Excel模板"""
        if not os.path.exists(self.template_file):
            return []

        try:
            wb = openpyxl.load_workbook(self.template_file)
            ws = wb.active

            db_list = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row[0] and row[1]:  # 名称和主机地址不能为空
                    db_info = {
                        'name': row[0],
                        'host': row[1],
                        'port': row[2] if row[2] else 1433,
                        'user': row[3] if row[3] else 'sa',
                        'password': row[4] if row[4] else '',
                        'database': row[5] if row[5] else 'master',
                        'inspector': row[6] if row[6] else 'DBA',
                        'ssh_host': row[7] if row[7] else None,
                        'ssh_port': row[8] if row[8] else 22,
                        'ssh_user': row[9] if row[9] else 'root',
                        'ssh_password': row[10] if row[10] else '',
                        'ssh_key_file': row[11] if row[11] else ''
                    }
                    db_list.append(db_info)

            print(_t("sqlserver_cli_excel_read_count").format(n=len(db_list)))
            return db_list
        except Exception as e:
            print(_t("sqlserver_cli_excel_read_fail").format(e=e))
            return []

# ── Markdown → Word 渲染器 ─────────────────────────────────────────────────
import re

def _render_markdown_to_doc(doc, text, default_size=11, ch8_prefix=False):
    """
    将 Markdown 文本渲染为 Word 段落，支持：
    - **加粗**、*斜体*、`行内代码`
    - ## 二级标题（ch8_prefix=True 时自动加 8.X 序号）→ Heading 2
    - ### 三级标题 → Heading 3（无序号）
    - - /*/• 列表项 → bullet paragraph
    - > 引用块 → indented paragraph
    - [text](url) → text（去掉链接）
    """
    CODE_FONT = 'Courier New'
    lines = text.strip().split('\n')
    in_code_block = False
    code_buf = []
    _h2_seq = 0  # 用于 ## 标题的 8.X 序号

    def _add_run(para, md_text, size):
        """解析 md_text 中的 **bold**、*italic*、`code` 并添加 Run"""
        # 先处理行内代码（优先级最高）
        parts = re.split(r'(``[^`]+``|`[^`]+`)', md_text)
        for part in parts:
            if re.match(r'`[^`]+`', part):
                run = para.add_run(part.strip('`'))
                run.font.name = CODE_FONT
                run.font.size = Pt(size - 1)
            else:
                # 处理 **bold** 和 *italic*
                sub_parts = re.split(r'(\*\*[^\*]+\*\*|\*[^\*]+\*)', part)
                for sp in sub_parts:
                    if sp.startswith('**') and sp.endswith('**'):
                        run = para.add_run(sp.strip('**'))
                        run.bold = True
                        run.font.size = Pt(size)
                    elif sp.startswith('*') and sp.endswith('*'):
                        run = para.add_run(sp.strip('*'))
                        run.italic = True
                        run.font.size = Pt(size)
                    elif sp:
                        run = para.add_run(sp)
                        run.font.size = Pt(size)

    for line in lines:
        stripped = line.strip()
        # 代码块
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_buf = []
            else:
                in_code_block = False
                code_para = doc.add_paragraph()
                code_para.style = 'No Spacing'
                code_para.paragraph_format.left_indent = Cm(0.5)
                code_para.paragraph_format.space_before = Pt(2)
                code_para.paragraph_format.space_after = Pt(2)
                for code_line in code_buf:
                    run = code_para.add_run(code_line + '\n')
                    run.font.name = CODE_FONT
                    run.font.size = Pt(size - 1)
                code_para.add_run('\n')
            continue
        if in_code_block:
            code_buf.append(stripped)
            continue
        # 空行 → 跳过（不生成空段落，避免多余间距）
        if not stripped:
            continue
        # 二级标题
        m = re.match(r'^##\s+(.*)', stripped)
        if m:
            _h2_seq += 1
            title_text = m.group(1).strip()
            if ch8_prefix:
                full_title = f"{ch8_prefix.replace('X', str(_h2_seq))} {title_text}"
            else:
                full_title = title_text
            p = doc.add_heading(full_title, level=2)
            p.runs[0].font.size = Pt(13)
            continue
        # 三级标题
        m = re.match(r'^###\s+(.*)', stripped)
        if m:
            title_text = m.group(1).strip()
            p = doc.add_heading(title_text, level=3)
            p.runs[0].font.size = Pt(11)
            continue
        # 列表项
        if re.match(r'^[-*•]\s+', stripped):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5)
            _add_run(p, re.sub(r'^[-*•]\s+', '', stripped), default_size)
            continue
        # 引用块
        if stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            _add_run(p, stripped.lstrip('> '), default_size)
            continue
        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75) if not stripped.startswith(('**', '*', '`')) else None
        # 去掉 [text](url) 中的链接
        clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', stripped)
        _add_run(p, clean_text, default_size)


# ═══════════════════════════════════════════════════════════════════════════
# SQL Server 巡检查询定义
# ═══════════════════════════════════════════════════════════════════════════

class SQLServerQueries:
    """SQL Server 巡检查询模板类"""

    # ── 1. 版本与实例信息 ──────────────────────────────────────────
    QUERY_VERSION = """
-- SQL Server 版本与实例信息
SELECT
    @@VERSION AS [full_version],
    SERVERPROPERTY('ProductVersion') AS [product_version],
    SERVERPROPERTY('ProductLevel') AS [service_pack],
    SERVERPROPERTY('Edition') AS [edition],
    SERVERPROPERTY('MachineName') AS [machine_name],
    SERVERPROPERTY('ServerName') AS [server_name],
    (SELECT sqlserver_start_time FROM sys.dm_os_sys_info) AS [startup_time]
"""

    QUERY_SERVER_CONFIG = """
-- 服务器配置信息
SELECT
    name AS [config_name],
    value AS [config_value],
    description AS [description]
FROM sys.configurations WITH (NOLOCK)
WHERE name IN (
    'max server memory (MB)',
    'min server memory (MB)',
    'max degree of parallelism',
    'cost threshold for parallelism',
    'query wait (s)',
    'fill factor (%)'
)
ORDER BY name
"""

    # ── 2. 数据库状态 ──────────────────────────────────────────────
    QUERY_DATABASES = """
-- 所有用户数据库状态
SELECT
    d.name AS [database_name],
    d.state_desc,
    d.recovery_model_desc,
    d.compatibility_level,
    d.create_date,
    d.collation_name,
    m.data_size_mb,
    m.log_size_mb,
    m.total_size_mb,
    m.data_file_path,
    m.log_file_path
FROM sys.databases d
LEFT JOIN (
    SELECT
        database_id,
        SUM(CASE WHEN type = 0 THEN size * 8.0 / 1024 ELSE 0 END) AS data_size_mb,
        SUM(CASE WHEN type = 1 THEN size * 8.0 / 1024 ELSE 0 END) AS log_size_mb,
        SUM(size * 8.0 / 1024) AS total_size_mb,
        MAX(CASE WHEN type = 0 THEN physical_name END) AS data_file_path,
        MAX(CASE WHEN type = 1 THEN physical_name END) AS log_file_path
    FROM sys.master_files WITH (NOLOCK)
    WHERE database_id > 4
    GROUP BY database_id
) m ON d.database_id = m.database_id
WHERE d.database_id > 4
ORDER BY d.name
"""

    QUERY_DB_FILE_STATS = """
-- 数据库文件读写统计
SELECT
    DB_NAME(database_id) AS [database_name],
    type_desc,
    num_of_reads,
    num_of_writes,
    bytes_read / 1048576.0 AS [read_mb],
    bytes_written / 1048576.0 AS [written_mb],
    io_stall_read_ms / 1000 AS [read_stall_sec],
    io_stall_write_ms / 1000 AS [write_stall_sec]
FROM sys.dm_io_virtual_file_stats(NULL, NULL) vf
INNER JOIN sys.master_files mf ON vf.database_id = mf.database_id AND vf.file_id = mf.file_id
WHERE database_id > 4
ORDER BY io_stall_read_ms + io_stall_write_ms DESC
"""

    # ── 3. 连接与会话 ──────────────────────────────────────────────
    QUERY_CONNECTIONS = """
-- 连接统计
SELECT
    COUNT(*) AS [total_connections],
    SUM(CASE WHEN status = 'running' THEN 1 ELSE 0 END) AS [active_sessions],
    SUM(CASE WHEN is_user_process = 1 THEN 1 ELSE 0 END) AS [user_sessions],
    SUM(CASE WHEN is_user_process = 0 THEN 1 ELSE 0 END) AS [system_sessions],
    @@MAX_CONNECTIONS AS [max_connections],
    CAST(COUNT(*) * 100.0 / @@MAX_CONNECTIONS AS DECIMAL(5,2)) AS [connection_usage_pct]
FROM sys.dm_exec_sessions WITH (NOLOCK)
WHERE session_id > 50
"""

    QUERY_ACTIVE_SESSIONS = """
-- 活动会话详情（TOP 50 按 CPU 时间排序）
SELECT TOP 50
    s.session_id AS [spid],
    s.login_name AS [login_name],
    s.host_name AS [host_name],
    s.program_name,
    s.status,
    s.cpu_time AS [cpu_time_ms],
    s.memory_usage * 8 AS [memory_kb],
    s.reads AS [logical_reads],
    s.writes AS [writes],
    s.total_elapsed_time AS [elapsed_time_ms],
    s.last_request_start_time,
    r.command,
    r.status AS [request_status],
    r.blocking_session_id AS [blocked_by],
    r.wait_type AS [wait_type],
    r.wait_time AS [wait_time_ms],
    r.percent_complete AS [progress_pct],
    SUBSTRING(t.text, 1, 500) AS [sql_text]
FROM sys.dm_exec_sessions s
LEFT JOIN sys.dm_exec_requests r ON s.session_id = r.session_id
OUTER APPLY sys.dm_exec_sql_text(r.sql_handle) t
WHERE s.session_id > 50
ORDER BY s.cpu_time DESC
"""

    QUERY_BLOCKING = """
-- 阻塞会话
SELECT
    BLOCKING.session_id AS [blocking_spid],
    BLOCKING.login_name AS [blocking_login],
    BLOCKING.host_name AS [blocking_host],
    BLOCKING.program_name AS [blocking_program],
    BLOCKING.status AS [blocking_status],
    BLOCKING.cpu_time AS [blocking_cpu_ms],
    WAIT.session_id AS [waiting_spid],
    WAIT.login_name AS [waiting_login],
    WAIT.host_name AS [waiting_host],
    WAIT.program_name AS [waiting_program],
    WAIT.status AS [waiting_status],
    WAIT.wait_type,
    WAIT.wait_time AS [wait_time_ms],
    WAIT.blocking_session_id,
    r.command AS [waiting_command],
    SUBSTRING(t.text, 1, 500) AS [waiting_sql]
FROM sys.dm_exec_sessions WAIT
INNER JOIN sys.dm_exec_sessions BLOCKING ON WAIT.blocking_session_id = BLOCKING.session_id
LEFT JOIN sys.dm_exec_requests r ON WAIT.session_id = r.session_id
OUTER APPLY sys.dm_exec_sql_text(r.sql_handle) t
WHERE WAIT.session_id > 50
ORDER BY WAIT.wait_time DESC
"""

    # ── 4. 等待统计 ────────────────────────────────────────────────
    QUERY_WAIT_STATS = """
-- TOP 20 等待统计
SELECT TOP 20
    wait_type,
    waiting_tasks_count,
    wait_time_ms,
    signal_wait_time_ms,
    (wait_time_ms - signal_wait_time_ms) AS [resource_wait_ms],
    CAST(wait_time_ms * 100.0 / NULLIF(SUM(wait_time_ms) OVER(), 0) AS DECIMAL(10,2)) AS [wait_pct]
FROM sys.dm_os_wait_stats WITH (NOLOCK)
WHERE wait_type NOT IN (
    'RESOURCE_QUEUE', 'SQLTRACE_INCREMENTAL_FLUSH_SLEEP',
    'LOGMGR_QUEUE', 'CHECKPOINT_QUEUE', 'LAZYWRITER_SLEEP',
    'SLEEP_TASK', 'SLEEP_SYSTEMTASK', 'BROKER_TASK_STOP',
    'BROKER_TO_FLUSH', 'BROKER_EVENTHANDLER', 'BROKER_INIT',
    'BROKER_REGISTRATION', 'BROKER_TRANSMITTER', 'CHECKPOINT_QUEUE',
    'CLR_AUTO_EVENT', 'CLR_MANUAL_EVENT', 'CLR_SEMAPHORE',
    'DBMIRROR_DBM_EVENT', 'DBMIRROR_EVENTS_QUEUE', 'DBMIRROR_WORKER_QUEUE',
    'DBMIRRORING_CMD', 'DISPATCHER_QUEUE_SEMAPHORE', 'FT_IFTS_SCHEDULER_IDLE_WAIT',
    'FT_IFTS_WORKER_THREAD', 'HADR_CLUSAPI_CALL', 'HADR_FILESTREAM_IOMGR',
    'HADR_TIMER_TASK', 'HADR_WORK_QUEUE', 'KSOURCE_WAKEUP', 'LAZYWRITER_SLEEP',
    'LOGMGR_QUEUE', 'ONDEMAND_TASK_QUEUE', 'PWAIT_ALL_COMPONENTS_ENABLED',
    'QDS_PERSIST_TASK_MAIN_LOOP_SLEEP', 'QDS_CLEANUP_STALE_TASK',
    'REQUEST_FOR_DEADLOCK_SEARCH', 'RESOURCE_BROKER_EVENTHANDLER',
    'RESOURCE_BROKER_EXECUTION_SEMAPHORE', 'RESOURCE_GOVERNOR_IDLE',
    'RESOURCE_QUEUE', 'SERVER_IDLE_CHECK', 'SLEEP_BPOOL_FLUSH',
    'SLEEP_BPOOL_STATS', 'SLEEP_DBSTARTUP', 'SLEEP_DCOMSTARTUP',
    'SLEEP_MASTERDBREADY', 'SLEEP_MASTERMDREADY', 'SLEEP_MASTERUPGRADED',
    'SLEEP_MSDBSTARTUP', 'SLEEP_SYSTEMTASK', 'SLEEP_TASK',
    'SLEEP_TEMPDBSTARTUP', 'SLEEP_WORKITEM_ASYNC', 'SP_SERVER_DIAGNOSTICS',
    'SQLTRACE_BUFFER_FLUSH', 'SQLTRACE_FILE_BUFFER', 'SQLTRACE_FLUSH',
    'SQLTRACE_INCREMENTAL_FLUSH_SLEEP', 'SQLTRACE_WAIT_ENTRIES', 'WAIT_FOR_RESULTS',
    'WAITFOR_TASKHALT_EVENT', 'WAIT_XTP_HOST_WAIT', 'WAIT_XTP_OFFLINE_CKPT',
    'WAIT_XTP_CKPT_CLOSE', 'XE_DISPATCHER_WAIT', 'XE_LIVE_TARGET_TVF',
    'XE_TIMER_EVENT'
)
ORDER BY wait_time_ms DESC
"""

    # ── 5. 锁与事务 ────────────────────────────────────────────────
    QUERY_LOCKS = """
-- 当前锁统计
SELECT
    DB_NAME(resource_database_id) AS [database_name],
    resource_type,
    resource_description,
    request_mode AS [lock_mode],
    request_status AS [lock_status],
    request_owner_type,
    COUNT(*) AS [lock_count]
FROM sys.dm_tran_locks WITH (NOLOCK)
WHERE resource_database_id > 4
GROUP BY resource_database_id, resource_type, resource_description,
         request_mode, request_status, request_owner_type
ORDER BY lock_count DESC
"""

    QUERY_LOCK_DETAILS = """
-- 锁详情（TOP 100）
SELECT TOP 100
    l.request_session_id AS [spid],
    DB_NAME(l.resource_database_id) AS [database_name],
    l.resource_type,
    l.resource_description,
    l.request_mode AS [lock_mode],
    l.request_status AS [lock_status],
    s.login_name,
    s.host_name,
    s.program_name,
    s.status AS [session_status],
    t.transaction_id,
    t.transaction_begin_time,
    t.transaction_type_desc,
    CASE t.transaction_state
        WHEN 0 THEN 'Initializing'
        WHEN 1 THEN 'Initialized'
        WHEN 2 THEN 'Active'
        WHEN 3 THEN 'Ended'
        WHEN 4 THEN 'Commit Started'
        WHEN 5 THEN 'Prepared'
        WHEN 6 THEN 'Committed'
        WHEN 7 THEN 'Rolling Back'
        WHEN 8 THEN 'Rolled Back'
        ELSE 'Unknown'
    END AS [transaction_state]
FROM sys.dm_tran_locks l
INNER JOIN sys.dm_exec_sessions s ON l.request_session_id = s.session_id
LEFT JOIN sys.dm_tran_session_transactions t ON l.request_session_id = t.session_id
WHERE l.resource_database_id > 4
ORDER BY l.request_session_id
"""

    QUERY_DEADLOCK = """
-- 最近死锁信息（需要启用跟踪标志 1222 或 1224）
SELECT TOP 10
    XEvent.query('(event/data/value)[1]') AS [deadlock_xml],
    XEvent.timestamp
FROM (
    SELECT timestamp, object_name, CAST(target_data AS XML) AS target_data
    FROM sys.dm_xe_session_targets t
    INNER JOIN sys.dm_xe_sessions s ON t.event_session_address = s.address
    WHERE s.name = 'system_health'
    AND t.target_name = 'event_file'
) AS XEvent
CROSS APPLY target_data.nodes('//event[@name="xml_deadlock_report"]') AS Deadlock(XEvent)
ORDER BY XEvent.timestamp DESC
"""

    # ── 6. 备份与恢复 ──────────────────────────────────────────────
    QUERY_BACKUPS = """
-- 最近备份记录
SELECT TOP 20
    bs.database_name,
    bs.backup_start_date,
    bs.backup_finish_date,
    DATEDIFF(MINUTE, bs.backup_start_date, bs.backup_finish_date) AS [duration_min],
    CAST(bs.backup_size / 1048576.0 AS DECIMAL(10,2)) AS [backup_size_mb],
    CAST(bs.backup_size / 1048576.0 / NULLIF(DATEDIFF(MINUTE, bs.backup_start_date, bs.backup_finish_date), 0) AS DECIMAL(10,2)) AS [backup_speed_mbpm],
    bs.type AS [backup_type],
    CASE bs.type
        WHEN 'D' THEN 'Full'
        WHEN 'I' THEN 'Differential'
        WHEN 'L' THEN 'Log'
        WHEN 'F' THEN 'File'
        WHEN 'G' THEN 'File Differential'
        WHEN 'P' THEN 'Partial'
        WHEN 'Q' THEN 'Partial Differential'
        ELSE 'Unknown'
    END AS [backup_type_desc],
    bs.user_name AS [operator],
    bmf.physical_device_name,
    bs.server_name,
    bs.recovery_model,
    bs.is_copy_only,
    bs.is_password_protected,
    bs.encryption_enabled
FROM msdb.dbo.backupset bs
INNER JOIN msdb.dbo.backupmediafamily bmf ON bs.media_set_id = bmf.media_set_id
WITH (NOLOCK)
WHERE bs.backup_start_date > DATEADD(DAY, -30, GETDATE())
ORDER BY bs.backup_start_date DESC
"""

    QUERY_BACKUP_MISSING = """
-- 缺失备份的数据库（超过24小时未备份）
SELECT
    d.name AS [database_name],
    d.recovery_model_desc,
    d.state_desc,
    COALESCE(bs.last_backup_date, 'Never') AS [last_backup_date],
    CASE
        WHEN bs.last_backup_date IS NULL THEN 'No Backup'
        WHEN DATEDIFF(HOUR, bs.last_backup_date, GETDATE()) > 24 THEN 'Backup Overdue'
        ELSE 'OK'
    END AS [backup_status]
FROM sys.databases d
LEFT JOIN (
    SELECT database_name, MAX(backup_start_date) AS last_backup_date
    FROM msdb.dbo.backupset WITH (NOLOCK)
    GROUP BY database_name
) bs ON d.name = bs.database_name
WHERE d.database_id > 4
AND d.state = 0
AND (
    bs.last_backup_date IS NULL
    OR DATEDIFF(HOUR, bs.last_backup_date, GETDATE()) > 24
)
ORDER BY backup_status DESC, d.name
"""

    # ── 7. 性能指标 ────────────────────────────────────────────────
    QUERY_MEMORY_CLERKS = """
-- 内存使用统计
SELECT
    type AS [clerk_type],
    SUM(pages_kb) / 1024.0 AS [pages_mb],
    SUM(awe_allocated_kb) / 1024.0 AS [awe_mb],
    SUM(workflow_mem_kb) / 1024.0 AS [workflow_mb]
FROM sys.dm_os_memory_clerks WITH (NOLOCK)
WHERE pages_kb > 0
GROUP BY type
ORDER BY SUM(pages_kb) DESC
"""

    QUERY_BUFFER_POOL = """
-- 缓冲池统计
SELECT
    COUNT(*) * 8 / 1024.0 AS [total_buffer_mb],
    SUM(CASE WHEN is_modified = 1 THEN 1 ELSE 0 END) * 8 / 1024.0 AS [dirty_buffer_mb],
    SUM(CASE WHEN is_modified = 0 AND is_clean = 1 THEN 1 ELSE 0 END) * 8 / 1024.0 AS [clean_buffer_mb],
    SUM(CASE WHEN is_modified = 0 AND is_clean = 0 THEN 1 ELSE 0 END) * 8 / 1024.0 AS [free_buffer_mb],
    CAST(SUM(CASE WHEN is_modified = 1 THEN 1 ELSE 0 END) * 100.0 / NULLIF(COUNT(*), 0) AS DECIMAL(5,2)) AS [dirty_pct]
FROM sys.dm_os_buffer_descriptors WITH (NOLOCK)
"""

    QUERY_TOP_QUERIES = """
-- TOP 20 高开销查询
SELECT TOP 20
    qs.execution_count,
    qs.total_elapsed_time / 1000 AS [total_elapsed_sec],
    qs.total_elapsed_time / 1000 / NULLIF(qs.execution_count, 0) AS [avg_elapsed_sec],
    qs.total_logical_reads / 1000 AS [total_logical_reads_k],
    qs.total_logical_reads / 1000 / NULLIF(qs.execution_count, 0) AS [avg_logical_reads_k],
    qs.total_physical_reads / 1000 AS [total_physical_reads_k],
    qs.total_physical_reads / 1000 / NULLIF(qs.execution_count, 0) AS [avg_physical_reads_k],
    qs.total_worker_time / 1000 AS [total_cpu_ms],
    qs.total_worker_time / 1000 / NULLIF(qs.execution_count, 0) AS [avg_cpu_ms],
    SUBSTRING(st.text, 1, 500) AS [sql_text],
    DB_NAME(st.dbid) AS [database_name],
    OBJECT_NAME(st.objectid, st.dbid) AS [object_name]
FROM sys.dm_exec_query_stats qs
CROSS APPLY sys.dm_exec_sql_text(qs.sql_handle) st
WHERE st.dbid IS NOT NULL
ORDER BY qs.total_elapsed_time DESC
"""

    QUERY_INDEX_STATS = """
-- 缺失索引统计
SELECT TOP 20
    DB_NAME(database_id) AS [database_name],
    OBJECT_NAME(object_id, database_id) AS [table_name],
    equality_columns,
    inequality_columns,
    included_columns,
    unique_compiles,
    last_user_seek,
    avg_total_user_cost,
    avg_user_impact,
    user_seeks,
    user_scans
FROM sys.dm_db_missing_index_details mid
INNER JOIN sys.dm_db_missing_index_groups mig ON mid.index_handle = mig.index_handle
INNER JOIN sys.dm_db_missing_index_group_stats migs ON mig.index_group_handle = migs.group_handle
ORDER BY avg_total_user_cost * avg_user_impact DESC
"""

    # ── 8. 可用性/复制 ─────────────────────────────────────────────
    QUERY_AG = """
-- AlwaysOn 可用性组状态（如适用）
SELECT
    ag.name AS [ag_name],
    ag.group_id,
    ag.is_distributed,
    ar.replica_server_name,
    ar.role_desc,
    dbs.is_local,
    dbs.state_desc,
    ar.availability_mode_desc,
    ar.failover_mode_desc,
    ar.primary_recovery_health,
    ar.secondary_recovery_health
FROM sys.availability_groups ag
INNER JOIN sys.availability_replicas ar ON ag.group_id = ar.group_id
INNER JOIN sys.dm_hadr_availability_replica_states ars ON ar.replica_id = ars.replica_id
INNER JOIN sys.dm_hadr_availability_group_states dbs ON ag.group_id = dbs.group_id
ORDER BY ag.name, ar.replica_server_name
"""

    # ── 9. 错误日志摘要 ────────────────────────────────────────────
    QUERY_ERROR_LOG = """
-- 最近错误和警告（从当前日志）
WITH CTE_ErrorLog AS (
    SELECT
        LogDate,
        Severity,
        Message,
        ROW_NUMBER() OVER (ORDER BY LogDate DESC) AS rn
    FROM (
        SELECT CAST(LogDate AS DATETIME) AS LogDate, Severity, CAST(TextData AS NVARCHAR(4000)) AS Message
        FROM sys.messages WITH (NOLOCK)
        WHERE language_id = 1033
        UNION ALL
        SELECT GETDATE(), 0, 'Error log analysis requires xp_readerrorlog'
    ) AS ErrorLog
)
SELECT TOP 20
    LogDate,
    CASE Severity
        WHEN 0 THEN 'Info'
        WHEN 10 THEN 'Info'
        WHEN 11 THEN 'Warning'
        WHEN 12 THEN 'Warning'
        WHEN 13 THEN 'Error'
        WHEN 16 THEN 'Error'
        WHEN 17 THEN 'Error'
        WHEN 18 THEN 'Error'
        WHEN 19 THEN 'Error'
        WHEN 20 THEN 'Fatal'
        WHEN 21 THEN 'Fatal'
        WHEN 22 THEN 'Fatal'
        WHEN 23 THEN 'Fatal'
        WHEN 24 THEN 'Fatal'
        ELSE CAST(Severity AS VARCHAR)
    END AS [severity_desc],
    Message
FROM CTE_ErrorLog
WHERE LogDate > DATEADD(HOUR, -24, GETDATE())
AND (Severity >= 16 OR Message LIKE '%error%' OR Message LIKE '%fail%' OR Message LIKE '%deadlock%')
ORDER BY LogDate DESC
"""


# ═══════════════════════════════════════════════════════════════════════════
# Word 模板生成器
# ═══════════════════════════════════════════════════════════════════════════

class WordTemplateGeneratorSQLServer:
    """SQL Server 巡检报告 Word 模板生成器"""

    def __init__(self, data):
        self.data = data

    def _t(self, key):
        return _t(key)

    def _render_version_table(self, doc, version_data):
        """渲染版本信息表"""
        p = doc.add_heading(_t('sqlserver.chapter_version'), level=1)
        p.runs[0].font.size = Pt(14)

        if not version_data:
            doc.add_paragraph(_t('sqlserver.no_data'))
            return

        table = doc.add_table(rows=len(version_data[0]) if version_data else 1, cols=2)
        table.style = 'Table Grid'

        # 表头
        hdr_cells = table.rows[0].cells
        headers = [_t('sqlserver.config_name'), _t('sqlserver.config_value')]
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # 数据行
        for row_idx, row in enumerate(version_data):
            if row_idx == 0:
                continue  # 跳过表头行（如果有）
            cells = table.rows[row_idx].cells
            for col_idx, val in enumerate(row):
                cells[col_idx].text = str(val) if val else 'N/A'

    def _render_table(self, doc, data, headers, title):
        """通用表格渲染"""
        p = doc.add_heading(title, level=2)
        p.runs[0].font.size = Pt(12)

        if not data:
            doc.add_paragraph(_t('sqlserver.no_data'))
            return

        rows_count = min(len(data) + 1, 100)  # 限制最大100行
        table = doc.add_table(rows=rows_count, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = True

        # 表头
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # 数据行
        for row_idx, row in enumerate(data[:rows_count-1]):
            cells = table.rows[row_idx + 1].cells
            for col_idx, val in enumerate(row):
                if col_idx < len(cells):
                    cells[col_idx].text = str(val) if val is not None else ''

    def generate(self, output_path):
        """生成 Word 报告"""
        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # Logo 图片
        logo_path = os.path.join(os.path.dirname(__file__), 'dbcheck_logo.png')
        if os.path.exists(logo_path):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Cm(3.5))

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self._t('sqlserver.report_title'))
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 75, 135)

        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(_t('sqlserver.report_title'))
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(100, 100, 100)
        sub_run.font.italic = True

        # 装饰分隔线
        doc.add_paragraph()
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run('━' * 50)
        line_run.font.color.rgb = RGBColor(15, 75, 135)
        line_run.font.size = Pt(8)
        doc.add_paragraph()

        # 封面信息
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.autofit = False
        info_table.columns[0].width = Cm(3)
        info_table.columns[1].width = Cm(10)
        info = [
            [_t('sqlserver.label'), self.data.get('label', 'N/A')],
            [_t('sqlserver.inspector'), self.data.get('inspector', 'N/A')],
            [_t('sqlserver.server'), self.data.get('server', 'N/A')],
            [_t('sqlserver.report_time'), self.data.get('report_time', 'N/A')],
        ]
        for i, (k, v) in enumerate(info):
            info_table.rows[i].cells[0].text = k
            info_table.rows[i].cells[1].text = v

        doc.add_page_break()

        # 版本信息
        if self.data.get('version'):
            self._render_table(doc, self.data['version'],
                [_t('sqlserver.config_name'), _t('sqlserver.config_value')],
                _t('sqlserver.chapter_version'))

        # 数据库状态
        if self.data.get('databases'):
            headers = [_t('sqlserver.db_name'), _t('sqlserver.state'),
                      _t('sqlserver.recovery_model'), _t('sqlserver.data_size_mb'),
                      _t('sqlserver.log_size_mb')]
            self._render_table(doc, self.data['databases'], headers,
                _t('sqlserver.chapter_databases'))

        # 连接统计
        if self.data.get('connections'):
            self._render_table(doc, self.data['connections'],
                [_t('sqlserver.metric'), _t('sqlserver.value')],
                _t('sqlserver.chapter_connections'))

        # 活动会话
        if self.data.get('sessions'):
            headers = [_t('sqlserver.spid'), _t('sqlserver.login'),
                      _t('sqlserver.host'), _t('sqlserver.program'),
                      _t('sqlserver.status'), _t('sqlserver.cpu_ms')]
            self._render_table(doc, self.data['sessions'], headers,
                _t('sqlserver.chapter_sessions'))

        # 等待统计
        if self.data.get('wait_stats'):
            headers = [_t('sqlserver.wait_type'), _t('sqlserver.waiting_tasks'),
                      _t('sqlserver.wait_time_ms'), _t('sqlserver.wait_pct')]
            self._render_table(doc, self.data['wait_stats'], headers,
                _t('sqlserver.chapter_wait_stats'))

        # 锁信息
        if self.data.get('locks'):
            headers = [_t('sqlserver.db_name'), _t('sqlserver.lock_type'),
                      _t('sqlserver.lock_mode'), _t('sqlserver.lock_count')]
            self._render_table(doc, self.data['locks'], headers,
                _t('sqlserver.chapter_locks'))

        # 备份状态
        if self.data.get('backups'):
            headers = [_t('sqlserver.db_name'), _t('sqlserver.backup_date'),
                      _t('sqlserver.backup_type'), _t('sqlserver.size_mb'),
                      _t('sqlserver.duration_min')]
            self._render_table(doc, self.data['backups'], headers,
                _t('sqlserver.chapter_backups'))

        # 健康总结
        if self.data.get('summary'):
            p = doc.add_heading(_t('sqlserver.chapter_summary'), level=1)
            _render_markdown_to_doc(doc, self.data['summary'])

        doc.save(output_path)
        return output_path


# ═══════════════════════════════════════════════════════════════════════════
# SSH 远程系统信息采集器
# ═══════════════════════════════════════════════════════════════════════════

class RemoteSystemInfoCollector:
    """远程系统信息收集器 - 通过SSH连接获取远程主机信息（Windows/Linux）"""

    def __init__(self, host, port=22, username='root', password=None, key_file=None):
        """
        初始化远程系统信息收集器。

        :param host: 远程主机 IP 地址或主机名
        :param port: SSH 端口，默认 22
        :param username: SSH 登录用户名，默认 root
        :param password: SSH 登录密码（与 key_file 二选一）
        :param key_file: SSH 私钥文件路径（与 password 二选一）
        """
        self.host = host
        self.port = port
        self.username = username
        self.password = password
        self.key_file = key_file
        self.ssh_client = None

    def connect(self):
        """
        建立 SSH 连接。优先使用私钥文件认证，若无私钥则使用密码认证。
        自动接受远程主机密钥（AutoAddPolicy）。
        """
        try:
            import paramiko
            self.ssh_client = paramiko.SSHClient()
            self.ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            if self.key_file:
                private_key = paramiko.RSAKey.from_private_key_file(self.key_file)
                self.ssh_client.connect(hostname=self.host, port=self.port,
                                       username=self.username, pkey=private_key, timeout=10)
            else:
                self.ssh_client.connect(hostname=self.host, port=self.port,
                                       username=self.username, password=self.password, timeout=10)
            return True
        except Exception as e:
            print(_t("sqlserver_remote_ssh_fail") % (self.host, self.port, e))
            return False

    def disconnect(self):
        """断开 SSH 连接，释放资源。"""
        if self.ssh_client:
            self.ssh_client.close()

    def execute_command(self, command):
        """
        在远程主机上执行 Shell 命令。

        :return: (stdout 输出内容, stderr 错误内容) 的元组
        """
        try:
            stdin, stdout, stderr = self.ssh_client.exec_command(command, timeout=10)
            output = stdout.read().decode('utf-8').strip()
            error = stderr.read().decode('utf-8').strip()
            return output, error
        except Exception as e:
            print(_t("sqlserver_remote_cmd_fail") % (command, e))
            return "", str(e)

    def get_cpu_info(self):
        """通过远程命令采集 CPU 信息（支持 Linux 和 Windows）"""
        try:
            import platform
            sys_info = platform.system().lower()

            if sys_info == 'windows':
                # Windows: 使用 wmic
                cmd = 'wmic cpu get LoadPercentage /value'
                output, _ = self.execute_command(cmd)
                cpu_percent = 0.0
                if output:
                    for line in output.split('\n'):
                        if 'LoadPercentage' in line:
                            cpu_percent = float(line.split('=')[1].strip())
                cmd = 'wmic computersystem get NumberOfCores,NumberOfLogicalProcessors /value'
                output, _ = self.execute_command(cmd)
                physical_cores = logical_cores = 0
                if output:
                    for line in output.split('\n'):
                        if 'NumberOfCores' in line:
                            physical_cores = int(line.split('=')[1].strip())
                        elif 'NumberOfLogicalProcessors' in line:
                            logical_cores = int(line.split('=')[1].strip())
                return {
                    'usage_percent': cpu_percent,
                    'physical_cores': physical_cores,
                    'logical_cores': logical_cores,
                }
            else:
                # Linux: 使用 top 等命令
                cmd = "top -bn1 | grep 'Cpu(s)' | awk '{print $2}' | cut -d'%' -f1"
                output, _ = self.execute_command(cmd)
                cpu_percent = float(output) if output else 0.0
                cmd = "nproc"
                output, _ = self.execute_command(cmd)
                logical_cores = int(output) if output else 0
                cmd = "lscpu | grep 'Core(s) per socket' | awk '{print $4}'"
                output, _ = self.execute_command(cmd)
                cores_per_socket = int(output) if output else 1
                cmd = "lscpu | grep 'Socket(s)' | awk '{print $2}'"
                output, _ = self.execute_command(cmd)
                sockets = int(output) if output else 1
                physical_cores = cores_per_socket * sockets
                return {
                    'usage_percent': cpu_percent,
                    'physical_cores': physical_cores,
                    'logical_cores': logical_cores,
                }
        except Exception as e:
            print(_t("sqlserver_remote_cpu_fail") % e)
            return {}

    def get_memory_info(self):
        """通过远程命令采集内存信息（支持 Linux 和 Windows）"""
        try:
            import platform
            sys_info = platform.system().lower()

            if sys_info == 'windows':
                # Windows: 使用 wmic
                cmd = 'wmic OS get FreePhysicalMemory,TotalVisibleMemorySize /value'
                output, _ = self.execute_command(cmd)
                free_mb = total_mb = 0
                if output:
                    for line in output.split('\n'):
                        if 'FreePhysicalMemory' in line:
                            free_mb = int(line.split('=')[1].strip()) / 1024
                        elif 'TotalVisibleMemorySize' in line:
                            total_mb = int(line.split('=')[1].strip()) / 1024
                used_mb = total_mb - free_mb
                usage_percent = (used_mb / total_mb * 100) if total_mb > 0 else 0
                return {
                    'total_gb': round(total_mb / 1024, 2),
                    'available_gb': round(free_mb / 1024, 2),
                    'used_gb': round(used_mb / 1024, 2),
                    'usage_percent': round(usage_percent, 2),
                }
            else:
                # Linux: 使用 free -b
                cmd = "free -b | grep Mem"
                output, _ = self.execute_command(cmd)
                if output:
                    parts = output.split()
                    total_bytes = int(parts[1])
                    used_bytes = int(parts[2])
                    free_bytes = int(parts[3])
                    available_bytes = int(parts[6]) if len(parts) > 6 else free_bytes
                    usage_percent = (used_bytes / total_bytes) * 100 if total_bytes > 0 else 0
                    memory_info = {
                        'total_gb': round(total_bytes / (1024**3), 2),
                        'available_gb': round(available_bytes / (1024**3), 2),
                        'used_gb': round(used_bytes / (1024**3), 2),
                        'usage_percent': round(usage_percent, 2)
                    }
                    cmd = "free -b | grep Swap"
                    output, _ = self.execute_command(cmd)
                    if output:
                        parts = output.split()
                        swap_total = int(parts[1])
                        swap_used = int(parts[2])
                        swap_usage_percent = (swap_used / swap_total) * 100 if swap_total > 0 else 0
                        memory_info.update({
                            'swap_total_gb': round(swap_total / (1024**3), 2),
                            'swap_used_gb': round(swap_used / (1024**3), 2),
                            'swap_usage_percent': round(swap_usage_percent, 2)
                        })
                    return memory_info
                return {}
        except Exception as e:
            print(_t("sqlserver_remote_mem_fail") % e)
            return {}

    def get_disk_info(self):
        """通过远程命令采集磁盘信息（支持 Linux 和 Windows）"""
        try:
            import platform
            sys_info = platform.system().lower()
            disk_data = []

            if sys_info == 'windows':
                # Windows: 使用 wmic
                cmd = 'wmic logicaldisk get DeviceID,Size,FreeSpace,FileSystem /value'
                output, _ = self.execute_command(cmd)
                if output:
                    for line in output.split('\n'):
                        if 'DeviceID=' in line and 'DeviceID=.' not in line:
                            parts = line.strip().split('=')
                            if len(parts) >= 2:
                                device = parts[1].strip()
                                continue
                        if 'FreeSpace=' in line:
                            free_bytes = int(line.split('=')[1].strip())
                        elif 'Size=' in line:
                            total_bytes = int(line.split('=')[1].strip())
                        elif 'FileSystem=' in line:
                            fstype = line.split('=')[1].strip()
                            used_bytes = total_bytes - free_bytes
                            total_gb = round(total_bytes / (1024**3), 2)
                            used_gb = round(used_bytes / (1024**3), 2)
                            free_gb = round(free_bytes / (1024**3), 2)
                            usage_percent = (used_bytes / total_bytes * 100) if total_bytes > 0 else 0
                            disk_data.append({
                                'device': device,
                                'mountpoint': device,
                                'fstype': fstype,
                                'total_gb': total_gb,
                                'used_gb': used_gb,
                                'free_gb': free_gb,
                                'usage_percent': round(usage_percent, 2)
                            })
            else:
                # Linux: 使用 df -h
                IGNORE_PATTERN = "|".join([
                    "/mnt/iso", "/iso", "/media", "/run/media", "/cdrom",
                    "/mnt/iso/", "/mnt/media/", "/run/media/", "/iso/", "/cdrom/"
                ])
                ISO_FILTER = "mnt/iso|/iso|/media/|/run/media/|/cdrom"
                cmd = f"df -h | grep -vE 'tmpfs|devtmpfs' | grep -vE '{ISO_FILTER}' | tail -n +2"
                output, _ = self.execute_command(cmd)
                if output:
                    lines = output.strip().split('\n')
                    for line in lines:
                        parts = line.split()
                        if len(parts) >= 6:
                            device = parts[0]
                            mountpoint = parts[5]
                            if any(vfs in device for vfs in ['tmpfs', 'devtmpfs', 'overlay']):
                                continue
                            size_str = parts[1]
                            used_str = parts[2]
                            avail_str = parts[3]
                            usage_percent_str = parts[4].rstrip('%')

                            def to_gb(s):
                                if not s:
                                    return 0.0
                                s = s.strip().upper()
                                if s.endswith('G'):
                                    return round(float(s[:-1]), 2)
                                elif s.endswith('M'):
                                    return round(float(s[:-1]) / 1024, 2)
                                elif s.endswith('T'):
                                    return round(float(s[:-1]) * 1024, 2)
                                elif s.endswith('K'):
                                    return round(float(s[:-1]) / (1024**2), 2)
                                else:
                                    try:
                                        return round(float(s), 2)
                                    except:
                                        return 0.0

                            total_gb = to_gb(size_str)
                            used_gb = to_gb(used_str)
                            free_gb = to_gb(avail_str)
                            try:
                                usage_percent = float(usage_percent_str)
                            except:
                                usage_percent = 0.0

                            disk_data.append({
                                'device': device,
                                'mountpoint': mountpoint,
                                'fstype': "ext4",
                                'total_gb': total_gb,
                                'used_gb': used_gb,
                                'free_gb': free_gb,
                                'usage_percent': usage_percent
                            })
            return disk_data
        except Exception as e:
            print(_t("sqlserver_remote_disk_fail") % e)
            return []

    def get_system_info(self):
        """
        聚合采集远程主机的全部系统信息。

        :return: 包含系统信息的字典，字段：cpu、memory、disk_list、hostname、platform、boot_time
        """
        if not self.connect():
            return {}
        try:
            import platform
            system_info = {
                'cpu': {},
                'memory': {},
                'disk_list': [],
                'hostname': '',
                'platform': '',
                'boot_time': '',
            }

            system_info['cpu'] = self.get_cpu_info()
            system_info['memory'] = self.get_memory_info()
            system_info['disk_list'] = self.get_disk_info()

            sys_info = platform.system().lower()
            if sys_info == 'windows':
                cmd = "hostname"
                output, _ = self.execute_command(cmd)
                if output:
                    system_info['hostname'] = output.strip()
                cmd = "systeminfo | findstr /B /C:\"OS Name\" /C:\"OS Version\""
                output, _ = self.execute_command(cmd)
                if output:
                    system_info['platform'] = output.strip()
            else:
                cmd = "hostname"
                output, _ = self.execute_command(cmd)
                if output:
                    system_info['hostname'] = output.strip()
                cmd = "uname -a"
                output, _ = self.execute_command(cmd)
                if output:
                    system_info['platform'] = output.strip()
                cmd = "who -b | awk '{print $3 \" \" $4}'"
                output, _ = self.execute_command(cmd)
                if output:
                    system_info['boot_time'] = output.strip()

            return system_info
        finally:
            self.disconnect()


class LocalSystemInfoCollector:
    """本地系统信息收集器 - 使用 psutil 库采集当前主机系统信息"""

    def __init__(self):
        pass

    def get_cpu_info(self):
        """采集本机 CPU 信息。"""
        try:
            cpu_percent = psutil.cpu_percent(interval=1)
            return {
                'usage_percent': cpu_percent,
                'physical_cores': psutil.cpu_count(logical=False) or 0,
                'logical_cores': psutil.cpu_count(logical=True) or 0,
            }
        except Exception:
            return {}

    def get_memory_info(self):
        """采集本机内存信息。"""
        try:
            mem = psutil.virtual_memory()
            return {
                'total_gb': round(mem.total / (1024**3), 2),
                'available_gb': round(mem.available / (1024**3), 2),
                'used_gb': round(mem.used / (1024**3), 2),
                'usage_percent': mem.percent,
            }
        except Exception:
            return {}

    def get_disk_info(self):
        """采集本机磁盘信息。"""
        try:
            disk_data = []
            partitions = psutil.disk_partitions()
            for partition in partitions:
                try:
                    usage = psutil.disk_usage(partition.mountpoint)
                    # 跳过虚拟文件系统
                    if partition.fstype in ('tmpfs', 'devtmpfs', 'overlay', 'none'):
                        continue
                    # 跳过忽略的挂载点
                    if any(partition.mountpoint.startswith(ign) for ign in IGNORE_MOUNTS):
                        continue
                    disk_data.append({
                        'device': partition.device,
                        'mountpoint': partition.mountpoint,
                        'fstype': partition.fstype,
                        'total_gb': round(usage.total / (1024**3), 2),
                        'used_gb': round(usage.used / (1024**3), 2),
                        'free_gb': round(usage.free / (1024**3), 2),
                        'usage_percent': usage.percent
                    })
                except Exception:
                    continue
            return disk_data
        except Exception:
            return []

    def get_system_info(self):
        """聚合采集本机全部系统信息。"""
        import platform as plat
        return {
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk_list': self.get_disk_info(),
            'hostname': plat.node(),
            'platform': f"{plat.system()} {plat.release()}",
            'boot_time': '',
        }


# ═══════════════════════════════════════════════════════════════════════════
# 主巡检类
# ═══════════════════════════════════════════════════════════════════════════

class DBCheckSQLServer:
    """SQL Server 数据库健康巡检主类"""

    def __init__(self, host, port, user, password, database=None, label=None,
                 inspector=None, ssh_host=None, ssh_user=None, ssh_password=None,
                 ssh_key_file=None, desensitize=False):
        self.host = host
        self.port = port or 1433
        self.user = user
        self.password = password
        self.database = database
        self.label = label or 'SQLServer'
        self.inspector = inspector or 'DBA'
        self.ssh_host = ssh_host
        self.ssh_user = ssh_user
        self.ssh_password = ssh_password
        self.ssh_key_file = ssh_key_file
        self.ssh_port = 22  # 默认 SSH 端口
        self.desensitize = desensitize

        self.conn = None
        self.cursor = None
        self.data = {}
        self.report_path = None
        self.ssh_client = None

        # 初始化语言
        self._lang = _SQLSERVER_LANG

    def _t(self, key):
        return _t(key)

    def _connect(self):
        """建立数据库连接"""
        import pyodbc

        conn_str = (
            f"DRIVER={{ODBC Driver 17 for SQL Server}};"
            f"SERVER={self.host},{self.port};"
            f"UID={self.user};"
            f"PWD={self.password};"
            f"TrustServerCertificate=yes;"
            f"Encrypt=yes;"
        )

        if self.database:
            conn_str += f"Database={self.database};"

        try:
            self.conn = pyodbc.connect(conn_str)
            self.cursor = self.conn.cursor()
            print(_t('sqlserver.conn_success'))
            return True
        except Exception as e:
            print(_t('sqlserver.conn_failed') % str(e))
            return False

    def _disconnect(self):
        """关闭数据库连接"""
        if self.cursor:
            self.cursor.close()
        if self.conn:
            self.conn.close()

    def _execute_query(self, query):
        """执行查询并返回结果"""
        try:
            self.cursor.execute(query)
            rows = self.cursor.fetchall()
            columns = [col[0] for col in self.cursor.description] if self.cursor.description else []
            return [dict(zip(columns, row)) for row in rows]
        except Exception as e:
            print(f"Query error: {e}")
            return []

    def _get_host_resources(self):
        """获取主机资源 - 支持 SSH 远程和本地采集"""
        try:
            # 如果配置了 SSH，使用远程采集
            if self.ssh_host and (self.ssh_password or self.ssh_key_file):
                print(_t('sqlserver.ssh_collect_start') % self.ssh_host)
                collector = RemoteSystemInfoCollector(
                    host=self.ssh_host,
                    port=int(self.ssh_port) if self.ssh_port else 22,
                    username=self.ssh_user or 'root',
                    password=self.ssh_password,
                    key_file=self.ssh_key_file
                )
                sys_info = collector.get_system_info()
                if sys_info:
                    print(_t('sqlserver.ssh_collect_ok'))
                    return {
                        'cpu_percent': sys_info.get('cpu', {}).get('usage_percent', 0),
                        'memory_total_gb': sys_info.get('memory', {}).get('total_gb', 0),
                        'memory_used_gb': sys_info.get('memory', {}).get('used_gb', 0),
                        'memory_percent': sys_info.get('memory', {}).get('usage_percent', 0),
                        'disk_total_gb': sys_info.get('disk_list', [{}])[0].get('total_gb', 0) if sys_info.get('disk_list') else 0,
                        'disk_used_gb': sys_info.get('disk_list', [{}])[0].get('used_gb', 0) if sys_info.get('disk_list') else 0,
                        'disk_percent': sys_info.get('disk_list', [{}])[0].get('usage_percent', 0) if sys_info.get('disk_list') else 0,
                        'system_info': sys_info,
                    }
                else:
                    print(_t('sqlserver.ssh_collect_fail'))
            # 本地采集
            collector = LocalSystemInfoCollector()
            sys_info = collector.get_system_info()
            return {
                'cpu_percent': sys_info.get('cpu', {}).get('usage_percent', 0),
                'memory_total_gb': sys_info.get('memory', {}).get('total_gb', 0),
                'memory_used_gb': sys_info.get('memory', {}).get('used_gb', 0),
                'memory_percent': sys_info.get('memory', {}).get('usage_percent', 0),
                'disk_total_gb': sys_info.get('disk_list', [{}])[0].get('total_gb', 0) if sys_info.get('disk_list') else 0,
                'disk_used_gb': sys_info.get('disk_list', [{}])[0].get('used_gb', 0) if sys_info.get('disk_list') else 0,
                'disk_percent': sys_info.get('disk_list', [{}])[0].get('usage_percent', 0) if sys_info.get('disk_list') else 0,
                'system_info': sys_info,
            }
        except Exception as e:
            return {'error': str(e)}

    def getData(self):
        """采集所有巡检数据"""
        if not self._connect():
            return False

        print(_t('sqlserver.start_collection'))

        try:
            # 1. 版本信息
            print(_t('sqlserver.collect_version'))
            version_data = self._execute_query(SQLServerQueries.QUERY_VERSION)
            self.data['version'] = version_data

            # 2. 服务器配置
            print(_t('sqlserver.collect_config'))
            config_data = self._execute_query(SQLServerQueries.QUERY_SERVER_CONFIG)
            self.data['config'] = config_data

            # 3. 数据库状态
            print(_t('sqlserver.collect_databases'))
            db_data = self._execute_query(SQLServerQueries.QUERY_DATABASES)
            self.data['databases'] = db_data

            # 4. 连接统计
            print(_t('sqlserver.collect_connections'))
            conn_data = self._execute_query(SQLServerQueries.QUERY_CONNECTIONS)
            self.data['connections'] = conn_data

            # 5. 活动会话
            print(_t('sqlserver.collect_sessions'))
            session_data = self._execute_query(SQLServerQueries.QUERY_ACTIVE_SESSIONS)
            self.data['sessions'] = session_data

            # 6. 阻塞信息
            print(_t('sqlserver.collect_blocking'))
            blocking_data = self._execute_query(SQLServerQueries.QUERY_BLOCKING)
            self.data['blocking'] = blocking_data

            # 7. 等待统计
            print(_t('sqlserver.collect_wait_stats'))
            wait_data = self._execute_query(SQLServerQueries.QUERY_WAIT_STATS)
            self.data['wait_stats'] = wait_data

            # 8. 锁统计
            print(_t('sqlserver.collect_locks'))
            lock_data = self._execute_query(SQLServerQueries.QUERY_LOCKS)
            self.data['locks'] = lock_data

            # 9. 备份信息
            print(_t('sqlserver.collect_backups'))
            backup_data = self._execute_query(SQLServerQueries.QUERY_BACKUPS)
            self.data['backups'] = backup_data

            # 10. 缺失备份
            print(_t('sqlserver.collect_backup_missing'))
            backup_missing = self._execute_query(SQLServerQueries.QUERY_BACKUP_MISSING)
            self.data['backup_missing'] = backup_missing

            # 11. 内存统计
            print(_t('sqlserver.collect_memory'))
            memory_data = self._execute_query(SQLServerQueries.QUERY_MEMORY_CLERKS)
            self.data['memory'] = memory_data

            # 12. 缓冲池
            print(_t('sqlserver.collect_buffer'))
            buffer_data = self._execute_query(SQLServerQueries.QUERY_BUFFER_POOL)
            self.data['buffer'] = buffer_data

            # 13. TOP 查询
            print(_t('sqlserver.collect_top_queries'))
            top_queries = self._execute_query(SQLServerQueries.QUERY_TOP_QUERIES)
            self.data['top_queries'] = top_queries

            # 14. 主机资源
            print(_t('sqlserver.collect_host'))
            self.data['host'] = self._get_host_resources()

            # 15. 巡检信息
            self.data['label'] = self.label
            self.data['inspector'] = self.inspector
            self.data['server'] = f"{self.host}:{self.port}"
            self.data['report_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

            print(_t('sqlserver.collection_complete'))
            return True

        except Exception as e:
            print(_t('sqlserver.collection_failed') % str(e))
            return False
        finally:
            self._disconnect()

    def _analyze(self):
        """智能分析巡检数据"""
        issues = []

        # 检查等待统计
        if self.data.get('wait_stats'):
            for wait in self.data['wait_stats'][:5]:
                wait_type = wait.get('wait_type', '')
                wait_time = wait.get('wait_time_ms', 0)

                if 'CXPACKET' in wait_type and wait_time > 10000:
                    issues.append({
                        'severity': 'medium',
                        'category': _t('sqlserver.wait_stats'),
                        'issue': f"CXPACKET wait: {wait_time} ms",
                        'suggestion': _t('sqlserver.suggest_cxpacket')
                    })

                if 'PAGEIOLATCH' in wait_type and wait_time > 5000:
                    issues.append({
                        'severity': 'high',
                        'category': _t('sqlserver.wait_stats'),
                        'issue': f"PAGEIOLATCH wait: {wait_time} ms",
                        'suggestion': _t('sqlserver.suggest_pageiolatch')
                    })

                if 'LOCK' in wait_type and wait_time > 5000:
                    issues.append({
                        'severity': 'high',
                        'category': _t('sqlserver.locks'),
                        'issue': f"Lock wait: {wait_time} ms",
                        'suggestion': _t('sqlserver.suggest_lock')
                    })

        # 检查阻塞
        if self.data.get('blocking'):
            issues.append({
                'severity': 'high',
                'category': _t('sqlserver.blocking'),
                'issue': f"{len(self.data['blocking'])} blocking sessions detected",
                'suggestion': _t('sqlserver.suggest_blocking')
            })

        # 检查备份缺失
        if self.data.get('backup_missing'):
            issues.append({
                'severity': 'high',
                'category': _t('sqlserver.backups'),
                'issue': f"{len(self.data['backup_missing'])} databases missing backup",
                'suggestion': _t('sqlserver.suggest_backup')
            })

        # 检查连接数
        if self.data.get('connections'):
            conn = self.data['connections'][0] if self.data['connections'] else {}
            usage_pct = conn.get('connection_usage_pct', 0)
            if usage_pct > 80:
                issues.append({
                    'severity': 'medium',
                    'category': _t('sqlserver.connections'),
                    'issue': f"Connection usage: {usage_pct}%",
                    'suggestion': _t('sqlserver.suggest_connections')
                })

        # 检查内存
        if self.data.get('host'):
            mem_pct = self.data['host'].get('memory_percent', 0)
            if mem_pct > 90:
                issues.append({
                    'severity': 'high',
                    'category': _t('sqlserver.memory'),
                    'issue': f"Memory usage: {mem_pct}%",
                    'suggestion': _t('sqlserver.suggest_memory')
                })

        self.data['issues'] = issues
        return issues

    def _generate_summary(self):
        """生成健康总结"""
        issues = self.data.get('issues', [])

        high_count = sum(1 for i in issues if i.get('severity') == 'high')
        medium_count = sum(1 for i in issues if i.get('severity') == 'medium')
        low_count = sum(1 for i in issues if i.get('severity') == 'low')

        if high_count == 0 and medium_count == 0:
            status = _t('sqlserver.status_healthy')
        elif high_count > 0:
            status = _t('sqlserver.status_critical')
        else:
            status = _t('sqlserver.status_warning')

        summary = f"""## {self._t('sqlserver.summary_title')}

### {self._t('sqlserver.health_status')}: {status}

### {self._t('sqlserver.issue_summary')}:
- **{self._t('sqlserver.high')}**: {high_count}
- **{self._t('sqlserver.medium')}**: {medium_count}
- **{self._t('sqlserver.low')}**: {low_count}

### {self._t('sqlserver.recommendations')}:
"""

        if issues:
            for idx, issue in enumerate(issues[:10], 1):
                summary += f"{idx}. [{issue.get('severity', 'info').upper()}] {issue.get('category', '')}: {issue.get('issue', '')}\n"
                summary += f"   - {issue.get('suggestion', '')}\n"
        else:
            summary += _t('sqlserver.no_issues')

        self.data['summary'] = summary
        return summary

    def checkdb(self):
        """执行完整巡检流程"""
        print(f"\n{'='*60}")
        print(f"{self._t('sqlserver.starting_inspection')}: {self.label}")
        print(f"{'='*60}\n")

        # 1. 采集数据
        if not self.getData():
            return False

        # 2. 智能分析
        print(_t('sqlserver.running_analysis'))
        issues = self._analyze()

        # 3. AI 诊断
        self.data['ai_advice'] = ''
        try:
            from analyzer import AIAdvisor
            import json as _json
            cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ai_config.json')
            ai_cfg = {}
            if os.path.exists(cfg_path):
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    ai_cfg = _json.load(f)
            advisor = AIAdvisor(
                backend=ai_cfg.get('backend'),
                api_key=ai_cfg.get('api_key'),
                api_url=ai_cfg.get('api_url'),
                model=ai_cfg.get('model')
            )
            if advisor.enabled:
                print(("\n🤖 " + _t('sqlserver.ai_calling')) % (advisor.backend, advisor.model))
                ai_advice = advisor.diagnose('sqlserver', self.label, self.data, issues, lang=self._lang)
                self.data['ai_advice'] = ai_advice
        except Exception as e:
            print(f"AI 诊断异常: {e}")

        # ── 慢查询深度分析（P2）──────────────────────────────
        self.data['slow_query_result'] = None
        try:
            from slow_query_analyzer import SQLServerSlowQueryAnalyzer
            if self.conn:
                analyzer = SQLServerSlowQueryAnalyzer()
                ai_advisor = None
                try:
                    from analyzer import AIAdvisor
                    import json as _json
                    cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ai_config.json')
                    ai_cfg = {}
                    if os.path.exists(cfg_path):
                        with open(cfg_path, 'r', encoding='utf-8') as f:
                            ai_cfg = _json.load(f)
                    ai_advisor = AIAdvisor(
                        backend=ai_cfg.get('backend'),
                        api_key=ai_cfg.get('api_key'),
                        api_url=ai_cfg.get('api_url'),
                        model=ai_cfg.get('model')
                    )
                except Exception:
                    pass
                print("\n\U0001f50d " + _t('sqlserver.slow_query_analyzing'))
                result = analyzer.analyze(self.conn, ai_advisor=ai_advisor, lang=self._lang)
                self.data['slow_query_result'] = result.to_dict()
                if result.is_empty():
                    print("  \u2139\ufe0f  " + _t('sqlserver.slow_query_unavailable'))
                else:
                    print("  \u2705  " + (_t('sqlserver.slow_query_ok') % len(result.top_sql_by_latency)))
        except ImportError:
            pass
        except Exception as e:
            print("\u26a0\ufe0f 慢查询深度分析失败: %s" % e)

        # 4. 生成总结
        self._generate_summary()

        # 5. 生成报告
        print(_t('sqlserver.generating_report'))
        self._save_report()

        print(f"\n{'='*60}")
        print(f"{self._t('sqlserver.inspection_complete')}")
        print(f"{self._t('sqlserver.report_saved')}: {self.report_path}")
        print(f"{'='*60}\n")

        return True

    def _save_report(self):
        """保存报告"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_label = re.sub(r'[^\w\u4e00-\u9fa5]', '_', self.label)

        # 创建报告目录
        reports_dir = os.path.join(os.path.dirname(__file__), 'reports', 'sqlserver')
        os.makedirs(reports_dir, exist_ok=True)

        filename = f"SQLServer_{safe_label}_{timestamp}.docx"
        self.report_path = os.path.join(reports_dir, filename)

        # ── 脱敏处理（如开启了脱敏导出）───────────────────────
        report_data = self.data
        if self.desensitize:
            try:
                from desensitize import apply_desensitization
                report_data = apply_desensitization(dict(self.data))
            except Exception:
                pass

        generator = WordTemplateGeneratorSQLServer(report_data)
        generator.generate(self.report_path)


# ═══════════════════════════════════════════════════════════════════════════
# CLI 主函数
# ═══════════════════════════════════════════════════════════════════════════

def show_main_menu():
    """显示程序主菜单"""
    print("\n" + "=" * 60)
    print("            " + _t("sqlserver_cli_banner") + " " + VER)
    print("=" * 60)
    print(_t("sqlserver_cli_menu_item1"))
    print(_t("sqlserver_cli_menu_item2"))
    print(_t("sqlserver_cli_menu_item3"))
    print(_t("sqlserver_cli_menu_item4"))
    print("=" * 60)
    while True:
        choice = input(_t("sqlserver_cli_choose_prompt")).strip()
        if choice in ['1', '2', '3', '4']:
            return choice
        else:
            print("\u274c " + _t("sqlserver_cli_invalid_choice"))


def input_db_info():
    """交互式输入数据库连接信息"""
    print("\n" + _t("sqlserver_cli_db_info_title"))
    host = input(_t("cli_db_host").format(default="localhost")).strip() or "localhost"
    port_input = input(_t("cli_db_port").format(default=1433)).strip()
    if not port_input:
        port = 1433
    else:
        try:
            port = int(port_input)
        except ValueError:
            print(_t("cli_db_port_invalid").format(default=1433))
            port = 1433
    user = input(_t("cli_db_user").format(default="sa")).strip() or "sa"
    import getpass
    password = getpass.getpass(_t("cli_db_password")).strip()
    db_name = input(_t("cli_db_name").format(default="master")).strip() or "master"

    print("\n" + _t("cli_ssh_config_title"))
    print(_t("cli_ssh_config_note"))
    enable_ssh = input(_t("cli_ssh_enable")).strip().lower()
    ssh_info = {}
    if enable_ssh in ['y', 'yes']:
        ssh_host = input(_t("cli_ssh_host").format(default=host)).strip() or host
        ssh_port_input = input(_t("cli_ssh_port").format(default=22)).strip()
        if not ssh_port_input:
            ssh_port = 22
        else:
            try:
                ssh_port = int(ssh_port_input)
            except ValueError:
                print(_t("cli_ssh_port_invalid").format(default=22))
                ssh_port = 22
        ssh_user = input(_t("cli_ssh_user").format(default="root")).strip() or "root"
        auth_choice = input(_t("cli_ssh_auth_method")).strip()
        if auth_choice == '2':
            ssh_key_file = input(_t("cli_ssh_key_path")).strip()
            ssh_password = ""
            if not os.path.exists(ssh_key_file):
                print(_t("cli_ssh_key_not_exist").format(path=ssh_key_file))
                retry = input(_t("cli_retry_yes")).strip().lower()
                if retry in ['', 'y', 'yes']:
                    return input_db_info()
                else:
                    ssh_key_file = ""
        else:
            ssh_password = getpass.getpass(_t("cli_ssh_password")).strip()
            ssh_key_file = ""
        ssh_info = {
            'ssh_host': ssh_host,
            'ssh_port': ssh_port,
            'ssh_user': ssh_user,
            'ssh_password': ssh_password,
            'ssh_key_file': ssh_key_file
        }

    # 验证连接
    print("\n\U0001f50d " + _t("sqlserver_cli_verifying_sqlserver").format(host=host, port=port))
    try:
        import pyodbc
        conn_str = (
            f"DRIVER={{ODBC Driver 17 for SQL Server}};"
            f"SERVER={host},{port};"
            f"UID={user};"
            f"PWD={password};"
            f"TrustServerCertificate=yes;"
            f"Encrypt=yes;"
        )
        if db_name:
            conn_str += f"Database={db_name};"
        conn = pyodbc.connect(conn_str, timeout=10)
        conn.close()
        print(_t("sqlserver_cli_sqlserver_success").format(host=host, port=port))
    except Exception as e:
        print(_t("sqlserver_cli_sqlserver_fail").format(e=e))
        retry = input(_t("cli_retry_no")).strip().lower()
        if retry == 'y':
            return input_db_info()
        else:
            return None

    # 验证 SSH
    if ssh_info:
        print("\U0001f50d " + _t("pg_cli_verifying_ssh").format(host=ssh_info['ssh_host'], port=ssh_info['ssh_port']))
        try:
            import paramiko
            ssh_client = paramiko.SSHClient()
            ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            if ssh_info['ssh_key_file']:
                ssh_client.connect(ssh_info['ssh_host'], ssh_info['ssh_port'], ssh_info['ssh_user'],
                                   key_filename=ssh_info['ssh_key_file'], timeout=10)
            else:
                ssh_client.connect(ssh_info['ssh_host'], ssh_info['ssh_port'], ssh_info['ssh_user'],
                                   password=ssh_info['ssh_password'], timeout=10)
            ssh_client.close()
            print(_t("pg_cli_ssh_success"))
        except Exception as e:
            print(_t("pg_cli_ssh_fail").format(e=e))
            retry = input(_t("cli_retry_no")).strip().lower()
            if retry == 'y':
                return input_db_info()
            else:
                ssh_info = {}

    inspector_name = input(_t("sqlserver_cli_inspector_prompt")).strip() or "DBA"

    return {
        'host': host,
        'port': port,
        'user': user,
        'password': password,
        'database': db_name,
        'label': db_name,
        'inspector': inspector_name,
        'ssh_info': ssh_info
    }


def single_inspection():
    """单机巡检"""
    print("\n=== " + _t("sqlserver_cli_single_mode") + " ===")
    db_info = input_db_info()
    if not db_info:
        return False

    ssh_info = db_info.pop('ssh_info', {})
    inspector = DBCheckSQLServer(
        host=db_info['host'],
        port=db_info['port'],
        user=db_info['user'],
        password=db_info['password'],
        database=db_info.get('database'),
        label=db_info.get('label'),
        inspector=db_info.get('inspector'),
        ssh_host=ssh_info.get('ssh_host'),
        ssh_user=ssh_info.get('ssh_user'),
        ssh_password=ssh_info.get('ssh_password'),
        ssh_key_file=ssh_info.get('ssh_key_file')
    )

    return inspector.checkdb()


def batch_inspection():
    """批量巡检"""
    print("\n=== " + _t("sqlserver_cli_batch_mode") + " ===")
    excel_manager = ExcelTemplateManager()
    if not os.path.exists(excel_manager.template_file):
        print("\u274c " + _t("sqlserver_cli_excel_not_exist"))
        create_template = input(_t("sqlserver_cli_create_template_now")).strip().lower()
        if create_template in ['', 'y', 'yes']:
            excel_manager.create_template()
        return
    db_list = excel_manager.read_template()
    if not db_list:
        return
    print("\n\U0001f4cb " + _t("sqlserver_cli_will_inspect_n").format(n=len(db_list)))
    for i, db in enumerate(db_list, 1):
        ssh_suffix = " " + _t("cli_ssh_suffix") if db.get("ssh_host") and (db.get("ssh_password") or db.get("ssh_key_file")) else ""
        print("  " + str(i) + ". " + db["name"] + " - " + db["host"] + ":" + str(db.get("port", 1433)) + ssh_suffix)
    confirm = input("\n" + _t("sqlserver_cli_confirm_batch")).strip().lower()
    if confirm in ['', 'y', 'yes']:
        total_dbs = len(db_list)
        success_count = 0
        for i, db_info in enumerate(db_list, 1):
            print("\n[" + str(i) + "/" + str(total_dbs) + "] " + _t("sqlserver_cli_start_inspect_n").format(name=db_info["name"]))
            ssh_info = {}
            if db_info.get("ssh_host"):
                ssh_info = {
                    'ssh_host': db_info.get('ssh_host'),
                    'ssh_port': db_info.get('ssh_port', 22),
                    'ssh_user': db_info.get('ssh_user', 'root'),
                    'ssh_password': db_info.get('ssh_password', ''),
                    'ssh_key_file': db_info.get('ssh_key_file', '')
                }
            inspector = DBCheckSQLServer(
                host=db_info['host'],
                port=db_info.get('port', 1433),
                user=db_info['user'],
                password=db_info['password'],
                database=db_info.get('database'),
                label=db_info.get('name'),
                inspector=db_info.get('inspector', 'DBA'),
                ssh_host=ssh_info.get('ssh_host'),
                ssh_user=ssh_info.get('ssh_user'),
                ssh_password=ssh_info.get('ssh_password'),
                ssh_key_file=ssh_info.get('ssh_key_file')
            )
            if inspector.checkdb():
                success_count += 1
        print("\n=== " + _t("sqlserver_cli_batch_done") + " ===")
        print(_t("sqlserver_cli_success_count").format(s=success_count, t=total_dbs))


def create_excel_template():
    """创建Excel配置模板"""
    print("\n=== " + _t("pg_cli_create_excel") + " ===")
    excel_manager = ExcelTemplateManager()
    excel_manager.create_template()


def main():
    """程序主入口"""
    import time
    start_time = time.time()

    # 支持从主入口通过 --template 直接生成 Excel 模板
    if len(sys.argv) > 1 and sys.argv[1] == '--template':
        create_excel_template()
        return

    print_banner()
    while True:
        choice = show_main_menu()
        if choice == '1':
            single_inspection()
        elif choice == '2':
            batch_inspection()
        elif choice == '3':
            create_excel_template()
        elif choice == '4':
            print("\n" + _t("sqlserver_cli_thanks"))
            break
        if choice != '4':
            continue_choice = input("\n" + _t("sqlserver_cli_back_menu")).strip().lower()
            if continue_choice in ['', 'y', 'yes']:
                continue
            else:
                print("\n" + _t("sqlserver_cli_thanks"))
                break
    end_time = time.time()
    print("\n" + _t("sqlserver_cli_total_time").format(t=end_time - start_time))


if __name__ == '__main__':
    main()
