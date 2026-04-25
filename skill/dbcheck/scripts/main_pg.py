#!/usr/bin/env python3
# -*- coding:utf-8 -*-
from version import __version__ as VER

"""
PostgreSQL 数据库自动化健康巡检工具 {VER}
依赖: psycopg2, python-docx, docxtpl, openpyxl, psutil, paramiko
"""
import warnings
warnings.filterwarnings("ignore")
import itertools
import math
import sys
import datetime
import argparse
import subprocess
import logging
import logging.handlers
import socket
import re
import time
from pathlib import Path
import sys, getopt, os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate
import configparser
import importlib
import subprocess
import json
import hashlib
import base64
from datetime import datetime, timedelta
import platform
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import tempfile
import io
import psutil
import shutil
import paramiko

try:
    import psycopg2
    import psycopg2.extras
except ImportError:
    print("缺少 psycopg2 库，请执行 pip install psycopg2-binary")
    sys.exit(1)

importlib.reload(sys)

# ── i18n setup for CLI ─────────────────────────────────────────────
try:
    from i18n import get_lang
    _PG_LANG = get_lang()
except Exception:
    _PG_LANG = 'zh'

def _t(key):
    try:
        from i18n import t as _tt
        return _tt(key, _PG_LANG)
    except Exception:
        return key

# ============================================================
# 内置 PostgreSQL 巡检 SQL 模板
# ============================================================
PG_SQL_TEMPLATES_CONTENT = """
[report]
name = PostgreSQL HealthCheck Report
template = ./templates/pg_wordtemplates_v1.0.docx
output = /tmp/PGCheckReport.docx

[variables]
pg_version       = SELECT version() AS version;
pg_uptime        = SELECT now() - pg_postmaster_start_time() AS uptime, pg_postmaster_start_time() AS start_time;
pg_connections   = SELECT count(*) AS total_connections, max_conn.setting::int AS max_connections, round(count(*)*100.0/max_conn.setting::int,2) AS usage_percent FROM pg_stat_activity, (SELECT setting FROM pg_settings WHERE name='max_connections') max_conn GROUP BY max_conn.setting;
pg_conn_detail   = SELECT state, count(*) AS count FROM pg_stat_activity WHERE state IS NOT NULL GROUP BY state ORDER BY count DESC;
pg_wait_events   = SELECT wait_event_type, wait_event, count(*) AS count FROM pg_stat_activity WHERE wait_event IS NOT NULL GROUP BY wait_event_type, wait_event ORDER BY count DESC LIMIT 10;
pg_long_queries  = SELECT pid, now()-query_start AS duration, state, left(query,120) AS query FROM pg_stat_activity WHERE state NOT IN ('idle') AND query_start IS NOT NULL AND now()-query_start > interval '30 seconds' ORDER BY duration DESC LIMIT 10;
pg_lock_info     = SELECT count(*) AS total_locks, sum(CASE WHEN granted THEN 1 ELSE 0 END) AS granted_locks, sum(CASE WHEN NOT granted THEN 1 ELSE 0 END) AS waiting_locks FROM pg_locks;
pg_db_size       = SELECT datname AS database_name, pg_size_pretty(pg_database_size(datname)) AS size, pg_database_size(datname) AS size_bytes FROM pg_database WHERE datistemplate=false ORDER BY size_bytes DESC;
pg_table_stats   = SELECT schemaname, relname AS tablename, n_live_tup AS live_rows, n_dead_tup AS dead_rows, round(n_dead_tup*100.0/NULLIF(n_live_tup+n_dead_tup,0),2) AS dead_ratio, last_vacuum, last_autovacuum, last_analyze, last_autoanalyze FROM pg_stat_user_tables ORDER BY n_dead_tup DESC LIMIT 15;
pg_index_usage   = SELECT schemaname, relname AS tablename, indexrelname AS indexname, idx_scan, idx_tup_read, idx_tup_fetch FROM pg_stat_user_indexes ORDER BY idx_scan ASC LIMIT 15;
pg_replication   = SELECT pid, usename, application_name, client_addr, state, sent_lsn, write_lsn, flush_lsn, replay_lsn, sync_state FROM pg_stat_replication;
pg_cache_hit     = SELECT datname, blks_read, blks_hit, round(blks_hit*100.0/NULLIF(blks_hit+blks_read,0),2) AS cache_hit_ratio, tup_returned, tup_fetched, tup_inserted, tup_updated, tup_deleted FROM pg_stat_database WHERE datname NOT IN ('template0','template1') ORDER BY blks_hit DESC;
pg_bgwriter      = SELECT checkpoints_timed, checkpoints_req, buffers_checkpoint, buffers_clean, buffers_backend FROM pg_stat_bgwriter;
pg_settings_key  = SELECT name, setting, unit, short_desc FROM pg_settings WHERE name IN ('max_connections','shared_buffers','work_mem','maintenance_work_mem','effective_cache_size','wal_level','archive_mode','max_wal_size','checkpoint_completion_target','random_page_cost','log_min_duration_statement','autovacuum','autovacuum_vacuum_scale_factor','autovacuum_analyze_scale_factor') ORDER BY name;
pg_users         = SELECT rolname AS username, rolsuper AS is_superuser, rolcreatedb AS can_createdb, rolcreaterole AS can_createrole, rolvaliduntil AS password_expiry FROM pg_roles ORDER BY rolname;
pg_databases     = SELECT datname, pg_encoding_to_char(encoding) AS encoding, datcollate, datctype, datallowconn, datconnlimit FROM pg_database WHERE datistemplate=false ORDER BY datname;
pg_extensions    = SELECT name, default_version, installed_version, comment FROM pg_available_extensions WHERE installed_version IS NOT NULL ORDER BY name;
pg_processlist   = SELECT pid, usename, datname, application_name, client_addr, state, left(query,100) AS query, now()-query_start AS duration FROM pg_stat_activity WHERE pid <> pg_backend_pid() ORDER BY duration DESC NULLS LAST LIMIT 20;
"""


class RemoteSystemInfoCollector:
    """远程系统信息收集器 - 通过SSH连接获取远程主机信息"""

    def __init__(self, host, port=22, username='root', password=None, key_file=None):
        """
        初始化远程系统信息收集器。

        :param host: 远程主机 IP 地址或主机名
        :param port: SSH 端口，默认 22
        :param username: SSH 登录用户名，默认 root
        :param password: SSH 登录密码（与 key_file 二选一）
        :param key_file: SSH 私钥文件路径（与 password 二选一）
        """
        self.host = host
        self.port = port
        self.username = username
        self.password = password
        self.key_file = key_file
        self.ssh_client = None

    def connect(self):
        """
        建立 SSH 连接。

        优先使用私钥文件认证，若无私钥则使用密码认证。
        自动接受远程主机密钥（AutoAddPolicy）。

        :return: 连接成功返回 True，失败返回 False
        """
        try:
            self.ssh_client = paramiko.SSHClient()
            self.ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            if self.key_file:
                private_key = paramiko.RSAKey.from_private_key_file(self.key_file)
                self.ssh_client.connect(hostname=self.host, port=self.port,
                                        username=self.username, pkey=private_key, timeout=10)
            else:
                self.ssh_client.connect(hostname=self.host, port=self.port,
                                        username=self.username, password=self.password, timeout=10)
            return True
        except Exception as e:
            print("SSH连接失败 %s:%s: %s" % (self.host, self.port, str(e)))
            return False

    def disconnect(self):
        """断开 SSH 连接，释放资源。"""
        if self.ssh_client:
            self.ssh_client.close()

    def execute_command(self, command):
        """
        在远程主机上执行 Shell 命令。

        :param command: 要执行的 Shell 命令字符串
        :return: (stdout 输出内容, stderr 错误内容) 的元组；执行失败时返回 ("", 错误信息)
        """
        try:
            stdin, stdout, stderr = self.ssh_client.exec_command(command, timeout=10)
            output = stdout.read().decode('utf-8').strip()
            error = stderr.read().decode('utf-8').strip()
            return output, error
        except Exception as e:
            print("执行命令失败: %s, 错误: %s" % (command, str(e)))
            return "", str(e)

    def get_cpu_info(self):
        """
        通过远程 Shell 命令采集 CPU 信息。

        :return: 包含 usage_percent、physical_cores、logical_cores、
                 current_frequency、max_frequency 的字典；失败时返回空字典
        """
        try:
            cmd = "top -bn1 | grep 'Cpu(s)' | awk '{print $2}' | cut -d'%' -f1"
            output, _ = self.execute_command(cmd)
            cpu_percent = float(output) if output else 0.0
            cmd = "nproc"
            output, _ = self.execute_command(cmd)
            logical_cores = int(output) if output else 0
            cmd = "lscpu | grep 'Core(s) per socket' | awk '{print $4}'"
            output, _ = self.execute_command(cmd)
            cores_per_socket = int(output) if output else 1
            cmd = "lscpu | grep 'Socket(s)' | awk '{print $2}'"
            output, _ = self.execute_command(cmd)
            sockets = int(output) if output else 1
            physical_cores = cores_per_socket * sockets
            cmd = "lscpu | grep 'CPU MHz' | awk '{print $3}'"
            output, _ = self.execute_command(cmd)
            current_frequency = float(output) / 1000 if output else 0.0
            cmd = "lscpu | grep 'CPU max MHz' | awk '{print $4}'"
            output, _ = self.execute_command(cmd)
            max_frequency = float(output) / 1000 if output else current_frequency
            return {
                'usage_percent': cpu_percent,
                'physical_cores': physical_cores,
                'logical_cores': logical_cores,
                'current_frequency': round(current_frequency, 2),
                'max_frequency': round(max_frequency, 2)
            }
        except Exception as e:
            print("获取CPU信息失败: %s" % str(e))
            return {}

    def get_memory_info(self):
        """
        通过远程 Shell 命令采集内存信息。

        :return: 包含 total_gb、available_gb、used_gb、usage_percent、
                 swap_total_gb、swap_used_gb、swap_usage_percent 的字典；失败时返回空字典
        """
        try:
            cmd = "free -b | grep Mem"
            output, _ = self.execute_command(cmd)
            if output:
                parts = output.split()
                total_bytes = int(parts[1])
                used_bytes = int(parts[2])
                free_bytes = int(parts[3])
                available_bytes = int(parts[6]) if len(parts) > 6 else free_bytes
                usage_percent = (used_bytes / total_bytes) * 100 if total_bytes > 0 else 0
                memory_info = {
                    'total_gb': round(total_bytes / (1024**3), 2),
                    'available_gb': round(available_bytes / (1024**3), 2),
                    'used_gb': round(used_bytes / (1024**3), 2),
                    'usage_percent': round(usage_percent, 2)
                }
                cmd = "free -b | grep Swap"
                output, _ = self.execute_command(cmd)
                if output:
                    parts = output.split()
                    swap_total = int(parts[1])
                    swap_used = int(parts[2])
                    swap_usage_percent = (swap_used / swap_total) * 100 if swap_total > 0 else 0
                    memory_info.update({
                        'swap_total_gb': round(swap_total / (1024**3), 2),
                        'swap_used_gb': round(swap_used / (1024**3), 2),
                        'swap_usage_percent': round(swap_usage_percent, 2)
                    })
                return memory_info
            return {}
        except Exception as e:
            print("获取内存信息失败: %s" % str(e))
            return {}

    def get_disk_info(self):
        """
        通过远程 Shell 命令采集磁盘使用信息，过滤虚拟文件系统。

        :return: 磁盘分区信息列表，每项包含 device、mountpoint、fstype、
                 total_gb、used_gb、free_gb、usage_percent；失败时返回空列表
        """
        try:
            cmd = "df -h | grep -v 'tmpfs' | grep -v 'devtmpfs' | tail -n +2"
            output, _ = self.execute_command(cmd)
            disk_data = []
            if output:
                lines = output.strip().split('\n')
                for line in lines:
                    parts = line.split()
                    if len(parts) >= 6:
                        device = parts[0]
                        mountpoint = parts[5]
                        if any(vfs in device for vfs in ['tmpfs', 'devtmpfs', 'overlay']):
                            continue
                        if mountpoint.startswith('/mnt/') and not mountpoint.startswith('/mnt/share'): continue
                        size_str = parts[1]
                        used_str = parts[2]
                        avail_str = parts[3]
                        usage_percent_str = parts[4].rstrip('%')

                        def to_gb(s):
                            if not s: return 0.0
                            s = s.strip().upper()
                            if s.endswith('G'): return round(float(s[:-1]), 2)
                            elif s.endswith('M'): return round(float(s[:-1]) / 1024, 2)
                            elif s.endswith('T'): return round(float(s[:-1]) * 1024, 2)
                            elif s.endswith('K'): return round(float(s[:-1]) / (1024**2), 2)
                            else:
                                try: return round(float(s), 2)
                                except: return 0.0

                        total_gb = to_gb(size_str)
                        used_gb = to_gb(used_str)
                        free_gb = to_gb(avail_str)
                        try: usage_percent = float(usage_percent_str)
                        except: usage_percent = 0.0
                        disk_data.append({
                            'device': device, 'mountpoint': mountpoint, 'fstype': "ext4",
                            'total_gb': total_gb, 'used_gb': used_gb,
                            'free_gb': free_gb, 'usage_percent': usage_percent
                        })
            return disk_data
        except Exception as e:
            print("获取磁盘信息失败: %r" % e)
            return []

    def get_system_info(self):
        """
        聚合采集远程主机全部系统信息（CPU/内存/磁盘/主机名/平台/启动时间）。

        :return: 系统信息字典；SSH 连接失败时返回空字典
        """
        if not self.connect():
            return {}
        try:
            system_info = {
                'cpu': self.get_cpu_info(),
                'memory': self.get_memory_info(),
                'disk': self.get_disk_info(),
                'hostname': "",
                'platform': "",
                'boot_time': ""
            }
            cmd = "hostname"
            output, _ = self.execute_command(cmd)
            if output: system_info['hostname'] = output.strip()
            cmd = "uname -a"
            output, _ = self.execute_command(cmd)
            if output: system_info['platform'] = output.strip()
            cmd = "who -b | awk '{print $3 \" \" $4}'"
            output, _ = self.execute_command(cmd)
            if output: system_info['boot_time'] = output.strip()
            return system_info
        finally:
            self.disconnect()


class LocalSystemInfoCollector:
    """本地系统信息收集器 - 使用 psutil 库采集当前主机系统信息，无需 SSH"""

    def __init__(self):
        """初始化本地系统信息收集器（无需任何参数）。"""
        pass

    def get_cpu_info(self):
        """
        采集本机 CPU 信息（使用率、核心数、频率）。

        :return: CPU 信息字典；失败时返回空字典
        """
        try:
            cpu_percent = psutil.cpu_percent(interval=1)
            cpu_count = psutil.cpu_count(logical=False)
            cpu_count_logical = psutil.cpu_count(logical=True)
            cpu_freq = psutil.cpu_freq()
            return {
                'usage_percent': cpu_percent,
                'physical_cores': cpu_count,
                'logical_cores': cpu_count_logical,
                'current_frequency': round(cpu_freq.current, 2) if cpu_freq else 'N/A',
                'max_frequency': round(cpu_freq.max, 2) if cpu_freq else 'N/A'
            }
        except Exception as e:
            print("获取CPU信息失败: %s" % str(e))
            return {}

    def get_memory_info(self):
        """
        采集本机内存及 Swap 使用情况，单位 GB。

        :return: 内存信息字典；失败时返回空字典
        """
        try:
            memory = psutil.virtual_memory()
            swap = psutil.swap_memory()
            return {
                'total_gb': round(memory.total / (1024**3), 2),
                'available_gb': round(memory.available / (1024**3), 2),
                'used_gb': round(memory.used / (1024**3), 2),
                'usage_percent': memory.percent,
                'swap_total_gb': round(swap.total / (1024**3), 2),
                'swap_used_gb': round(swap.used / (1024**3), 2),
                'swap_usage_percent': swap.percent
            }
        except Exception as e:
            print("获取内存信息失败: %s" % str(e))
            return {}

    def get_disk_info(self):
        """
        采集本机磁盘分区信息，并额外检查常见 PostgreSQL 数据目录。

        :return: 以挂载点为键的磁盘信息字典；失败时返回空字典
        """
        try:
            disk_info = {}
            partitions = psutil.disk_partitions()
            IGNORE_PREFIXES = ('/mnt/', '/media', '/run/media', '/snap')
            for partition in partitions:
                mp = partition.mountpoint
                if partition.fstype and 'loop' not in partition.device and not mp.startswith(IGNORE_PREFIXES):
                    try:
                        usage = psutil.disk_usage(partition.mountpoint)
                        disk_info[partition.mountpoint] = {
                            'device': partition.device, 'mountpoint': partition.mountpoint,
                            'fstype': partition.fstype,
                            'total_gb': round(usage.total / (1024**3), 2),
                            'used_gb': round(usage.used / (1024**3), 2),
                            'free_gb': round(usage.free / (1024**3), 2),
                            'usage_percent': usage.percent
                        }
                    except PermissionError:
                        continue
            # 额外检查常见 PostgreSQL 数据目录
            pg_paths = ['/var/lib/postgresql', '/data/postgresql', '/usr/local/pgsql/data']
            for path in pg_paths:
                if os.path.exists(path):
                    try:
                        usage = psutil.disk_usage(path)
                        disk_info[f'pg_data_{path}'] = {
                            'device': 'PG Data', 'mountpoint': path, 'fstype': 'N/A',
                            'total_gb': round(usage.total / (1024**3), 2),
                            'used_gb': round(usage.used / (1024**3), 2),
                            'free_gb': round(usage.free / (1024**3), 2),
                            'usage_percent': usage.percent
                        }
                    except Exception:
                        pass
            return disk_info
        except Exception as e:
            print("获取磁盘信息失败: %r" % e)
            return {}

    def get_system_info(self):
        """
        聚合采集本机全部系统信息（CPU/内存/磁盘/主机名/平台/启动时间）。

        :return: 系统信息字典
        """
        return {
            'cpu': self.get_cpu_info(),
            'memory': self.get_memory_info(),
            'disk': self.get_disk_info(),
            'hostname': socket.gethostname(),
            'platform': platform.platform(),
            'boot_time': datetime.fromtimestamp(psutil.boot_time()).strftime('%Y-%m-%d %H:%M:%S')
        }


class SystemInfoCollector:
    """系统信息收集器工厂类 - 根据主机类型创建对应的采集器实例"""

    @staticmethod
    def create_collector(host_type='local', **kwargs):
        """
        工厂方法：根据 host_type 返回合适的系统信息采集器。

        :param host_type: 'local' 使用本地采集，'remote' 使用 SSH 远程采集
        :param kwargs: remote 模式下的 SSH 连接参数
        :return: LocalSystemInfoCollector 或 RemoteSystemInfoCollector 实例
        """
        if host_type == 'remote':
            return RemoteSystemInfoCollector(**kwargs)
        else:
            return LocalSystemInfoCollector()


def get_host_disk_usage():
    """
    跨平台获取主机磁盘使用情况的备用函数。

    Windows 使用 wmic，Linux/macOS 使用 df -h。

    :return: 磁盘分区信息列表；失败时返回空列表
    """
    try:
        disk_data = []
        if platform.system() == "Windows":
            result = subprocess.Popen(
                ["wmic", "logicaldisk", "get", "deviceid,size,freespace", "/format:csv"],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            stdout, stderr = result.communicate(timeout=10)
            output = stdout.decode('utf-8', errors='ignore')
            lines = output.strip().split('\n')[1:]
            for line in lines:
                if line.strip():
                    parts = line.split(',')
                    if len(parts) >= 4:
                        device = parts[1]
                        total_bytes = int(parts[2]) if parts[2] else 0
                        free_bytes = int(parts[3]) if parts[3] else 0
                        used_bytes = total_bytes - free_bytes
                        total_gb = round(total_bytes / (1024**3), 2)
                        used_gb = round(used_bytes / (1024**3), 2)
                        free_gb = round(free_bytes / (1024**3), 2)
                        usage_percent = round((used_bytes / total_bytes) * 100, 2) if total_bytes > 0 else 0
                        disk_data.append({
                            'device': device, 'mountpoint': device + "\\", 'fstype': "NTFS",
                            'total_gb': total_gb, 'used_gb': used_gb,
                            'free_gb': free_gb, 'usage_percent': usage_percent
                        })
        else:
            result = subprocess.Popen(["df", "-h"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            stdout, stderr = result.communicate(timeout=10)
            output = stdout.decode('utf-8', errors='ignore')
            lines = output.strip().split('\n')[1:]
            important_mounts = ['/', '/boot', '/home', '/var', '/usr', '/opt', '/tmp']
            for line in lines:
                parts = line.split()
                if len(parts) >= 6:
                    device = parts[0]
                    mountpoint = parts[5]
                    if mountpoint in important_mounts:
                        def to_gb(s):
                            if not s: return 0.0
                            s = s.strip().upper()
                            if s.endswith('G'): return round(float(s[:-1]), 2)
                            elif s.endswith('M'): return round(float(s[:-1]) / 1024, 2)
                            elif s.endswith('T'): return round(float(s[:-1]) * 1024, 2)
                            elif s.endswith('K'): return round(float(s[:-1]) / (1024**2), 2)
                            else:
                                try: return round(float(s), 2)
                                except: return 0.0
                        disk_data.append({
                            'device': device, 'mountpoint': mountpoint, 'fstype': "ext4",
                            'total_gb': to_gb(parts[1]), 'used_gb': to_gb(parts[2]),
                            'free_gb': to_gb(parts[3]),
                            'usage_percent': float(parts[4].rstrip('%')) if parts[4].rstrip('%') else 0.0
                        })
            if not disk_data:
                for line in lines:
                    parts = line.split()
                    if len(parts) >= 6:
                        device = parts[0]
                        mountpoint = parts[5]
                        if any(vfs in device for vfs in ['tmpfs', 'devtmpfs', 'overlay']):
                            continue
                        if mountpoint.startswith('/mnt/') and not mountpoint.startswith('/mnt/share'): continue
                        def to_gb(s):
                            if not s: return 0.0
                            s = s.strip().upper()
                            if s.endswith('G'): return round(float(s[:-1]), 2)
                            elif s.endswith('M'): return round(float(s[:-1]) / 1024, 2)
                            elif s.endswith('T'): return round(float(s[:-1]) * 1024, 2)
                            elif s.endswith('K'): return round(float(s[:-1]) / (1024**2), 2)
                            else:
                                try: return round(float(s), 2)
                                except: return 0.0
                        disk_data.append({
                            'device': device, 'mountpoint': mountpoint, 'fstype': "ext4",
                            'total_gb': to_gb(parts[1]), 'used_gb': to_gb(parts[2]),
                            'free_gb': to_gb(parts[3]),
                            'usage_percent': float(parts[4].rstrip('%')) if parts[4].rstrip('%') else 0.0
                        })
        return disk_data
    except Exception as e:
        print(f"获取磁盘使用率失败: {str(e)}")
        return []





class WordTemplateGenerator:
    """Word 模板生成器 - 动态生成包含 Jinja2 模板语法的巡检报告 Word 模板文件"""

    def __init__(self, inspector_name="Jack"):
        """
        初始化 Word 模板生成器。

        创建新的 Document 对象并调用 _setup_document() 设置页面边距。
        """
        self.doc = Document()
        self.inspector_name = inspector_name
        self._setup_document()

        # Pre-compute all translated strings for use in template sections
        # Title page
        self._l = {
            'db_name':         self._t('report.fallback_db_name'),
            'server_addr':     self._t('report.fallback_server_addr'),
            'pg_version':      self._t('report.fallback_pg_version'),
            'hostname':        self._t('report.fallback_hostname'),
            'instance_time':   self._t('report.fallback_instance_time'),
            'inspector':       self._t('report.fallback_inspector'),
            'platform':        self._t('report.fallback_platform'),
            'report_time':     self._t('report.fallback_report_time'),
            'pg_title':        self._t('report.pg_title'),
            'overall_health':  self._t('report.fallback_overall_health'),
            'issue_count':    self._t('report.fallback_issue_count'),
            # Chapter headings
            'ch3':             self._t('report.pg_ch3'),
            'ch31':            self._t('report.pg_ch31'),
            'ch32':            self._t('report.pg_ch32'),
            'ch33':            self._t('report.pg_ch33'),
            'ch4':             self._t('report.pg_ch4'),
            'ch41':            self._t('report.pg_ch41'),
            'ch42':            self._t('report.pg_ch42'),
            'ch43':            self._t('report.pg_ch43'),
            'ch44':            self._t('report.pg_ch44'),
            'ch5':             self._t('report.pg_ch5'),
            'ch51':            self._t('report.pg_ch51'),
            'ch6':             self._t('report.pg_ch6'),
            'ch7':             self._t('report.pg_ch7'),
            'ch8':             self._t('report.pg_ch8'),
            # Table headers
            'hdr_config':      self._t('report.pg_hdr_config_item'),
            'hdr_cur_val':     self._t('report.pg_hdr_current_value'),
            'hdr_desc':        self._t('report.pg_hdr_description'),
            'hdr_metric':      self._t('report.pg_hdr_metric'),
            'hdr_value':       self._t('report.pg_hdr_value'),
            'hdr_state':       self._t('report.pg_hdr_state'),
            'hdr_conn_cnt':    self._t('report.pg_hdr_conn_count'),
            'hdr_db':          self._t('report.pg_hdr_database'),
            'hdr_size':        self._t('report.pg_hdr_size_gb'),
            'hdr_total':       self._t('report.pg_hdr_total'),
            'hdr_user':        self._t('report.pg_hdr_user'),
            'hdr_member_of':   self._t('report.pg_hdr_member_of'),
            'hdr_super':       self._t('report.pg_hdr_super'),
            'hdr_locked':      self._t('report.pg_hdr_locked'),
            # 3.1 Connection items
            'max_conn':        self._t('report.pg_max_connections'),
            'max_conn_desc':   self._t('report.pg_max_connections_desc'),
            'shared_buf':      self._t('report.pg_shared_buffers'),
            'shared_buf_desc': self._t('report.pg_shared_buffers_desc'),
            'eff_cache':       self._t('report.pg_effective_cache'),
            'eff_cache_desc':  self._t('report.pg_effective_cache_desc'),
            'work_mem':        self._t('report.pg_work_mem'),
            'work_mem_desc':   self._t('report.pg_work_mem_desc'),
            # 3.2 Memory items
            'maint_mem':       self._t('report.pg_maint_work_mem'),
            'maint_mem_desc':  self._t('report.pg_maint_work_mem_desc'),
            'wal_level':       self._t('report.pg_wal_level'),
            'wal_level_desc':  self._t('report.pg_wal_level_desc'),
            'max_wal':         self._t('report.pg_max_wal_size'),
            'max_wal_desc':    self._t('report.pg_max_wal_size_desc'),
            'ckpt_target':     self._t('report.pg_checkpoint_target'),
            'ckpt_target_desc':self._t('report.pg_checkpoint_target_desc'),
            'rnd_page_cost':   self._t('report.pg_random_page_cost'),
            'rnd_page_cost_desc': self._t('report.pg_random_page_cost_desc'),
            # 3.3 Autovacuum items
            'autovacuum':      self._t('report.pg_autovacuum'),
            'autovacuum_desc': self._t('report.pg_autovacuum_desc'),
            'vacuum_factor':   self._t('report.pg_autovacuum_vacuum_factor'),
            'vacuum_factor_desc': self._t('report.pg_autovacuum_vacuum_factor_desc'),
            'analyze_factor':  self._t('report.pg_autovacuum_analyze_factor'),
            'analyze_factor_desc': self._t('report.pg_autovacuum_analyze_factor_desc'),
            'slow_query':      self._t('report.pg_slow_query_threshold'),
            'slow_query_desc': self._t('report.pg_slow_query_threshold_desc'),
            # 4.1 Connection status
            'cur_conn':        self._t('report.pg_current_connections'),
            'max_conn_hdr':    self._t('report.pg_max_connections_hdr'),
            'conn_usage':      self._t('report.pg_conn_usage'),
            # 4.3 Cache hit
            'dbname':          self._t('report.pg_dbname'),
            'blks_hit':        self._t('report.pg_blks_hit'),
            'blks_read':       self._t('report.pg_blks_read'),
            'hit_ratio':       self._t('report.pg_hit_ratio'),
            # 5.1 Security
            'superuser':       self._t('report.pg_superuser'),
            'locked':          self._t('report.pg已锁定'),
            # Chapter 2 - System Resources
            'ch2_title':       self._t('report.pg_ch2_title'),
            'ch2_cpu':         self._t('report.pg_ch2_cpu'),
            'ch2_mem':         self._t('report.pg_ch2_mem'),
            'ch2_disk':        self._t('report.pg_ch2_disk'),
            'hdr_cpu_usage':   self._t('report.pg_cpu_usage_hdr'),
            'hdr_physical':    self._t('report.pg_physical_cores_hdr'),
            'hdr_logical':     self._t('report.pg_logical_cores_hdr'),
            'hdr_freq':        self._t('report.pg_freq_hdr'),
            'hdr_total_mem':   self._t('report.pg_total_mem_hdr'),
            'hdr_used_mem':    self._t('report.pg_used_mem_hdr'),
            'hdr_avail_mem':   self._t('report.pg_avail_mem_hdr'),
            'hdr_mem_usage':   self._t('report.pg_mem_usage_hdr'),
            'hdr_mountpoint':  self._t('report.pg_mountpoint_hdr'),
            'hdr_disk_usage':  self._t('report.fallback_usage_pct'),
            'na':              self._t('report.pg_fallback_na'),
        }

    def _t(self, key):
        try:
            from i18n import t
            return t(key, _PG_LANG)
        except Exception:
            return key

    def _setup_document(self):
        """
        设置 Word 文档的页面边距。

        将文档所有节（Section）的上、下、左、右边距统一设置为 2.54 cm（标准 A4 边距）。
        """
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    def create_template(self):
        """
        按顺序组装完整的巡检报告模板，包含封面及 7 个章节。

        依次调用以下方法：
          封面 → 健康状态概览 → 系统资源检查 → PostgreSQL 配置检查
          → 性能分析 → 数据库信息 → 安全信息 → 报告说明

        :return: 组装完成的 Document 对象
        """
        self._add_title_page()
        self._add_summary_section()
        self._add_system_info_section()
        self._add_pg_config_section()
        self._add_performance_section()
        self._add_database_info_section()
        self._add_security_section()
        self._add_notes_section()
        return self.doc

    def _add_title_page(self):
        """
        生成报告封面页。

        包含标题（PostgreSQL数据库健康巡检报告）和一个 8 行 2 列的信息表，
        表格中使用 Jinja2 模板变量填充：数据库名称、服务器地址、PostgreSQL版本、
        服务器主机名、实例启动时间、巡检人员、服务器平台、报告生成时间。
        封面末尾插入分页符。
        """
        title = self.doc.add_heading(self._l['pg_title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        self.doc.add_paragraph()
        table = self.doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        cells = table.rows[0].cells
        cells[0].text = self._l['db_name']
        cells[1].text = "{{ co_name[0]['CO_NAME'] }}"
        cells = table.rows[1].cells
        cells[0].text = self._l['server_addr']
        cells[1].text = "{{ ip[0]['IP'] }}:{{ port[0]['PORT'] }}"
        cells = table.rows[2].cells
        cells[0].text = self._l['pg_version']
        cells[1].text = "{{ myversion[0]['version'] }}"
        cells = table.rows[3].cells
        cells[0].text = self._l['hostname']
        cells[1].text = "{{ system_info.hostname }}"
        cells = table.rows[4].cells
        cells[0].text = self._l['instance_time']
        cells[1].text = "{% if instancetime %}{{ instancetime[0]['started_at'] }}{% else %}N/A{% endif %}"
        cells = table.rows[5].cells
        cells[0].text = self._l['inspector']
        cells[1].text = "{{ inspector_name }}"
        cells = table.rows[6].cells
        cells[0].text = self._l['platform']
        cells[1].text = "{% if platform_info and platform_info|length > 0 %}{% for item in platform_info %}{% if item.variable_name == 'version_compile_os' %}{{ item.variable_value }}{% endif %}{% endfor %}{% else %}N/A{% endif %}"
        cells = table.rows[7].cells
        cells[0].text = self._l['report_time']
        cells[1].text = "{{ report_time }}"
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_page_break()

    def _add_summary_section(self):
        """
        生成第 1 章「健康状态概览」。

        包含一个 2 行 2 列的状态表（总体健康状态、发现问题数量）
        以及健康总结段落，所有值均使用 Jinja2 模板变量占位。
        """
        heading = self.doc.add_heading('1. ' + self._l['overall_health'], level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        cells = table.rows[0].cells
        cells[0].text = self._l['overall_health']
        cells[1].text = "{{ health_status }}"
        cells = table.rows[1].cells
        cells[0].text = self._l['issue_count']
        cells[1].text = "{{ problem_count }} 个"
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()
        p = self.doc.add_paragraph(self._t('report.fallback_health_summary') + ": ")
        p.add_run("{{ health_summary[0]['health_summary'] }}").bold = True
        p.runs[0].font.size = Pt(11)
        p.runs[1].font.size = Pt(11)
    def _add_system_info_section(self):
        """
        生成第 2 章「系统资源检查」，包含三个小节：

        - 2.1 CPU 信息：4 列表格（使用率、物理/逻辑核心数、当前频率）
        - 2.2 内存信息：4 列表格（总量、已用、可用、使用率）
        - 2.3 磁盘信息：2 列表格（挂载点、使用率），使用 Jinja2 循环语法动态生成行
        """
        heading = self.doc.add_heading(self._l['ch2_title'], level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        sub_heading = self.doc.add_heading(self._l['ch2_cpu'], level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        # CPU table: 4 cols, auto-fit to window
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.autofit = True  # 根据窗口自动调整
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = self._l['hdr_cpu_usage']
        hdr_cells[1].text = self._l['hdr_physical']
        hdr_cells[2].text = self._l['hdr_logical']
        hdr_cells[3].text = self._l['hdr_freq']
        na = self._l['na']
        data_cells = table.rows[1].cells
        data_cells[0].text = "{% if system_info.cpu and system_info.cpu.usage_percent is defined %}{{ '%.2f'|format(system_info.cpu.usage_percent) }}%{% else %}" + na + "{% endif %}"
        data_cells[1].text = "{% if system_info.cpu and system_info.cpu.physical_cores is defined %}{{ system_info.cpu.physical_cores }}{% else %}" + na + "{% endif %}"
        data_cells[2].text = "{% if system_info.cpu and system_info.cpu.logical_cores is defined %}{{ system_info.cpu.logical_cores }}{% else %}" + na + "{% endif %}"
        data_cells[3].text = "{% if system_info.cpu and system_info.cpu.current_frequency != 'N/A' %}{{ '%.2f'|format(system_info.cpu.current_frequency/1000) }}{% else %}" + na + "{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading(self._l['ch2_mem'], level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        # Memory table: 4 cols, auto-fit to window
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.autofit = True  # 根据窗口自动调整
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = self._l['hdr_total_mem']
        hdr_cells[1].text = self._l['hdr_used_mem']
        hdr_cells[2].text = self._l['hdr_avail_mem']
        hdr_cells[3].text = self._l['hdr_mem_usage']
        data_cells = table.rows[1].cells
        data_cells[0].text = "{% if system_info.memory and system_info.memory.total_gb is defined %}{{ '%.2f'|format(system_info.memory.total_gb) }}{% else %}" + na + "{% endif %}"
        data_cells[1].text = "{% if system_info.memory and system_info.memory.used_gb is defined %}{{ '%.2f'|format(system_info.memory.used_gb) }}{% else %}" + na + "{% endif %}"
        data_cells[2].text = "{% if system_info.memory and system_info.memory.available_gb is defined %}{{ '%.2f'|format(system_info.memory.available_gb) }}{% else %}" + na + "{% endif %}"
        data_cells[3].text = "{% if system_info.memory and system_info.memory.usage_percent is defined %}{{ '%.2f'|format(system_info.memory.usage_percent) }}%{% else %}" + na + "{% endif %}"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
        sub_heading = self.doc.add_heading(self._l['ch2_disk'], level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        # Disk table: 2 cols, auto-fit to window
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True  # 根据窗口自动调整
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = self._l['hdr_mountpoint']
        hdr_cells[1].text = self._l['hdr_disk_usage']
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells = table.add_row().cells
        row_cells[0].text = "{% for disk in system_info.disk_list %}{{ disk.mountpoint }}{% if not loop.last %}\n{% endif %}{% endfor %}"
        row_cells[1].text = "{% for disk in system_info.disk_list %}{% if disk.usage_percent is defined %}{{ '%.2f'|format(disk.usage_percent) }}%{% else %}" + na + "{% endif %}{% if not loop.last %}\n{% endif %}{% endfor %}"
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
    def _add_pg_config_section(self):
        """
        生成第 3 章「PostgreSQL 配置检查」，包含三个小节。

        - 3.1 连接配置：最大连接数、当前连接数、使用率
        - 3.2 内存配置：shared_buffers、work_mem、maintenance_work_mem、effective_cache_size
        - 3.3 核心参数：wal_level、archive_mode、checkpoint_completion_target、autovacuum 等
        每小节均使用「配置项 / 当前值 / 说明」三列表格，值通过 Jinja2 遍历 pg_settings_key 列表填充。
        """
        heading = self.doc.add_heading('3. ' + self._l['ch3'], level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True

        # ---- 3.1 连接配置 ----
        sub_heading = self.doc.add_heading('3.1 ' + self._l['ch31'], level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(6)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_config']
        hdr[1].text = self._l['hdr_cur_val']
        hdr[2].text = self._l['hdr_desc']
        conn_items = [
            ('max_connections',     self._l['max_conn'],        self._l['max_conn_desc']),
            ('shared_buffers',     self._l['shared_buf'],      self._l['shared_buf_desc']),
            ('effective_cache_size', self._l['eff_cache'],     self._l['eff_cache_desc']),
            ('work_mem',           self._l['work_mem'],        self._l['work_mem_desc']),
        ]
        for idx, (name, label, desc) in enumerate(conn_items, 1):
            cells = table.rows[idx].cells
            cells[0].text = label
            cells[1].text = f"{{{{ pg_settings_key|selectattr('name', 'equalto', '{name}')|map(attribute='setting')|first|default('') }}}}"
            cells[2].text = desc
            for cell in cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()

        # ---- 3.2 内存与存储配置 ----
        sub_heading = self.doc.add_heading('3.2 ' + self._l['ch32'], level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(6)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_config']
        hdr[1].text = self._l['hdr_cur_val']
        hdr[2].text = self._l['hdr_desc']
        mem_items = [
            ('maintenance_work_mem',        self._l['maint_mem'],          self._l['maint_mem_desc']),
            ('wal_level',                   self._l['wal_level'],          self._l['wal_level_desc']),
            ('max_wal_size',               self._l['max_wal'],            self._l['max_wal_desc']),
            ('checkpoint_completion_target', self._l['ckpt_target'],        self._l['ckpt_target_desc']),
            ('random_page_cost',            self._l['rnd_page_cost'],      self._l['rnd_page_cost_desc']),
        ]
        for idx, (name, label, desc) in enumerate(mem_items, 1):
            cells = table.rows[idx].cells
            cells[0].text = label
            cells[1].text = f"{{{{ pg_settings_key|selectattr('name', 'equalto', '{name}')|map(attribute='setting')|first|default('') }}}}"
            cells[2].text = desc
            for cell in cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()

        # ---- 3.3 自动清理与日志配置 ----
        sub_heading = self.doc.add_heading('3.3 ' + self._l['ch33'], level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(6)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_config']
        hdr[1].text = self._l['hdr_cur_val']
        hdr[2].text = self._l['hdr_desc']
        other_items = [
            ('autovacuum',                       self._l['autovacuum'],          self._l['autovacuum_desc']),
            ('autovacuum_vacuum_scale_factor',   self._l['vacuum_factor'],        self._l['vacuum_factor_desc']),
            ('autovacuum_analyze_scale_factor',  self._l['analyze_factor'],       self._l['analyze_factor_desc']),
            ('log_min_duration_statement',        self._l['slow_query'],           self._l['slow_query_desc']),
        ]
        for idx, (name, label, desc) in enumerate(other_items, 1):
            cells = table.rows[idx].cells
            cells[0].text = label
            cells[1].text = f"{{{{ pg_settings_key|selectattr('name', 'equalto', '{name}')|map(attribute='setting')|first|default('') }}}}"
            cells[2].text = desc
            for cell in cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def _add_performance_section(self):
        """
        生成第 4 章「性能分析」，包含四个小节：

        - 4.1 连接状态：当前连接数、最大连接数、使用率
        - 4.2 缓存命中率：各数据库的共享缓冲区命中率
        - 4.3 数据库大小：各数据库占用空间
        - 4.4 后台写入器：检查点统计和缓冲区分配信息
        所有值均通过 Jinja2 遍历对应列表填充。
        """
        heading = self.doc.add_heading('4. ' + self._l['ch4'], level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True

        # ---- 4.1 连接状态 ----
        sub_heading = self.doc.add_heading('4.1 ' + self._l['ch41'], level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(10)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_metric']
        hdr[1].text = self._l['hdr_value']
        conn_data = [
            (self._l['cur_conn'], "{% if pg_connections and pg_connections[0] %}{{ pg_connections[0]['total_connections'] }}{% endif %}"),
            (self._l['max_conn_hdr'], "{% if pg_connections and pg_connections[0] %}{{ pg_connections[0]['max_connections'] }}{% endif %}"),
            (self._l['conn_usage'], "{% if pg_connections and pg_connections[0] %}{{ pg_connections[0]['usage_percent'] }}%{% endif %}"),
        ]
        for idx, (label, val_tpl) in enumerate(conn_data, 1):
            cells = table.rows[idx].cells
            cells[0].text = label
            cells[1].text = val_tpl
            for cell in cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()

        # ---- 4.2 连接详情 ----
        sub_heading = self.doc.add_heading('4.2 ' + self._l['ch42'], level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(4)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_state']
        hdr[1].text = self._l['hdr_conn_cnt']
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(8):
            row_cells = table.add_row().cells
            state_tpl = "{{{{ pg_conn_detail[{}].state if pg_conn_detail and pg_conn_detail[{}] else '' }}}}".format(i, i)
            cnt_tpl = "{{{{ pg_conn_detail[{}].count if pg_conn_detail and pg_conn_detail[{}] else '' }}}}".format(i, i)
            row_cells[0].text = state_tpl
            row_cells[1].text = cnt_tpl
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()

        # ---- 4.3 缓存命中率 ----
        sub_heading = self.doc.add_heading('4.3 ' + self._l['ch43'], level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(2)
        table.columns[3].width = Cm(2)
        table.columns[4].width = Cm(5)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_db']
        hdr[1].text = self._l['blks_hit']
        hdr[2].text = self._l['blks_read']
        hdr[3].text = self._l['hit_ratio']
        hdr[4].text = self._l['hdr_desc']
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(8):
            row_cells = table.add_row().cells
            db_tpl = "{{{{ pg_cache_hit[{}].datname if pg_cache_hit and pg_cache_hit[{}] else '' }}}}".format(i, i)
            hit_tpl = "{{{{ pg_cache_hit[{}].blks_hit if pg_cache_hit and pg_cache_hit[{}] else '' }}}}".format(i, i)
            read_tpl = "{{{{ pg_cache_hit[{}].blks_read if pg_cache_hit and pg_cache_hit[{}] else '' }}}}".format(i, i)
            ratio_tpl = "{{{{ pg_cache_hit[{}].cache_hit_ratio if pg_cache_hit and pg_cache_hit[{}] else '' }}}}".format(i, i)
            desc_tpl = "{% if pg_cache_hit and pg_cache_hit[" + str(i) + "] %}{% if pg_cache_hit[" + str(i) + "]['cache_hit_ratio']|float >= 90 %}" + self._t('report.fallback_pg_cache_hit_good') + "{% elif pg_cache_hit[" + str(i) + "]['cache_hit_ratio']|float >= 70 %}" + self._t('report.fallback_pg_cache_hit_fair') + "{% else %}" + self._t('report.fallback_pg_cache_hit_poor') + "{% endif %}{% endif %}"
            row_cells[0].text = db_tpl
            row_cells[1].text = hit_tpl
            row_cells[2].text = read_tpl
            row_cells[3].text = ratio_tpl
            row_cells[4].text = desc_tpl
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()

        # ---- 4.4 数据库大小 ----
        sub_heading = self.doc.add_heading('4.4 ' + self._l['ch44'], level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(6)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_db']
        hdr[1].text = self._l['hdr_size']
        hdr[2].text = self._t('report.pg_fallback_pg_size_bytes')
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(8):
            row_cells = table.add_row().cells
            name_tpl = "{{{{ pg_db_size[{}].database_name if pg_db_size and pg_db_size[{}] else '' }}}}".format(i, i)
            size_tpl = "{{{{ pg_db_size[{}].size if pg_db_size and pg_db_size[{}] else '' }}}}".format(i, i)
            bytes_tpl = "{{{{ pg_db_size[{}].size_bytes if pg_db_size and pg_db_size[{}] else '' }}}}".format(i, i)
            row_cells[0].text = name_tpl
            row_cells[1].text = size_tpl
            row_cells[2].text = bytes_tpl
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def _add_database_info_section(self):
        """
        生成第 5 章「数据库信息」，包含三个小节：

        - 5.1 数据库列表：所有数据库的编码、排序规则、连接限制等信息
        - 5.2 当前进程列表：通过 pg_stat_activity 查看当前活动的连接
        - 5.3 扩展信息：已安装的 PostgreSQL 扩展
        所有值均通过 Jinja2 遍历对应列表填充。
        """
        heading = self.doc.add_heading('5. ' + self._l['ch5'], level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True

        # ---- 5.1 数据库列表 ----
        sub_heading = self.doc.add_heading('5.1 ' + self._t('report.pg_ch51'), level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(2)
        table.columns[3].width = Cm(2)
        table.columns[4].width = Cm(5)
        hdr = table.rows[0].cells
        hdr[0].text = self._l['hdr_db']
        hdr[1].text = self._t('report.pg_encoding')
        hdr[2].text = self._t('report.pg_collation')
        hdr[3].text = self._t('report.pg_allow_conn')
        hdr[4].text = self._t('report.pg_conn_limit')
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(10):
            row_cells = table.add_row().cells
            row_cells[0].text = "{{{{ pg_databases[{}].datname if pg_databases and pg_databases[{}] else '' }}}}".format(i, i)
            row_cells[1].text = "{{{{ pg_databases[{}].encoding if pg_databases and pg_databases[{}] else '' }}}}".format(i, i)
            row_cells[2].text = "{{{{ pg_databases[{}].datcollate if pg_databases and pg_databases[{}] else '' }}}}".format(i, i)
            row_cells[3].text = "{{{{ pg_databases[{}].datallowconn if pg_databases and pg_databases[{}] else '' }}}}".format(i, i)
            row_cells[4].text = "{{{{ pg_databases[{}].datconnlimit if pg_databases and pg_databases[{}] else '' }}}}".format(i, i)
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()

        # ---- 5.2 当前进程列表 ----
        sub_heading = self.doc.add_heading('5.2 ' + self._t('report.pg_current_processes'), level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(1.5)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(2.5)
        table.columns[3].width = Cm(2)
        table.columns[4].width = Cm(6)
        hdr = table.rows[0].cells
        hdr[0].text = 'PID'
        hdr[1].text = self._l['hdr_user']
        hdr[2].text = self._l['hdr_db']
        hdr[3].text = self._l['hdr_state']
        hdr[4].text = self._t('report.pg_current_sql')
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(15):
            row_cells = table.add_row().cells
            row_cells[0].text = "{{{{ pg_processlist[{}].pid if pg_processlist and pg_processlist[{}] else '' }}}}".format(i, i)
            row_cells[1].text = "{{{{ pg_processlist[{}].usename if pg_processlist and pg_processlist[{}] else '' }}}}".format(i, i)
            row_cells[2].text = "{{{{ pg_processlist[{}].datname if pg_processlist and pg_processlist[{}] else '' }}}}".format(i, i)
            row_cells[3].text = "{{{{ pg_processlist[{}].state if pg_processlist and pg_processlist[{}] else '' }}}}".format(i, i)
            row_cells[4].text = "{{{{ pg_processlist[{}].query if pg_processlist and pg_processlist[{}] else '' }}}}".format(i, i)
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.doc.add_paragraph()

        # ---- 5.3 已安装扩展 ----
        sub_heading = self.doc.add_heading('5.3 ' + self._t('report.pg_installed_extensions'), level=2)
        sub_heading.runs[0].font.size = Pt(12)
        sub_heading.runs[0].font.bold = True
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(3)
        table.columns[2].width = Cm(7)
        hdr = table.rows[0].cells
        hdr[0].text = self._t('report.pg_extension_name')
        hdr[1].text = self._t('report.pg_ext_version')
        hdr[2].text = self._l['hdr_desc']
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(10):
            row_cells = table.add_row().cells
            row_cells[0].text = "{{{{ pg_extensions[{}].name if pg_extensions and pg_extensions[{}] else '' }}}}".format(i, i)
            row_cells[1].text = "{{{{ pg_extensions[{}].installed_version if pg_extensions and pg_extensions[{}] else '' }}}}".format(i, i)
            row_cells[2].text = "{{{{ pg_extensions[{}].comment if pg_extensions and pg_extensions[{}] else '' }}}}".format(i, i)
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def _add_security_section(self):
        """
        生成第 6 章「安全信息」。

        包含 6.1「数据库用户信息」小节：5 列表格
        （用户名、超级用户、创建数据库权限、创建角色权限、密码过期时间），
        通过 Jinja2 按索引访问 pg_users 列表，最多展示 15 行。
        """
        heading = self.doc.add_heading('6. ' + self._l['ch5'], level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        sub_heading = self.doc.add_heading('6.1 ' + self._l['ch51'], level=2)
        sub_heading_run = sub_heading.runs[0]
        sub_heading_run.font.size = Pt(12)
        sub_heading_run.font.bold = True
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(2)
        table.columns[2].width = Cm(2)
        table.columns[3].width = Cm(2)
        table.columns[4].width = Cm(5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = self._l['hdr_user']
        hdr_cells[1].text = self._l['hdr_super']
        hdr_cells[2].text = self._t('report.pg_createdb')
        hdr_cells[3].text = self._t('report.pg_createrole')
        hdr_cells[4].text = self._t('report.pg_validuntil')
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(15):
            row_cells = table.add_row().cells
            row_cells[0].text = "{{{{ pg_users[{}].username if pg_users and pg_users[{}] else '' }}}}".format(i, i)
            row_cells[1].text = "{{{{ pg_users[{}].is_superuser if pg_users and pg_users[{}] else '' }}}}".format(i, i)
            row_cells[2].text = "{{{{ pg_users[{}].can_createdb if pg_users and pg_users[{}] else '' }}}}".format(i, i)
            row_cells[3].text = "{{{{ pg_users[{}].can_createrole if pg_users and pg_users[{}] else '' }}}}".format(i, i)
            row_cells[4].text = "{{{{ pg_users[{}].password_expiry if pg_users and pg_users[{}] else '' }}}}".format(i, i)
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def _add_notes_section(self):
        """
        生成第 7 章「报告说明」。

        以段落形式输出 5 条固定说明文字，包含：
        报告生成说明、空白项说明、磁盘信息范围说明、巡检结果免责说明、定期巡检建议。
        """
        heading = self.doc.add_heading('7. ' + self._l['ch6'], level=1)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        notes = [
            self._t('report.fallback_pg_note_1'),
            self._t('report.fallback_pg_note_2'),
            self._t('report.fallback_pg_note_3'),
            self._t('report.fallback_pg_note_4'),
            self._t('report.fallback_pg_note_5'),
        ]
        for note in notes:
            p = self.doc.add_paragraph()
            p.add_run(note)
            p.runs[0].font.size = Pt(10)

class SimpleCrypto:
    """简单加密解密工具类 - 基于 XOR + SHA256 密钥 + Base64 编码实现对称加密"""

    def __init__(self, secret="ODB_SECRET_2024"):
        """
        初始化加密工具。

        :param secret: 加密用的原始密钥字符串，默认为 "ODB_SECRET_2024"，
                       会被编码为 bytes 并通过 SHA256 派生实际加密密钥
        """
        self.secret = secret.encode('utf-8')

    def _xor_encrypt_decrypt(self, data, key):
        """
        使用 XOR 对数据进行加密或解密（XOR 可逆，加解密使用同一方法）。

        :param data: 需要处理的字节数据（bytes）
        :param key: 加密/解密使用的密钥字节（bytes），循环重复使用
        :return: 经过 XOR 处理后的字节数据（bytes）
        """
        key_len = len(key)
        result = bytearray()
        for i, byte in enumerate(data):
            result.append(byte ^ key[i % key_len])
        return bytes(result)

    def encrypt(self, text):
        """
        加密文本字符串。

        流程：UTF-8 编码 → SHA256 派生 32 字节密钥 → XOR 加密 → Base64 编码。

        :param text: 需要加密的明文字符串
        :return: Base64 编码后的密文字符串
        """
        data = text.encode('utf-8')
        key = hashlib.sha256(self.secret).digest()[:32]
        encrypted = self._xor_encrypt_decrypt(data, key)
        return base64.b64encode(encrypted).decode('utf-8')

    def decrypt(self, token):
        """
        解密 Base64 编码的密文字符串。

        流程：Base64 解码 → SHA256 派生 32 字节密钥 → XOR 解密 → UTF-8 解码。

        :param token: Base64 编码的密文字符串
        :return: 解密后的明文字符串
        :raises ValueError: 解密过程中发生异常时抛出，包含具体错误信息
        """
        try:
            encrypted = base64.b64decode(token.encode('utf-8'))
            key = hashlib.sha256(self.secret).digest()[:32]
            decrypted = self._xor_encrypt_decrypt(encrypted, key)
            return decrypted.decode('utf-8')
        except Exception as e:
            raise ValueError(f"解密失败: {e}")

class LicenseValidator:
    """许可证验证类 - 负责许可证文件的创建、读取、解密和有效性验证"""

    def __init__(self):
        """
        初始化许可证验证器。

        设置许可证文件路径（pg_inspector.lic）、试用期天数（36500 天，约 100 年），
        创建 SimpleCrypto 加密实例，并调用 _init_license_system() 确保许可证文件存在。
        """
        self.license_file = "pg_inspector.lic"
        self.trial_days = 36500
        self.crypto = SimpleCrypto()
        self._init_license_system()

    def _parse_datetime(self, date_string):
        """
        解析日期时间字符串，兼容 Python 3.6 及以上版本。

        优先使用 Python 3.7+ 的 datetime.fromisoformat()，
        若不支持则手动解析 ISO 格式（含 'T' 分隔符）或空格分隔格式。

        :param date_string: ISO 格式的日期时间字符串，如 "2026-01-01T00:00:00"
        :return: datetime 对象；解析失败时返回当前时间
        """
        try:
            return datetime.fromisoformat(date_string)
        except AttributeError:
            try:
                if 'T' in date_string:
                    date_part, time_part = date_string.split('T')
                    year, month, day = map(int, date_part.split('-'))
                    time_parts = time_part.split(':')
                    hour, minute = int(time_parts[0]), int(time_parts[1])
                    second = int(time_parts[2].split('.')[0]) if len(time_parts) > 2 else 0
                    return datetime(year, month, day, hour, minute, second)
                else:
                    date_part, time_part = date_string.split(' ')
                    year, month, day = map(int, date_part.split('-'))
                    hour, minute, second = map(int, time_part.split(':'))
                    return datetime(year, month, day, hour, minute, second)
            except Exception as e:
                print(f"日期解析错误: {e}")
                return datetime.now()

    def _format_datetime(self, dt):
        """
        将 datetime 对象格式化为 ISO 格式字符串（兼容 Python 3.6）。

        :param dt: datetime 对象
        :return: 格式为 "YYYY-MM-DDTHH:MM:SS" 的字符串
        """
        return dt.strftime('%Y-%m-%dT%H:%M:%S')

    def _init_license_system(self):
        """
        初始化许可证系统。

        检查许可证文件是否存在，若不存在则自动调用 _create_trial_license() 创建。
        """
        if not os.path.exists(self.license_file):
            self._create_trial_license()

    def _create_trial_license(self):
        """
        创建永久许可证文件（试用期为 100 年）。

        许可证数据包含：类型（PERMANENT）、创建时间、过期时间、机器 ID、数字签名。
        数据序列化为 JSON 后通过 SimpleCrypto 加密，写入 license_file。
        若过期时间年份超过 9999 则自动修正为 9999-12-31 或 2099-12-31。
        """
        create_time = datetime.now()
        try:
            expire_time = create_time + timedelta(days=self.trial_days)
            if expire_time.year > 9999:
                expire_time = datetime(9999, 12, 31)
                print("⚠️  许可证日期超出范围，已调整为9999-12-31")
        except OverflowError:
            expire_time = datetime(2099, 12, 31)
            print("⚠️  许可证日期溢出，已调整为2099-12-31")
        license_data = {
            "type": "PERMANENT",
            "create_time": self._format_datetime(create_time),
            "expire_time": self._format_datetime(expire_time),
            "machine_id": self._get_machine_id(),
            "signature": self._generate_signature("PERMANENT")
        }
        encrypted_data = self.crypto.encrypt(json.dumps(license_data))
        with open(self.license_file, 'w') as f:
            f.write(encrypted_data)

    def _get_machine_id(self):
        """
        获取当前主机的唯一标识符。

        通过组合主机名、操作系统名称和版本号生成字符串，
        再取 MD5 哈希的前 16 位作为机器 ID。

        :return: 16 位十六进制机器 ID 字符串；获取失败时返回 "unknown_machine"
        """
        try:
            machine_info = f"{platform.node()}-{platform.system()}-{platform.release()}"
            return hashlib.md5(machine_info.encode()).hexdigest()[:16]
        except:
            return "unknown_machine"

    def _generate_signature(self, license_type):
        """
        根据许可证类型生成数字签名。

        签名数据由许可证类型、当前日期和固定密钥拼接而成，
        使用 SHA256 哈希生成最终签名。

        :param license_type: 许可证类型字符串，"PERMANENT" 或其他值
        :return: SHA256 哈希的十六进制字符串
        """
        if license_type == "PERMANENT":
            key = "ODB2024PERM"
        else:
            key = "ODB2024TRL"
        signature_data = f"{license_type}-{datetime.now().strftime('%Y%m%d')}-{key}"
        return hashlib.sha256(signature_data.encode()).hexdigest()

    def _verify_signature(self, license_data):
        """
        验证许可证数据中的数字签名是否有效。

        重新生成预期签名并与许可证中的签名对比。

        :param license_data: 解密后的许可证数据字典，需包含 "type" 和 "signature" 字段
        :return: 签名匹配返回 True，不匹配或发生异常返回 False
        """
        try:
            expected_signature = self._generate_signature(license_data["type"])
            return license_data["signature"] == expected_signature
        except:
            return False

    def validate_license(self):
        """
        验证许可证文件的完整性和有效性。

        依次执行：文件存在性检查 → 读取并解密 → 签名验证 → 过期时间检查。

        :return: 三元组 (is_valid, message, remaining_days)
                 - is_valid (bool): 许可证是否有效
                 - message (str): 验证结果描述信息
                 - remaining_days (int): 剩余有效天数（永久版固定返回 99999）
        """
        try:
            if not os.path.exists(self.license_file):
                return False, "许可证文件不存在", 0
            with open(self.license_file, 'r') as f:
                encrypted_data = f.read().strip()
            decrypted_data = self.crypto.decrypt(encrypted_data)
            license_data = json.loads(decrypted_data)
            if not self._verify_signature(license_data):
                return False, "许可证签名无效", 0
            expire_time = self._parse_datetime(license_data["expire_time"])
            remaining_days = (expire_time - datetime.now()).days
            if remaining_days < 0:
                return False, "许可证已过期", 0
            license_type = license_data.get("type", "PERMANENT")
            if license_type == "PERMANENT":
                return True, "永久版许可证有效", 99999
            else:
                return True, f"{license_type}版许可证有效，剩余 {remaining_days} 天", remaining_days
        except Exception as e:
            return False, f"许可证验证失败: {str(e)}", 0

def getlogger():
    """
    获取全局日志记录器。

    创建名为 "pg_check" 的 Logger，日志级别为 INFO，
    若尚未添加处理器则附加一个输出到控制台（StreamHandler）的格式化处理器。
    该函数保证不重复添加处理器（幂等调用安全）。

    :return: 配置完成的 logging.Logger 实例
    """
    logger = logging.getLogger('pg_check')
    logger.setLevel(logging.INFO)
    if not logger.handlers:
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
    return logger

logger = getlogger()

class passArgu(object):
    """命令行参数解析类 - 负责解析程序启动时传入的命令行参数"""

    def get_argus(self):
        """
        解析命令行参数并返回解析结果对象。

        支持以下参数：
        - -C / --sqltemplates: SQL 模板文件路径，默认 templates/sqltemplates.ini
        - -L / --label: 单机巡检时的数据库标签名称
        - -B / --batch: 批量模式开关（store_true）

        :return: argparse.Namespace 对象，包含各参数的值
        """
        all_info = argparse.ArgumentParser(
            description="--example: python3 pg_autoDOC.py -C templates/sqltemplates.ini -L '标签名称'")
        all_info.add_argument('-C', '--sqltemplates', required=False, default='templates/sqltemplates.ini',
                              help='SQL sqltemplates.')
        all_info.add_argument('-L', '--label', required=False, help='Label used when health check single database.')
        all_info.add_argument('-B', '--batch', action='store_true', help='Batch mode (use interactive input for multiple DBs)')
        all_para = all_info.parse_args()
        return all_para

class ExcelTemplateManager:
    """Excel 模板管理器 - 负责批量巡检配置 Excel 模板的创建和读取"""

    def __init__(self):
        """
        初始化 Excel 模板管理器。

        设置默认模板文件名为 "pg_batch_template.xlsx"。
        """
        self.template_file = "pg_batch_template.xlsx"

    def create_template(self):
        """
        创建批量巡检 Excel 配置模板文件。

        模板包含两个工作表：
        - 「PostgreSQL数据库配置」：13 列表头，包含 PostgreSQL 连接信息和 SSH 连接信息，
          并预填 2 行示例数据；密码列用红色字体提醒
        - 「使用说明」：字段说明及注意事项

        :return: 成功返回 True，失败返回 False
        """
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "PostgreSQL数据库配置"
            headers = [
                "序号", "数据库标签", "主机地址", "端口", "用户名", 
                "密码", "数据库名称", 
                "SSH主机", "SSH端口", "SSH用户名", "SSH密码", "SSH密钥文件路径",
                "备注"
            ]
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            column_widths = [8, 20, 15, 8, 15, 20, 15, 15, 8, 15, 20, 25, 20]
            for col, width in enumerate(column_widths, 1):
                ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width
            example_data = [
                [1, "生产数据库", "localhost", 5432, "postgres", "password", "postgres",
                 "localhost", 22, "root", "ssh_password", "/path/to/private_key", "主数据库"],
                [2, "测试数据库", "localhost", 5432, "postgres", "test123", "test_db",
                 "", 22, "", "", "", "测试环境"],
            ]
            for row, data in enumerate(example_data, 2):
                for col, value in enumerate(data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    if col == 6 or col == 11:
                        cell.font = Font(color="FF0000")
            ws2 = wb.create_sheet("使用说明")
            instructions = [
                ["PostgreSQL批量巡检配置模板使用说明"],
                [""],
                ["字段说明:"],
                ["序号", "自动生成的序号，用于标识"],
                ["数据库标签", "用于报告标识的名称，如'生产数据库'"],
                ["主机地址", "PostgreSQL服务器IP地址或主机名"],
                ["端口", "PostgreSQL服务端口，默认5432"],
                ["用户名", "连接数据库的用户名"],
                ["密码", "连接数据库的密码"],
                ["数据库名称", "要连接的数据库名称(可选)"],
                ["SSH主机", "远程主机的IP地址或主机名（用于获取系统信息），如果为空则使用PostgreSQL主机地址"],
                ["SSH端口", "SSH服务端口，默认22"],
                ["SSH用户名", "SSH连接用户名"],
                ["SSH密码", "SSH连接密码（与密钥文件二选一）"],
                ["SSH密钥文件路径", "SSH私钥文件路径（与密码二选一）"],
                ["备注", "额外的说明信息"],
                [""],
                ["注意事项:"],
                ["1. 请确保PostgreSQL服务器允许远程连接"],
                ["2. 建议使用只读权限的账户进行巡检"],
                ["3. 密码会以明文保存，请妥善保管Excel文件"],
                ["4. 支持同时巡检多个不同的PostgreSQL实例"],
                ["5. 如需获取系统信息（CPU、内存、磁盘等），请填写SSH连接信息"],
                ["6. SSH连接支持密码和密钥两种认证方式，优先使用密钥"],
            ]
            for row, instruction in enumerate(instructions, 1):
                if len(instruction) > 1:
                    ws2.cell(row=row, column=1, value=instruction[0])
                    ws2.cell(row=row, column=2, value=instruction[1])
                else:
                    ws2.cell(row=row, column=1, value=instruction[0])
                if row == 1:
                    ws2.cell(row=row, column=1).font = Font(bold=True, size=14)
            wb.save(self.template_file)
            print(_t("pg_cli_excel_created").format(path=self.template_file))
            print(_t("pg_cli_excel_fill_note"))
            print(_t("pg_cli_excel_ssh_note"))
            return True
        except Exception as e:
            print(_t("pg_cli_excel_create_fail").format(e=e))
            return False
    def read_template(self, file_path=None):
        """
        读取 Excel 配置模板，解析并返回数据库连接信息列表。

        从「PostgreSQL数据库配置」工作表的第 2 行开始逐行读取，跳过标签列（B列）为空的行。
        自动处理端口、密码、SSH 信息的默认值及类型转换。
        若 SSH 主机列为空，则默认使用 PostgreSQL 主机地址作为 SSH 主机。

        :param file_path: Excel 文件路径，默认使用 self.template_file
        :return: 数据库配置字典列表，每项包含：
                 name、ip、port、user、password、database、
                 ssh_host、ssh_port、ssh_user、ssh_password、ssh_key_file、remark；
                 文件不存在或读取失败时返回 None
        """
        if file_path is None:
            file_path = self.template_file
        if not os.path.exists(file_path):
            print(_t("pg_cli_excel_not_exist_batch").format(path=file_path))
            return None
        try:
            wb = openpyxl.load_workbook(file_path)
            ws = wb["PostgreSQL数据库配置"]
            db_list = []
            for row in range(2, ws.max_row + 1):
                if not ws.cell(row=row, column=2).value:
                    continue
                port_value = ws.cell(row=row, column=4).value
                if port_value is None:
                    port_value = 5432
                else:
                    try:
                        port_value = int(port_value)
                    except (ValueError, TypeError):
                        port_value = 5432
                password_value = ws.cell(row=row, column=6).value
                if password_value is None:
                    password_value = ""
                else:
                    password_value = str(password_value)
                ssh_host = ws.cell(row=row, column=8).value
                if not ssh_host:
                    ssh_host = str(ws.cell(row=row, column=3).value) if ws.cell(row=row, column=3).value else "localhost"
                else:
                    ssh_host = str(ssh_host)
                ssh_port = ws.cell(row=row, column=9).value
                if ssh_port is None:
                    ssh_port = 22
                else:
                    try:
                        ssh_port = int(ssh_port)
                    except (ValueError, TypeError):
                        ssh_port = 22
                ssh_user = ws.cell(row=row, column=10).value
                if ssh_user is None:
                    ssh_user = "root"
                else:
                    ssh_user = str(ssh_user)
                ssh_password = ws.cell(row=row, column=11).value
                if ssh_password is None:
                    ssh_password = ""
                else:
                    ssh_password = str(ssh_password)
                ssh_key_file = ws.cell(row=row, column=12).value
                if ssh_key_file is None:
                    ssh_key_file = ""
                else:
                    ssh_key_file = str(ssh_key_file)
                db_info = {
                    'name': str(ws.cell(row=row, column=2).value) if ws.cell(row=row, column=2).value else f"DB_{row-1}",
                    'ip': str(ws.cell(row=row, column=3).value) if ws.cell(row=row, column=3).value else "localhost",
                    'port': port_value,
                    'user': str(ws.cell(row=row, column=5).value) if ws.cell(row=row, column=5).value else "postgres",
                    'password': password_value,
                    'database': str(ws.cell(row=row, column=7).value) if ws.cell(row=row, column=7).value else "",
                    'ssh_host': ssh_host,
                    'ssh_port': ssh_port,
                    'ssh_user': ssh_user,
                    'ssh_password': ssh_password,
                    'ssh_key_file': ssh_key_file,
                    'remark': str(ws.cell(row=row, column=13).value) if ws.cell(row=row, column=13).value else ""
                }
                db_list.append(db_info)
            if not db_list:
                print(_t("pg_cli_excel_no_valid_config"))
                return None
            print(_t("pg_cli_excel_read_count").format(n=len(db_list)))
            ssh_count = sum(1 for db in db_list if db['ssh_host'] and (db['ssh_password'] or db['ssh_key_file']))
            print(_t("pg_cli_excel_ssh_count").format(n=ssh_count))
            return db_list
        except Exception as e:
            print(_t("pg_cli_excel_read_fail").format(e=e))
            return None

def input_db_info():
    """
    通过命令行交互式输入数据库连接信息。

    依次提示用户输入：主机地址、端口、用户名、密码、数据库标签名称，
    以及可选的 SSH 连接信息（主机、端口、用户名、认证方式：密码或私钥文件）。
    输入完成后自动验证 PostgreSQL 连接；若配置了 SSH 信息，同时验证 SSH 连接。
    连接验证失败时可选择重新输入或退出。

    :return: 包含数据库连接信息的字典（含 SSH 信息字段）；
             用户放弃输入或连接验证失败且不重试时返回 None
    """
    print("\n" + _t("cli_db_info_title"))
    host = input(_t("cli_db_host").format(default="localhost")).strip() or "localhost"
    port_input = input(_t("cli_db_port").format(default=5432)).strip()
    if not port_input:
        port = '5432'
    else:
        try:
            port = int(port_input)
        except ValueError:
            print(_t("cli_db_port_invalid").format(default=5432))
            port = '5432'
    user = input(_t("cli_db_user").format(default="postgres")).strip() or "postgres"
    import getpass
    password = getpass.getpass(_t("cli_db_password")).strip()
    db_name = input(_t("cli_db_name").format(default="PostgreSQL_Server")).strip() or "PostgreSQL_Server"
    print("\n" + _t("cli_ssh_config_title"))
    print(_t("cli_ssh_config_note"))
    enable_ssh = input(_t("cli_ssh_enable")).strip().lower()
    ssh_info = {}
    if enable_ssh in ['y', 'yes']:
        ssh_host = input(_t("cli_ssh_host").format(default=host)).strip() or host
        ssh_port_input = input(_t("cli_ssh_port").format(default=22)).strip()
        if not ssh_port_input:
            ssh_port = 22
        else:
            try:
                ssh_port = int(ssh_port_input)
            except ValueError:
                print(_t("cli_ssh_port_invalid").format(default=22))
                ssh_port = 22
        ssh_user = input(_t("cli_ssh_user").format(default="root")).strip() or "root"
        auth_choice = input(_t("cli_ssh_auth_method")).strip()
        if auth_choice == '2':
            ssh_key_file = input(_t("cli_ssh_key_path")).strip()
            ssh_password = ""
            if not os.path.exists(ssh_key_file):
                print(_t("cli_ssh_key_not_exist").format(path=ssh_key_file))
                retry = input(_t("cli_retry_yes")).strip().lower()
                if retry in ['', 'y', 'yes']:
                    return input_db_info()
                else:
                    ssh_key_file = ""
        else:
            ssh_password = getpass.getpass(_t("cli_ssh_password")).strip()
            ssh_key_file = ""
        ssh_info = {
            'ssh_host': ssh_host,
            'ssh_port': ssh_port,
            'ssh_user': ssh_user,
            'ssh_password': ssh_password,
            'ssh_key_file': ssh_key_file
        }
    print("\n🔍 " + _t("pg_cli_verifying_pg").format(host=host, port=port))
    try:
        conn = psycopg2.connect(host=host, port=port, user=user, password=password, dbname=db_name, client_encoding='UTF8', connect_timeout=10)
        conn.close()
        print(_t("pg_cli_pg_success").format(host=host, port=port))
    except Exception as e:
        print(_t("pg_cli_pg_fail").format(e=e))
        retry = input(_t("pg_cli_retry_no")).strip().lower()
        if retry == 'y':
            return input_db_info()
        else:
            return None
    if ssh_info:
        print("🔍 " + _t("pg_cli_verifying_ssh").format(host=ssh_info['ssh_host'], port=ssh_info['ssh_port']))
        try:
            collector = RemoteSystemInfoCollector(
                host=ssh_info['ssh_host'], port=ssh_info['ssh_port'], username=ssh_info['ssh_user'],
                password=ssh_info['ssh_password'] if ssh_info['ssh_password'] else None,
                key_file=ssh_info['ssh_key_file'] if ssh_info['ssh_key_file'] else None
            )
            if collector.connect():
                print(_t("cli_ssh_success"))
                collector.disconnect()
            else:
                print(_t("cli_ssh_fail_no_msg"))
        except Exception as e:
            print(_t("cli_ssh_fail").format(e=e))
    db_info = {'name': db_name, 'ip': host, 'port': port, 'user': user, 'password': password, 'database': db_name}
    db_info.update(ssh_info)
    return db_info

def show_main_menu():
    """
    显示程序主菜单并等待用户选择。

    打印 PostgreSQL 数据库巡检工具 的主菜单，
    菜单选项：1 单机巡检、2 批量巡检、3 创建 Excel 模板、4 退出。
    循环接受输入，直到用户输入有效选项（1-4）为止。

    :return: 用户选择的菜单项字符串（"1"/"2"/"3"/"4"）
    """
    print("\n" + "=" * 60)
    print("            " + _t("pg_cli_banner") + " " + VER)
    print("=" * 60)
    print(_t("pg_cli_menu_item1"))
    print(_t("pg_cli_menu_item2"))
    print(_t("pg_cli_menu_item3"))
    print(_t("pg_cli_menu_item4"))
    print("=" * 60)
    while True:
        choice = input(_t("pg_cli_choose_prompt")).strip()
        if choice in ['1', '2', '3', '4']:
            return choice
        else:
            print("\u274c " + _t("pg_cli_invalid_choice"))

class getData(object):
    """数据采集类 - 负责连接 PostgreSQL 数据库并执行全量巡检 SQL，同步采集系统信息和风险分析"""

    def __init__(self, ip, port, user, password, database='postgres', ssh_info=None, label=None):
        """
        初始化数据采集实例并建立 PostgreSQL 连接。

        通过 passArgu 解析命令行参数获取标签名，
        使用 psycopg2 建立数据库连接；连接失败时 conn_db2 置为 None。
        初始化空的 context 字典，用于存储所有巡检结果。

        :param ip: PostgreSQL 服务器 IP 地址或主机名
        :param port: PostgreSQL 服务端口
        :param user: PostgreSQL 登录用户名
        :param password: PostgreSQL 登录密码
        :param ssh_info: SSH 连接信息字典（可选），含 ssh_host、ssh_port、ssh_user、
                         ssh_password、ssh_key_file 字段；为空则使用本地采集模式
        :param label: 巡检标签名（可选）；CLI模式通过 infos.label 传入，
                      直接调用时通过此参数传入
        """
        self.label = str(label if label is not None else infos.label) if 'infos' in dir() else str(label or db_info.get('name', 'pg_inspection'))
        self.H = ip
        self.P = int(port)
        self.user = user
        self.password = password
        self.database = database
        self.ssh_info = ssh_info or {}
        try:
            self.conn_db2 = psycopg2.connect(host=self.H, port=self.P, user=self.user, password=self.password, dbname=self.database, client_encoding='UTF8', connect_timeout=10)
        except Exception as e:
            print(_t("pg_cli_db_conn_fail").format(e=e))
            self.conn_db2 = None
        self.context = {}
        try:
            from i18n import get_lang
            self._lang = get_lang()
        except Exception:
            self._lang = 'zh'

    def _t(self, key):
        try:
            from i18n import t
            return t(key, self._lang)
        except Exception:
            return key

    def print_progress_bar(self, iteration, total, prefix='', suffix='', decimals=1, length=50, fill='█'):
        """
        在终端打印文字版进度条（覆盖当前行）。

        :param iteration: 当前步骤数（从 0 开始）
        :param total: 总步骤数
        :param prefix: 进度条前缀文字
        :param suffix: 进度条后缀文字
        :param decimals: 百分比小数位数，默认 1
        :param length: 进度条字符长度，默认 50
        :param fill: 已完成部分的填充字符，默认 '█'
        """
        percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
        filled_length = int(length * iteration // total)
        bar = fill * filled_length + '-' * (length - filled_length)
        print(f'\r{prefix} |{bar}| {percent}% {suffix}', end='\r')
        if iteration == total:
            print()
    def checkdb(self, sqlfile=''):
        """
        执行 PostgreSQL 数据库健康巡检，采集系统信息并进行风险分析。

        主要流程：
        1. 读取 SQL 模板（内置 builtin 模式 或 外部 .ini 文件）
        2. 逐条执行模板中的 SQL 语句，结果以字典列表形式存入 context
        3. 采集系统信息（SSH 远程或本地 psutil）并规范化磁盘信息格式
        4. 自动分析风险项（连接数使用率 > 80%、内存使用率 > 90%、磁盘使用率 > 90%）
        5. 全程打印进度条

        :param sqlfile: SQL 模板文件路径，传入 'builtin' 时使用内置模板（PG_SQL_TEMPLATES_CONTENT），
                        传入空字符串或文件路径时从文件加载
        :return: 包含所有巡检结果的 context 字典；连接异常或读取模板失败时返回当前已有内容
        """
        print("\n" + _t("pg_cli_starting"))
        total_steps = 15
        current_step = 0
        cfg = configparser.RawConfigParser()
        try:
            if sqlfile == 'builtin':
                cfg.read_string(PG_SQL_TEMPLATES_CONTENT)
            else:
                cfg.read(sqlfile, encoding='utf-8')
        except Exception as e:
            print(_t("pg_cli_sql_template_fail").format(e=e))
            return self.context
        init_keys = ["myversion", "pg_uptime", "pg_connections", "pg_conn_detail",
                    "pg_wait_events", "pg_long_queries", "pg_lock_info", "pg_db_size",
                    "pg_table_stats", "pg_index_usage", "pg_replication", "pg_cache_hit",
                    "pg_bgwriter", "pg_settings_key", "pg_users", "pg_databases",
                    "pg_extensions", "pg_processlist", "instancetime", "platform_info"]
        for key in init_keys:
            self.context.update({key: []})
        try:
            cursor_ver = self.conn_db2.cursor()
            cursor_ver.execute("SELECT version()")
            version_result = cursor_ver.fetchone()
            pg_version_str = version_result[0] if version_result else "Unknown"
            cursor_ver.close()
            self.context.update({"myversion": [{'version': pg_version_str}]})
            self.context.update({"health_summary": [{'health_summary': self._t("report.pg_fallback_health_ok")}]})
        except Exception as e:
            print(_t("pg_cli_version_fail").format(e=e))
            self.context.update({"myversion": [{'version': 'Unknown'}]})
            self.context.update({"health_summary": [{'health_summary': self._t("report.pg_fallback_health_ok")}]})
        try:
            cursor2 = self.conn_db2.cursor()
            variables_items = list(cfg.items("variables"))
            for i, (name, stmt) in enumerate(variables_items):
                try:
                    current_step = int((i / len(variables_items)) * total_steps)
                    self.print_progress_bar(current_step, total_steps, prefix=_t('pg_cli_progress_prefix'), suffix=_t('pg_cli_progress_step').format(i=i+1, total=len(variables_items)))
                    cursor2.execute(stmt.replace('\n', ' ').replace('\r', ' '))
                    result = [dict((cursor2.description[i][0], value) for i, value in enumerate(row)) for row in cursor2.fetchall()]
                    self.context[name] = result
                    time.sleep(0.05)
                except Exception as e:
                    print("\n⚠️  " + _t("pg_cli_step_fail").format(name=name, e=e))
                    self.context[name] = []
                    try:
                        self.conn_db2.rollback()
                    except Exception:
                        pass
                    time.sleep(0.05)
        except Exception as e:
            print("\n❌ " + _t("pg_cli_query_fail").format(e=e))
        finally:
            if 'cursor2' in locals():
                cursor2.close()
        current_step = total_steps - 2
        self.print_progress_bar(current_step, total_steps, prefix=_t('pg_cli_progress_prefix'), suffix=_t('pg_cli_sysinfo_suffix'))
        try:
            if self.ssh_info and self.ssh_info.get('ssh_host'):
                print("\n🔍 " + _t("cli_ssh_collecting").format(host=self.ssh_info['ssh_host']))
                collector = RemoteSystemInfoCollector(
                    host=self.ssh_info['ssh_host'], port=self.ssh_info.get('ssh_port', 22),
                    username=self.ssh_info.get('ssh_user', 'root'),
                    password=self.ssh_info.get('ssh_password'), key_file=self.ssh_info.get('ssh_key_file')
                )
            else:
                print("\n🔍 " + _t("pg_cli_local_sysinfo"))
                collector = LocalSystemInfoCollector()
            system_info = collector.get_system_info()
            if isinstance(system_info.get('disk'), dict):
                disk_list = list(system_info['disk'].values())
                system_info['disk_list'] = disk_list
            elif isinstance(system_info.get('disk'), list):
                system_info['disk_list'] = system_info['disk']
            else:
                disk_info = get_host_disk_usage()
                system_info['disk_list'] = disk_info
            self.context.update({"system_info": system_info})
        except Exception as e:
            print("\n❌ " + _t("pg_cli_sysinfo_fail").format(e=e))
            self.context.update({"system_info": {
                'hostname': '未知', 'platform': '未知', 'boot_time': '未知',
                'cpu': {}, 'memory': {},
                'disk_list': [{'device': '/dev/sda1', 'mountpoint': '/', 'fstype': 'ext4', 'total_gb': 0, 'used_gb': 0, 'free_gb': 0, 'usage_percent': 0}]
            }})
        current_step = total_steps - 1
        self.print_progress_bar(current_step, total_steps, prefix=_t('pg_cli_progress_prefix'), suffix=_t('pg_cli_risk_suffix'))
        self.context.update({"auto_analyze": []})
        try:
            # 使用增强智能分析模块（15+ 条规则）
            try:
                from analyzer import smart_analyze_pg
                issues = smart_analyze_pg(self.context)
                self.context['auto_analyze'] = issues
            except ImportError:
                # 降级：使用内置基础规则
                pg_conn = self.context.get('pg_connections', [])
                if pg_conn and pg_conn[0]:
                    usage_pct = float(pg_conn[0].get('usage_percent', 0))
                    if usage_pct > 80:
                        self.context['auto_analyze'].append({
                            'col1': self._t('report.pg_fallback_conn_usage_label'),
                            "col2": self._t('report.risk_high'),
                            "col3": self._t('report.pg_fallback_conn_usage_desc').format(pct=usage_pct),
                            "col4": self._t('report.pg_fallback_priority_high'),
                            "col5": self._t('report.pg_fallback_owner_dba'),
                            "fix_sql": ""
                        })
                if self.context.get('system_info', {}).get('memory', {}).get('usage_percent', 0) > 90:
                    self.context['auto_analyze'].append({
                        'col1': self._t('report.pg_fallback_mem_usage_label'),
                        'col2': self._t('report.risk_high'),
                        "col3": self._t('report.pg_fallback_mem_usage_desc'),
                        "col4": self._t('report.pg_fallback_priority_high'),
                        "col5": self._t('report.pg_fallback_owner_sysadmin'),
                        "fix_sql": ""
                    })
        except Exception as e:
            print("\n❌ " + _t("pg_cli_risk_fail").format(e=e))

        # AI 智能诊断（从 ai_config.json 读取配置，传递给 analyzer.AIAdvisor）
        self.context['ai_advice'] = ''
        try:
            from analyzer import AIAdvisor
            import json as _json
            cfg_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'ai_config.json')
            ai_cfg = {}
            if os.path.exists(cfg_path):
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    ai_cfg = _json.load(f)
            advisor = AIAdvisor(
                backend=ai_cfg.get('backend'),
                api_key=ai_cfg.get('api_key'),
                api_url=ai_cfg.get('api_url'),
                model=ai_cfg.get('model')
            )
            if advisor.enabled:
                label = self.context.get('co_name', [{}])[0].get('CO_NAME', 'PostgreSQL')
                print("\n🤖 " + _t("pg_cli_ai_calling").format(backend=advisor.backend, model=advisor.model))
                ai_advice = advisor.diagnose('pg', label, self.context, issues, lang=self._lang)
                self.context['ai_advice'] = ai_advice
        except Exception as e:
            self.context['ai_advice'] = ''

        self.print_progress_bar(total_steps, total_steps, prefix=_t('pg_cli_progress_prefix'), suffix=_t('pg_cli_complete_suffix'))
        return self.context

class saveDoc(object):
    """报告保存类 - 将巡检数据渲染到 Word 模板并输出最终报告文件"""

    def __init__(self, context, ofile, ifile, inspector_name="Jack"):
        """
        初始化报告保存实例。

        :param context: 包含所有巡检数据的字典（由 getData.checkdb() 返回）
        :param ofile: 输出的 Word 报告文件路径（目标路径）
        :param ifile: Word 模板文件路径（包含 Jinja2 占位符的 .docx 文件）
        :param inspector_name: 巡检人员姓名，默认 "Jack"
        """
        self.context = context
        self.ofile = ofile
        self.ifile = ifile
        self.inspector_name = inspector_name
        try:
            from i18n import get_lang
            self._lang = get_lang()
        except Exception:
            self._lang = 'zh'

    def _t(self, key):
        try:
            from i18n import t
            return t(key, self._lang)
        except Exception:
            return key

    def _set_cell_bg(self, cell, hex_color):
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        try:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        except Exception:
            pass

    def contextsave(self):
        """
        将巡检数据渲染并保存为 Word 报告文件。

        主要流程：
        1. 补全 context 中缺失的必要字段（health_summary、auto_analyze 等）
        2. 确保 disk_list 存在，否则用默认值填充
        3. 计算问题数量并根据数量评定健康状态（优秀/良好/一般/需关注）
        4. 优先使用 docxtpl 库进行模板渲染；若遇到 AttributeError 或其他异常，
           自动降级调用 _fallback_render() 直接构建报告

        :return: 保存成功返回 True，失败返回 False
        """
        try:
            required_keys = ['health_summary', 'auto_analyze', 'pg_version', 'co_name', 'port', 'ip', 'system_info']
            for key in required_keys:
                if key not in self.context:
                    if key == 'health_summary':
                        self.context[key] = [{'health_summary': self._t('report.running_ok')}]
                    elif key == 'auto_analyze':
                        self.context[key] = []
                    elif key == 'myversion':
                        self.context[key] = [{'version': 'Unknown'}]
                    elif key == 'system_info':
                        self.context[key] = {}
                    else:
                        self.context[key] = [{'placeholder': self._t('report.data_missing')}]

            if 'disk_list' not in self.context['system_info'] or not self.context['system_info']['disk_list']:
                self.context['system_info']['disk_list'] = [{
                    'device': '/dev/sda1', 'mountpoint': '/', 'fstype': 'ext4',
                    'total_gb': 50.0, 'used_gb': 25.0, 'free_gb': 25.0, 'usage_percent': 50.0
                }]

            list_keys = ['pg_db_size', 'pg_processlist', 'pg_users', 'pg_databases', 'pg_extensions', 'pg_cache_hit', 'pg_conn_detail', 'platform_info']
            for key in list_keys:
                if key not in self.context:
                    self.context[key] = []

            self.context.update({"report_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S')})
            self.context.update({"inspector_name": self.inspector_name})
            problem_count = len(self.context.get("auto_analyze", []))
            self.context.update({"problem_count": problem_count})

            if problem_count == 0:
                health_status = self._t("report.health_excellent")
            elif problem_count <= 3:
                health_status = self._t("report.health_good")
            elif problem_count <= 6:
                health_status = self._t("report.health_fair")
            else:
                health_status = self._t("report.health_attention")
            self.context.update({"health_status": health_status})

            # 尝试使用 docxtpl 正常渲染
            try:
                # ── 预处理：翻译 auto_analyze 中的 col1/col2 ───────────────
                _PG_ISSUE_KEY_MAP = {
                    'report.pg_issue_slow_query_log': 'report.pg_issue_slow_query_log',
                    'report.pg_issue_shared_buffers_small': 'report.pg_issue_shared_buffers_small',
                    'report.pg_issue_archive_mode_off': 'report.pg_issue_archive_mode_off',
                    'report.pg_issue_conn_usage_high': 'report.pg_issue_conn_usage_high',
                    'report.pg_issue_mem_usage_high': 'report.pg_issue_mem_usage_high',
                    'report.pg_issue_disk_usage_high': 'report.pg_issue_disk_usage_high',
                    'report.pg_issue_long_query': 'report.pg_issue_long_query',
                    'report.pg_issue_cache_hit_low': 'report.pg_issue_cache_hit_low',
                    'report.pg_issue_lock_wait': 'report.pg_issue_lock_wait',
                    'report.pg_issue_superuser_many': 'report.pg_issue_superuser_many',
                    'report.pg_issue_disk_warning': 'report.pg_issue_disk_warning',
                    'report.pg_issue_dead_tuples': 'report.pg_issue_dead_tuples',
                }
                _RISK_KEY_MAP = {
                    'report.risk_high': 'report.risk_high',
                    'report.risk_mid': 'report.risk_mid',
                    'report.risk_low': 'report.risk_low',
                    'report.risk_suggest': 'report.risk_suggest',
                }
                _PG_OWNER_MAP = {
                    'report.pg_fallback_owner_dba': 'report.pg_fallback_owner_dba',
                    'report.pg_fallback_owner_sysadmin': 'report.pg_fallback_owner_sysadmin',
                }
                _PG_PRIORITY_MAP = {
                    'report.pg_fallback_priority_high': 'report.pg_fallback_priority_high',
                    'report.pg_fallback_priority_mid': 'report.pg_fallback_priority_mid',
                    'report.pg_fallback_priority_low': 'report.pg_fallback_priority_low',
                }
                # col3/fix_sql 中文片段翻译映射（analyzer.py 硬编码的中文描述）
                _COL3_DESC_MAP = {
                    '（约 ': ' (approx. ',
                    ' GB），建议设为物理内存的 25%': ' GB), recommended to set to 25% of physical memory',
                    ' GB），建议设为物理内存的 30%': ' GB), recommended to set to 30% of physical memory',
                    ' GB），建议设为物理内存的 50%': ' GB), recommended to set to 50% of physical memory',
                    '），接近上限将拒绝新连接': '), close to limit, will reject new connections',
                    '），接近上限': '), close to limit',
                    '），建议关注': '), recommend to monitor',
                    '），大量数据从磁盘读取': '), large amount of data read from disk',
                    '（建议 > 99%）': ' (recommend > 99%)',
                    '（建议 > 95%）': ' (recommend > 95%)',
                    '），建议执行 VACUUM': '), recommend running VACUUM',
                    '），建议执行 VACUUM FULL': '), recommend running VACUUM FULL',
                    '），建议最小化权限': '), recommend minimizing privileges',
                    '），超级用户仅用于管理': '), superuser only for administration',
                    '），无法实现 PITR': '), PITR not possible',
                    '，生产环境建议开启': ', recommend enabling in production',
                    '），生产环境建议开启': '), recommend enabling in production',
                    '，生产环境建议开启': ', recommend enabling in production',
                    '），生产环境建议设为 1': '), recommend setting to 1 in production',
                    '），生产环境建议设置合理限制': '), recommend setting reasonable limits in production',
                    '），建议及时清理或扩容': '), recommend cleaning up or expanding capacity',
                    '），建议限制为本地': '), recommend restricting to localhost',
                    '），建议限制为本地或特定 IP': '), recommend restricting to localhost or specific IPs',
                    '），可能频繁开关文件句柄': '), may frequently open/close file handles',
                    '），复杂应用建议设为 500-2000': '), recommend setting to 500-2000 for complex applications',
                    '），避免 ORA-01000 错误': ') to avoid ORA-01000 error',
                    '），合规性要求建议开启': '), compliance requires enabling',
                    '），无法追踪敏感操作': '), cannot trace sensitive operations',
                    '），存在严重安全风险': '), serious security risk',
                    '），存在安全风险': '), security risk',
                    '），已关闭': '), is disabled',
                    '未开启': 'not enabled',
                    '已关闭': 'disabled',
                    '已移除': 'removed',
                    '未设置密码': 'no password set',
                    '允许从任意主机登录': 'allows login from any host',
                    '建议提前关注': 'recommend proactive monitoring',
                    '时间点恢复': 'point-in-time recovery',
                    '无法实现 PITR': 'PITR not possible',
                }
                _FIX_SQL_DESC_MAP = {
                    '-- 修改 postgresql.conf：': '-- Edit postgresql.conf:',
                    '-- 需要重启 PostgreSQL': '-- PostgreSQL restart required',
                    '-- 需要重启数据库': '-- Database restart required',
                    '-- 需要重启生效': '-- Restart required to take effect',
                    '-- 建议同时使用 PgBouncer 连接池': '-- Recommend using PgBouncer connection pool',
                    '-- 增大 shared_buffers（建议物理内存的 25%）：': '-- Increase shared_buffers (recommend 25% of physical memory):',
                    '-- 修改 postgresql.conf': '-- Edit postgresql.conf',
                    '-- 查看超级用户：': '-- View superusers:',
                    '-- 撤销多余超级权限：': '-- Revoke unnecessary superuser privileges:',
                    '-- 查找大表：': '-- Find large tables:',
                    '-- 或全库：': '-- Or for entire database:',
                    '-- 检查大对象：': '-- Check large objects:',
                    '-- 清理方案：': '-- Cleanup plan:',
                    '-- 1) 删除不需要的对象': '-- 1) Delete unnecessary objects',
                    '-- 2) TRUNCATE 大表': '-- 2) TRUNCATE large tables',
                    '-- 3) ALTER TABLESPACE': '-- 3) ALTER TABLESPACE',
                    '-- 查看 TEMP 使用者：': '-- View TEMP users:',
                    '-- 查看会话详情：': '-- View session details:',
                    '-- 调整参数：': '-- Adjust parameters:',
                    '-- 重启后生效': '-- Takes effect after restart',
                    '-- 调整 SGA_TARGET（需重启）：': '-- Adjust SGA_TARGET (requires restart):',
                    '-- 或启用 AMM': '-- Or enable AMM',
                    '-- 查看日志切换频率：': '-- View log switch frequency:',
                    '-- 新增日志组：': '-- Add new log groups:',
                    '-- 开启归档模式（需要重启到 MOUNT 状态）：': '-- Enable archive mode (requires restart to MOUNT state):',
                    '-- 设置归档路径：': '-- Set archive destination:',
                }
                for item in self.context.get('auto_analyze', []):
                    col1 = item.get('col1', '')
                    if col1 in _PG_ISSUE_KEY_MAP:
                        item['col1'] = self._t(_PG_ISSUE_KEY_MAP[col1])
                    col2 = item.get('col2', '')
                    if col2 in _RISK_KEY_MAP:
                        item['col2'] = self._t(_RISK_KEY_MAP[col2])
                    col3 = item.get('col3', '')
                    if col3:
                        # 翻译 col3 中的中文片段（英文环境）
                        if self._lang != 'zh':
                            for zh_frag, en_frag in _COL3_DESC_MAP.items():
                                col3 = col3.replace(zh_frag, en_frag)
                        item['col3'] = col3
                    col4 = item.get('col4', '')
                    if col4 in _PG_PRIORITY_MAP:
                        item['col4'] = self._t(_PG_PRIORITY_MAP[col4])
                    col5 = item.get('col5', '')
                    if col5 in _PG_OWNER_MAP:
                        item['col5'] = self._t(_PG_OWNER_MAP[col5])
                    # 翻译 fix_sql 中的中文注释（英文环境）
                    fix_sql = item.get('fix_sql', '')
                    if fix_sql and self._lang != 'zh':
                        for zh_frag, en_frag in _FIX_SQL_DESC_MAP.items():
                            fix_sql = fix_sql.replace(zh_frag, en_frag)
                        item['fix_sql'] = fix_sql
                # ── 渲染 ──────────────────────────────────────────────
                with open(self.ifile, 'rb') as f:
                    template_bytes = f.read()
                doc_stream = io.BytesIO(template_bytes)
                tpl = DocxTemplate(doc_stream)
                tpl.render(self.context)
                tpl.save(self.ofile)

                # ── 追加新章节（第7章 7.1/7.2 + 第8章 AI诊断）───────────────────
                # docxtpl 模板本身有旧的"7.报告说明"，先把它及之后的内容删掉，再追加新章节
                doc2 = Document(self.ofile)
                cutoff_idx = None
                for i, para in enumerate(doc2.paragraphs):
                    t = para.text.strip()
                    if t.startswith('7.') and (self._t('report.notes_chapter') in t or '报告说明' in t or 'Report Notes' in t):
                        cutoff_idx = i
                        break
                if cutoff_idx is not None:
                    body = doc2._element.body
                    for para in list(body.iterchildren())[cutoff_idx:]:
                        body.remove(para)

                auto_analyze = self.context.get('auto_analyze', [])
                high_risk = [i for i in auto_analyze if i.get('col2') == self._t('report.risk_high')]
                mid_risk  = [i for i in auto_analyze if i.get('col2') == self._t('report.risk_mid')]
                low_risk  = [i for i in auto_analyze if i.get('col2') in (self._t('report.risk_low'), self._t('report.risk_suggest'))]

                # 第 7 章 风险与建议
                doc2.add_heading('7. ' + self._t('report.risk_chapter'), level=1)
                p = doc2.add_paragraph()
                p.add_run(self._t('report.detected_prefix'))
                if high_risk:
                    r = p.add_run(self._t('report.high_risk_n').format(n=len(high_risk))); r.bold = True; r.font.color.rgb = RGBColor(0xC0,0x00,0x00)
                if mid_risk:
                    r = p.add_run(self._t('report.mid_risk_n').format(n=len(mid_risk))); r.bold = True; r.font.color.rgb = RGBColor(0xFF,0x78,0x00)
                if low_risk:
                    r = p.add_run(self._t('report.low_risk_n').format(n=len(low_risk))); r.bold = True; r.font.color.rgb = RGBColor(0x37,0x86,0x10)
                p.add_run(self._t('report.detected_suffix').format(c=len(auto_analyze)))

                # 7.1 问题明细
                if auto_analyze:
                    doc2.add_heading('7.1 ' + self._t('report.risk_detail_chapter'), level=2)
                    tbl = doc2.add_table(rows=1+len(auto_analyze), cols=7)
                    tbl.style = 'Table Grid'
                    tbl.autofit = True  # 根据窗口自动调整表格宽度
                    hdrs = [self._t('report.col_seq'), self._t('report.col_risk_item'), self._t('report.col_level'), self._t('report.col_desc'), self._t('report.col_priority'), self._t('report.col_owner'), self._t('report.col_fix')]
                    for j,(cell,ht) in enumerate(zip(tbl.rows[0].cells, hdrs)):
                        cell.text = ht
                        self._set_cell_bg(cell, '336699')
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(9)
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for idx,item in enumerate(auto_analyze,1):
                        row = tbl.rows[idx].cells
                        row[0].text = str(idx)
                        row[1].text = item.get('col1','')
                        row[2].text = item.get('col2','')
                        row[3].text = item.get('col3','')
                        row[4].text = item.get('col4','')
                        row[5].text = item.get('col5','')
                        fix_sql = item.get('fix_sql','').strip()
                        row[6].text = fix_sql if fix_sql else self._t('report.fallback_fix_sql_placeholder')
                        for j,cell in enumerate(row):
                            for para in cell.paragraphs:
                                for run in para.runs: run.font.size = Pt(9)
                        lvl = item.get('col2','')
                        cm = {
                            self._t('report.risk_high'): RGBColor(0xC0,0x00,0x00),
                            self._t('report.risk_mid'): RGBColor(0xFF,0x78,0x00),
                            self._t('report.risk_low'): RGBColor(0x37,0x86,0x10),
                            self._t('report.risk_suggest'): RGBColor(0x00,0x70,0xC0)
                        }
                        if lvl in cm:
                            row[2].paragraphs[0].runs[0].font.color.rgb = cm[lvl]
                            row[2].paragraphs[0].runs[0].bold = True
                else:
                    doc2.add_paragraph(self._t('report.no_risk_found_pg'))

                # 7.2 修复速查
                fix_items = [i for i in auto_analyze if i.get('fix_sql','').strip()]
                if fix_items:
                    doc2.add_heading('7.2 ' + self._t('report.fix_chapter'), level=2)
                    for idx,item in enumerate(fix_items,1):
                        p = doc2.add_paragraph()
                        p.add_run(f'{idx}. [{item.get("col1")}] {item.get("col3","")[:60]}').bold = True
                        qp = doc2.add_paragraph(item.get('fix_sql','').strip())
                        qp.style = 'Quote'
                        if qp.runs: qp.runs[0].font.size = Pt(9)

                # 第 8 章 AI 智能诊断建议
                ai_advice = self.context.get('ai_advice','').strip()
                doc2.add_heading('8. ' + self._t('report.ai_chapter'), level=1)
                if ai_advice:
                    p = doc2.add_paragraph()
                    p.add_run(self._t('report.ai_disclaimer')).italic = True
                    doc2.add_paragraph()
                    for line in ai_advice.split('\n'):
                        line = line.strip()
                        if not line:
                            doc2.add_paragraph()
                        elif line.startswith(('- ','* ','• ')):
                            bp = doc2.add_paragraph(style='List Bullet')
                            bp.add_run(line[2:]).font.size = Pt(11)
                        else:
                            np = doc2.add_paragraph(line)
                            if np.runs: np.runs[0].font.size = Pt(11)
                else:
                    p = doc2.add_paragraph()
                    p.add_run(self._t('report.ai_disabled')).italic = True

                # 第 9 章 报告说明
                doc2.add_heading('9. ' + self._t('report.notes_chapter'), level=1)
                notes = [
                    self._t("report.note_1_pg"),
                    self._t("report.fallback_note_2"),
                    self._t("report.fallback_note_3"),
                    self._t("report.fallback_note_4"),
                    self._t("report.fallback_note_5"),
                    self._t("report.fallback_note_6")
                ]
                for note in notes:
                    doc2.add_paragraph(note)

                doc2.save(self.ofile)
                return True
            except AttributeError as ae:
                if 'part' in str(ae):
                    pass  # 静默降级到备用渲染
                    return self._fallback_render()
                else:
                    raise
            except Exception as e:
                pass  # 静默降级到备用渲染
                return self._fallback_render()

        except Exception as e:
            print("❌ " + _t("pg_cli_word_fail_gen").format(e=e))
            import traceback
            traceback.print_exc()
            return False

    def _fallback_render(self):
        """
        备用报告渲染方法 - 当 docxtpl 渲染失败时使用 python-docx 直接构建报告。

        不依赖模板文件，直接通过代码构建包含 8 个章节的完整 Word 文档：
        封面、健康状态概览、系统资源检查（CPU/内存/磁盘）、
        PostgreSQL 配置检查（连接/内存/日志）、性能分析（QPS/锁/异常连接）、
        数据库信息（大小/进程列表）、安全信息（用户信息）、风险与建议、报告说明。

        所有数据均从 self.context 中直接提取，无 Jinja2 模板变量。

        :return: 渲染并保存成功返回 True，失败返回 False
        """
        # col3/fix_sql 中文片段翻译映射（与 contextsave 保持一致）
        _COL3_DESC_MAP = {
            '（约 ': ' (approx. ',
            ' GB），建议设为物理内存的 25%': ' GB), recommended to set to 25% of physical memory',
            ' GB），建议设为物理内存的 30%': ' GB), recommended to set to 30% of physical memory',
            ' GB），建议设为物理内存的 50%': ' GB), recommended to set to 50% of physical memory',
            '），接近上限将拒绝新连接': '), close to limit, will reject new connections',
            '），接近上限': '), close to limit',
            '），建议关注': '), recommend to monitor',
            '），大量数据从磁盘读取': '), large amount of data read from disk',
            '（建议 > 99%）': ' (recommend > 99%)',
            '（建议 > 95%）': ' (recommend > 95%)',
            '），建议执行 VACUUM': '), recommend running VACUUM',
            '），建议执行 VACUUM FULL': '), recommend running VACUUM FULL',
            '），建议最小化权限': '), recommend minimizing privileges',
            '），超级用户仅用于管理': '), superuser only for administration',
            '），无法实现 PITR': '), PITR not possible',
            '，生产环境建议开启': ', recommend enabling in production',
            '），生产环境建议开启': '), recommend enabling in production',
            '），生产环境建议设为 1': '), recommend setting to 1 in production',
            '），生产环境建议设置合理限制': '), recommend setting reasonable limits in production',
            '），建议及时清理或扩容': '), recommend cleaning up or expanding capacity',
            '），建议限制为本地': '), recommend restricting to localhost',
            '），建议限制为本地或特定 IP': '), recommend restricting to localhost or specific IPs',
            '），可能频繁开关文件句柄': '), may frequently open/close file handles',
            '），复杂应用建议设为 500-2000': '), recommend setting to 500-2000 for complex applications',
            '），避免 ORA-01000 错误': ') to avoid ORA-01000 error',
            '），合规性要求建议开启': '), compliance requires enabling',
            '），无法追踪敏感操作': '), cannot trace sensitive operations',
            '），存在严重安全风险': '), serious security risk',
            '），存在安全风险': '), security risk',
            '），已关闭': '), is disabled',
            '未开启': 'not enabled',
            '已关闭': 'disabled',
            '已移除': 'removed',
            '未设置密码': 'no password set',
            '允许从任意主机登录': 'allows login from any host',
            '建议提前关注': 'recommend proactive monitoring',
            '时间点恢复': 'point-in-time recovery',
            '无法实现 PITR': 'PITR not possible',
        }
        _FIX_SQL_DESC_MAP = {
            '-- 修改 postgresql.conf：': '-- Edit postgresql.conf:',
            '-- 需要重启 PostgreSQL': '-- PostgreSQL restart required',
            '-- 需要重启数据库': '-- Database restart required',
            '-- 需要重启生效': '-- Restart required to take effect',
            '-- 建议同时使用 PgBouncer 连接池': '-- Recommend using PgBouncer connection pool',
            '-- 增大 shared_buffers（建议物理内存的 25%）：': '-- Increase shared_buffers (recommend 25% of physical memory):',
            '-- 修改 postgresql.conf': '-- Edit postgresql.conf',
            '-- 查看超级用户：': '-- View superusers:',
            '-- 撤销多余超级权限：': '-- Revoke unnecessary superuser privileges:',
            '-- 查找大表：': '-- Find large tables:',
            '-- 或全库：': '-- Or for entire database:',
            '-- 检查大对象：': '-- Check large objects:',
            '-- 清理方案：': '-- Cleanup plan:',
            '-- 1) 删除不需要的对象': '-- 1) Delete unnecessary objects',
            '-- 2) TRUNCATE 大表': '-- 2) TRUNCATE large tables',
            '-- 3) ALTER TABLESPACE': '-- 3) ALTER TABLESPACE',
            '-- 查看 TEMP 使用者：': '-- View TEMP users:',
            '-- 查看会话详情：': '-- View session details:',
            '-- 调整参数：': '-- Adjust parameters:',
            '-- 重启后生效': '-- Takes effect after restart',
            '-- 调整 SGA_TARGET（需重启）：': '-- Adjust SGA_TARGET (requires restart):',
            '-- 或启用 AMM': '-- Or enable AMM',
            '-- 查看日志切换频率：': '-- View log switch frequency:',
            '-- 新增日志组：': '-- Add new log groups:',
            '-- 开启归档模式（需要重启到 MOUNT 状态）：': '-- Enable archive mode (requires restart to MOUNT state):',
            '-- 设置归档路径：': '-- Set archive destination:',
        }
        try:
            doc = Document()
            title = doc.add_heading(self._t('report.pg_title'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.size = Pt(20)
            title_run.font.bold = True
            doc.add_paragraph()
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(10)
            data_map = [
                (self._t("report.fallback_db_name"), self.context.get('co_name', [{}])[0].get('CO_NAME', 'N/A')),
                (self._t("report.fallback_server_addr"), f"{self.context.get('ip', [{}])[0].get('IP', 'N/A')}:{self.context.get('port', [{}])[0].get('PORT', 'N/A')}"),
                (self._t("report.fallback_pg_version"), self.context.get('myversion', [{}])[0].get('version', 'N/A')),
                (self._t("report.fallback_hostname"), self.context.get('system_info', {}).get('hostname', 'N/A')),
                (self._t("report.fallback_pg_start_time"), self.context.get('pg_uptime', [{}])[0].get('started_at', 'N/A') if self.context.get('pg_uptime') else 'N/A'),
                (self._t("report.fallback_inspector"), self.inspector_name),
                (self._t("report.fallback_platform"), self._get_platform_info()),
                (self._t("report.fallback_report_time"), self.context.get('report_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
            ]
            for i, (label, value) in enumerate(data_map):
                cells = table.rows[i].cells
                cells[0].text = label
                cells[1].text = str(value)
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
            doc.add_page_break()
            doc.add_heading('1. ' + self._t('report.fallback_health_overview'), level=1)
            col_w = [Cm(4), Cm(10)]
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr_texts = [self._t("report.fallback_item_col"), self._t("report.fallback_value_col")]
            for j, (cell, ht) in enumerate(zip(hdr, hdr_texts)):
                cell.text = ht
                self._set_cell_bg(cell, '336699')
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.width = col_w[j]
            # 第1行数据
            cells = table.rows[1].cells
            cells[0].text = self._t("report.fallback_overall_health")
            cells[1].text = self.context.get('health_status', 'N/A')
            for j, c in enumerate(cells):
                c.width = col_w[j]
                for para in c.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
            # 第2行数据
            cells = table.rows[2].cells
            cells[0].text = self._t("report.fallback_issue_count")
            cells[1].text = f"{self.context.get('problem_count', 0)}"
            for j, c in enumerate(cells):
                c.width = col_w[j]
                for para in c.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
            doc.add_paragraph()
            p = doc.add_paragraph(self._t("report.fallback_health_summary") + ": ")
            p.add_run(self.context.get('health_summary', [{}])[0].get('health_summary', self._t('report.running_ok'))).bold = True

            doc.add_heading('2. ' + self._t('report.fallback_system_check'), level=1)
            cpu = self.context.get('system_info', {}).get('cpu', {})
            mem = self.context.get('system_info', {}).get('memory', {})
            doc.add_heading('2.1 ' + self._t('report.fallback_cpu_info'), level=2)
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = self._t('report.fallback_cpu_usage')
            hdr[1].text = self._t('report.fallback_physical_cores')
            hdr[2].text = self._t('report.fallback_logical_cores')
            hdr[3].text = self._t('report.fallback_freq_ghz')
            row = table.rows[1].cells
            row[0].text = f"{cpu.get('usage_percent', 'N/A')}%"
            row[1].text = str(cpu.get('physical_cores', 'N/A'))
            row[2].text = str(cpu.get('logical_cores', 'N/A'))
            freq = cpu.get('current_frequency', 0)
            row[3].text = f"{freq/1000:.2f}" if isinstance(freq, (int, float)) and freq > 100 else str(freq)
            doc.add_paragraph()
            doc.add_heading('2.2 ' + self._t('report.fallback_memory_info'), level=2)
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = self._t('report.fallback_total_gb')
            hdr[1].text = self._t('report.fallback_used_gb')
            hdr[2].text = self._t('report.fallback_available_gb')
            hdr[3].text = self._t('report.fallback_usage_pct')
            row = table.rows[1].cells
            row[0].text = f"{mem.get('total_gb', 'N/A')}"
            row[1].text = f"{mem.get('used_gb', 'N/A')}"
            row[2].text = f"{mem.get('available_gb', 'N/A')}"
            row[3].text = f"{mem.get('usage_percent', 'N/A')}%"
            doc.add_paragraph()
            doc.add_heading('2.3 ' + self._t('report.fallback_disk_info'), level=2)
            disk_list = self.context.get('system_info', {}).get('disk_list', [])
            table = doc.add_table(rows=1+len(disk_list), cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Cm(8)
            table.columns[1].width = Cm(4)
            hdr = table.rows[0].cells
            hdr[0].text = self._t('report.fallback_mountpoint')
            hdr[1].text = self._t('report.fallback_usage_pct')
            for i, disk in enumerate(disk_list, 1):
                cells = table.rows[i].cells
                cells[0].text = disk.get('mountpoint', 'N/A')
                cells[1].text = f"{disk.get('usage_percent', 0):.2f}%"

            doc.add_heading('3. ' + self._t('report.fallback_pg_config'), level=1)
            # 从 pg_settings_key 列表中提取配置值
            pg_sk = self.context.get('pg_settings_key', [])
            def _pg_setting(name):
                for item in pg_sk:
                    if item.get('name') == name:
                        return item.get('setting', 'N/A')
                return 'N/A'

            doc.add_heading('3.1 ' + self._t('report.fallback_pg_conn_mem'), level=2)
            pg_conn = self.context.get('pg_connections', [])
            conn_total = pg_conn[0].get('total_connections', 'N/A') if pg_conn else 'N/A'
            conn_max = pg_conn[0].get('max_connections', 'N/A') if pg_conn else 'N/A'
            self._add_config_table_pg(doc, [
                (self._t('report.fallback_pg_max_conn'), conn_max),
                (self._t('report.fallback_pg_cur_conn'), conn_total),
                (self._t('report.fallback_pg_shared_buf'), _pg_setting('shared_buffers')),
                (self._t('report.fallback_pg_work_mem'), _pg_setting('work_mem')),
                (self._t('report.fallback_pg_maint_mem'), _pg_setting('maintenance_work_mem')),
                (self._t('report.fallback_pg_eff_cache'), _pg_setting('effective_cache_size')),
            ])
            doc.add_heading('3.2 ' + self._t('report.fallback_pg_wal_checkpoint'), level=2)
            self._add_config_table_pg(doc, [
                (self._t('report.fallback_pg_wal_level'), _pg_setting('wal_level')),
                (self._t('report.fallback_pg_max_wal'), _pg_setting('max_wal_size')),
                (self._t('report.fallback_pg_ckpt_target'), _pg_setting('checkpoint_completion_target')),
                (self._t('report.fallback_pg_random_page'), _pg_setting('random_page_cost')),
            ])
            doc.add_heading('3.3 ' + self._t('report.fallback_pg_autovacuum'), level=2)
            self._add_config_table_pg(doc, [
                (self._t('report.fallback_pg_autovac'), _pg_setting('autovacuum')),
                (self._t('report.fallback_pg_vac_scale'), _pg_setting('autovacuum_vacuum_scale_factor')),
                (self._t('report.fallback_pg_analyze_scale'), _pg_setting('autovacuum_analyze_scale_factor')),
                (self._t('report.fallback_pg_slow_query'), _pg_setting('log_min_duration_statement')),
            ])

            doc.add_heading('4. ' + self._t('report.fallback_pg_perf'), level=1)
            doc.add_heading('4.1 ' + self._t('report.fallback_pg_conn_status'), level=2)
            self._add_config_table_pg(doc, [
                (self._t('report.fallback_pg_cur_conn'), conn_total),
                (self._t('report.fallback_pg_max_conn'), conn_max),
                (self._t('report.fallback_pg_conn_usage'), pg_conn[0].get('usage_percent', 'N/A') if pg_conn else 'N/A'),
            ])
            doc.add_heading('4.2 ' + self._t('report.fallback_pg_cache_hit'), level=2)
            pg_ch = self.context.get('pg_cache_hit', [])
            cache_rows = [(self._t('report.fallback_pg_dbname'), self._t('report.fallback_pg_cache_hit'))] + [(d.get('datname', ''), d.get('cache_hit_ratio', '')) for d in pg_ch[:8]]
            for row_data in cache_rows:
                p = doc.add_paragraph(f"{row_data[0]}: {row_data[1]}%")
                if p.runs:
                    p.runs[0].font.size = Pt(10)
            doc.add_heading('4.3 ' + self._t('report.fallback_pg_db_size'), level=2)
            pg_db = self.context.get('pg_db_size', [])
            for d in pg_db[:8]:
                p = doc.add_paragraph(f"  {d.get('database_name', '')}: {d.get('size', '')}")
                if p.runs:
                    p.runs[0].font.size = Pt(10)

            doc.add_heading('5. ' + self._t('report.fallback_pg_db_info'), level=1)
            doc.add_heading('5.1 ' + self._t('report.fallback_pg_db_list'), level=2)
            pg_db2 = self.context.get('pg_databases', [])
            table = doc.add_table(rows=1+min(len(pg_db2), 10), cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = self._t('report.fallback_pg_dbname')
            hdr[1].text = self._t('report.fallback_pg_encoding')
            hdr[2].text = self._t('report.fallback_pg_allow_conn')
            for i, db in enumerate(pg_db2[:10], 1):
                cells = table.rows[i].cells
                cells[0].text = str(db.get('datname', ''))
                cells[1].text = str(db.get('encoding', ''))
                cells[2].text = str(db.get('datallowconn', ''))
                for cell in cells:
                    for par in cell.paragraphs:
                        for r in par.runs:
                            r.font.size = Pt(10)
            doc.add_heading('5.2 ' + self._t('report.fallback_pg_proc_list'), level=2)
            proc = self.context.get('pg_processlist', [])
            table = doc.add_table(rows=1+min(len(proc), 15), cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = self._t('report.fallback_pg_pid')
            hdr[1].text = self._t('report.fallback_pg_pg_user')
            hdr[2].text = self._t('report.fallback_pg_pg_db')
            hdr[3].text = self._t('report.fallback_pg_pg_state')
            hdr[4].text = self._t('report.fallback_pg_curr_sql')
            for i, p in enumerate(proc[:15], 1):
                cells = table.rows[i].cells
                cells[0].text = str(p.get('pid', ''))
                cells[1].text = str(p.get('usename', ''))
                cells[2].text = str(p.get('datname', ''))
                cells[3].text = str(p.get('state', ''))
                cells[4].text = str(p.get('query', ''))[:80]
                for cell in cells:
                    for par in cell.paragraphs:
                        for r in par.runs:
                            r.font.size = Pt(9)

            doc.add_heading('6. ' + self._t('report.fallback_pg_security'), level=1)
            doc.add_heading('6.1 ' + self._t('report.fallback_pg_db_users'), level=2)
            users = self.context.get('pg_users', [])
            table = doc.add_table(rows=1+min(len(users), 15), cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = self._t('report.fallback_pg_username')
            hdr[1].text = self._t('report.fallback_pg_superuser')
            hdr[2].text = self._t('report.fallback_pg_can_createdb')
            hdr[3].text = self._t('report.fallback_pg_can_createrole')
            hdr[4].text = self._t('report.fallback_pg_pass_expire')
            for i, u in enumerate(users[:15], 1):
                cells = table.rows[i].cells
                cells[0].text = str(u.get('username', ''))
                cells[1].text = str(u.get('is_superuser', ''))
                cells[2].text = str(u.get('can_createdb', ''))
                cells[3].text = str(u.get('can_createrole', ''))
                cells[4].text = str(u.get('password_expiry', ''))
                for cell in cells:
                    for par in cell.paragraphs:
                        for r in par.runs:
                            r.font.size = Pt(9)

            doc.add_heading('7. ' + self._t('report.fallback_pg_risk_chapter'), level=1)

            # ── 7.1 概览统计 ──
            auto_analyze = self.context.get('auto_analyze', [])
            high_risk  = [i for i in auto_analyze if i.get('col2') == self._t('report.risk_high')]
            mid_risk   = [i for i in auto_analyze if i.get('col2') == self._t('report.risk_mid')]
            low_risk   = [i for i in auto_analyze if i.get('col2') in (self._t('report.risk_low'), self._t('report.risk_suggest'))]
            p = doc.add_paragraph()
            p.add_run(self._t('report.fallback_detected_prefix')).bold = False
            if high_risk:
                run_h = p.add_run(f'{len(high_risk)}' + self._t('report.fallback_high_risk_n'))
                run_h.bold = True; run_h.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            if mid_risk:
                run_m = p.add_run(f' {len(mid_risk)}' + self._t('report.fallback_mid_risk_n'))
                run_m.bold = True; run_m.font.color.rgb = RGBColor(0xFF, 0x78, 0x00)
            if low_risk:
                run_l = p.add_run(f' {len(low_risk)}' + self._t('report.fallback_low_risk_n'))
                run_l.bold = True; run_l.font.color.rgb = RGBColor(0x37, 0x86, 0x10)
            p.add_run(self._t('report.fallback_detected_suffix2').format(c=len(auto_analyze)))

            # ── 7.2 风险明细表格 ──
            if auto_analyze:
                doc.add_heading('7.1 ' + self._t('report.fallback_pg_issue_detail'), level=2)
                tbl = doc.add_table(rows=1 + len(auto_analyze), cols=7)
                tbl.style = 'Table Grid'
                tbl.autofit = True  # 根据窗口自动调整表格宽度
                hdr = tbl.rows[0].cells
                headers = [
                    self._t('report.fallback_seq'),
                    self._t('report.fallback_risk_item'),
                    self._t('report.fallback_level'),
                    self._t('report.fallback_desc'),
                    self._t('report.fallback_priority'),
                    self._t('report.fallback_owner'),
                    self._t('report.fallback_fix_suggest')
                ]
                for j, (cell, hdr_text) in enumerate(zip(hdr, headers)):
                    cell.text = hdr_text
                    self._set_cell_bg(cell, '336699')
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for idx, item in enumerate(auto_analyze, 1):
                    row = tbl.rows[idx].cells
                    row[0].text = str(idx)
                    # 翻译 col1/col2/col4/col5（analyzer.py 使用翻译 key）
                    col1_val = item.get('col1', '')
                    if col1_val.startswith('report.'):
                        col1_val = self._t(col1_val)
                    row[1].text = col1_val
                    col2_val = item.get('col2', '')
                    if col2_val.startswith('report.'):
                        col2_val = self._t(col2_val)
                    row[2].text = col2_val
                    # 翻译 col3（analyzer.py 硬编码的中文描述）
                    col3_val = item.get('col3', '')
                    if col3_val and self._lang != 'zh':
                        for zh_frag, en_frag in _COL3_DESC_MAP.items():
                            col3_val = col3_val.replace(zh_frag, en_frag)
                    row[3].text = col3_val
                    col4_val = item.get('col4', '')
                    if col4_val.startswith('report.'):
                        col4_val = self._t(col4_val)
                    row[4].text = col4_val
                    col5_val = item.get('col5', '')
                    if col5_val.startswith('report.'):
                        col5_val = self._t(col5_val)
                    row[5].text = col5_val
                    fix_sql = item.get('fix_sql', '').strip()
                    if fix_sql and self._lang != 'zh':
                        for zh_frag, en_frag in _FIX_SQL_DESC_MAP.items():
                            fix_sql = fix_sql.replace(zh_frag, en_frag)
                    row[6].text = fix_sql if fix_sql else self._t('report.fallback_fix_sql_placeholder')
                    for j, cell in enumerate(row):
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)
                    # col2 is already translated, compare with translated risk strings
                    color_map = {
                        self._t('report.risk_high'): RGBColor(0xC0, 0x00, 0x00),
                        self._t('report.risk_mid'): RGBColor(0xFF, 0x78, 0x00),
                        self._t('report.risk_low'): RGBColor(0x37, 0x86, 0x10),
                        self._t('report.risk_suggest'): RGBColor(0x00, 0x70, 0xC0)
                    }
                    if col2_val in color_map:
                        row[2].paragraphs[0].runs[0].font.color.rgb = color_map[col2_val]
                        row[2].paragraphs[0].runs[0].bold = True
            else:
                doc.add_paragraph(self._t('report.fallback_pg_no_risk_found'))

            # ── 7.3 修复速查 ──
            fix_items = [i for i in auto_analyze if i.get('fix_sql', '').strip()]
            if fix_items:
                doc.add_heading('7.2 ' + self._t('report.fallback_pg_fix_chapter'), level=2)
                for idx, item in enumerate(fix_items, 1):
                    p = doc.add_paragraph()
                    col1_val = item.get('col1', '')
                    if col1_val.startswith('report.'):
                        col1_val = self._t(col1_val)
                    col3_raw = item.get('col3', '')[:60]
                    if self._lang != 'zh':
                        for zh_frag, en_frag in _COL3_DESC_MAP.items():
                            col3_raw = col3_raw.replace(zh_frag, en_frag)
                    p.add_run(f'{idx}. [{col1_val}] {col3_raw}').bold = True
                    fix_sql_raw = item.get('fix_sql', '').strip()
                    if self._lang != 'zh':
                        for zh_frag, en_frag in _FIX_SQL_DESC_MAP.items():
                            fix_sql_raw = fix_sql_raw.replace(zh_frag, en_frag)
                    code_p = doc.add_paragraph(fix_sql_raw)
                    code_p.style = 'Quote'
                    if code_p.runs:
                        code_p.runs[0].font.size = Pt(9)

            # ── 8. AI 智能诊断建议 ──
            ai_advice = self.context.get('ai_advice', '').strip()
            doc.add_heading('8. ' + self._t('report.fallback_pg_ai_chapter'), level=1)
            if ai_advice:
                p = doc.add_paragraph()
                p.add_run(self._t('report.fallback_pg_ai_disclaimer')).italic = True
                doc.add_paragraph()
                for line in ai_advice.split('\n'):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                    elif line.startswith(('- ', '* ', '• ')):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(line[2:]).font.size = Pt(11)
                    else:
                        if doc.add_paragraph(line).runs:
                            doc.paragraphs[-1].runs[0].font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                p.add_run(self._t('report.ai_disabled')).italic = True

            # ── 9. 报告说明 ──
            doc.add_heading('9. ' + self._t('report.fallback_pg_notes_chapter'), level=1)
            notes = [
                self._t("report.fallback_pg_note_1"),
                self._t("report.fallback_pg_note_2"),
                self._t("report.fallback_pg_note_3"),
                self._t("report.fallback_pg_note_4"),
                self._t("report.fallback_pg_note_5"),
                self._t("report.fallback_pg_note_6")
            ]
            for note in notes:
                doc.add_paragraph(note)

            doc.save(self.ofile)
            pass  # 备用渲染成功
            return True
        except Exception as e:
            print(f"{self._t('report.fallback_pg_render_fail')}: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _get_platform_info(self):
        """
        从 context 中提取服务器操作系统平台信息。

        遍历 platform_info 列表，查找 variable_name 为 'version_compile_os' 的条目。

        :return: 操作系统编译平台字符串（如 "Linux"），未找到时返回 'N/A'
        """
        pf = self.context.get('platform_info', [])
        for item in pf:
            if item.get('variable_name') == 'version_compile_os':
                return item.get('variable_value', 'N/A')
        return 'N/A'

    def _add_config_table_pg(self, doc, items, col1_width=4, col2_width=10):
        """
        向 Word 文档中添加 PostgreSQL 配置键值表格（备用渲染辅助方法）。

        创建「配置项 / 当前值」二列表格，直接使用传入的值填充。
        """
        table = doc.add_table(rows=1+len(items), cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(col1_width)
        table.columns[1].width = Cm(col2_width)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = '配置项', '当前值'
        for i, (label, value) in enumerate(items, 1):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = str(value) if value else 'N/A'
            for cell in cells:
                for par in cell.paragraphs:
                    for r in par.runs:
                        r.font.size = Pt(10)
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_config_table(self, doc, items, col1_width=4, col2_width=10):
        """
        向 Word 文档中添加配置键值表格（通用辅助方法）。

        创建「配置项 / 当前值」二列表格，并从 pg_settings_key 中查找对应配置项的值。
        数据不存在时填入 'N/A'。

        :param doc: 目标 Document 对象
        :param items: 配置项列表，每项为 (显示标签, pg_setting_name) 元组
        :param col1_width: 第 1 列宽度（cm），默认 4
        :param col2_width: 第 2 列宽度（cm），默认 10
        """
        pg_sk = self.context.get('pg_settings_key', [])
        def _get_setting(name):
            for item in pg_sk:
                if item.get('name') == name:
                    return item.get('setting', 'N/A')
            return 'N/A'
        table = doc.add_table(rows=1+len(items), cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(col1_width)
        table.columns[1].width = Cm(col2_width)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = '配置项', '当前值'
        for i, (label, setting_name) in enumerate(items, 1):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = _get_setting(setting_name)
            for cell in cells:
                for par in cell.paragraphs:
                    for r in par.runs:
                        r.font.size = Pt(10)
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT

def print_banner():
    """
    打印程序启动横幅（彩色 ASCII Art）。
    """
    try:
        from i18n import get_lang, t as _tt
        _lang = get_lang()
    except Exception:
        _lang = 'zh'

    def _t(key):
        try:
            return _tt(key, _lang)
        except Exception:
            return key

    try:
        import shutil
        cols = shutil.get_terminal_size((80, 20)).columns
    except Exception:
        cols = 80

    BLUE   = "\033[94m"
    GREEN  = "\033[92m"
    YELLOW = "\033[93m"
    BOLD   = "\033[1m"
    DIM    = "\033[2m"
    RESET  = "\033[0m"

    try:
        import os, ctypes
        if os.name == "nt":
            kernel32 = ctypes.windll.kernel32
            kernel32.SetConsoleMode(kernel32.GetStdHandle(-11), 7)
    except Exception:
        pass

    art = f"""
{BLUE}{BOLD}  ██████╗ ██████╗  ██████╗██╗  ██╗███████╗ ██████╗██╗  ██╗
  ██╔══██╗██╔══██╗██╔════╝██║  ██║██╔════╝██╔════╝██║ ██╔╝
  ██║  ██║██████╔╝██║     ███████║█████╗  ██║     █████╔╝
  ██║  ██║██╔══██╗██║     ██╔══██║██╔══╝  ██║     ██╔═██╗
  ██████╔╝██████╔╝╚██████╗██║  ██║███████╗╚██████╗██║  ██╗
  ╚═════╝ ╚═════╝  ╚═════╝╚═╝  ╚═╝╚══════╝ ╚═════╝╚═╝  ╚═╝{RESET}
{GREEN}{BOLD}             🐘  {_t('pg_cli_banner_title')}  {VER}{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
{YELLOW}  {_t('pg_cli_banner_subtitle')}{RESET}
{DIM}  ──────────────────────────────────────────────────────────{RESET}
"""
    print(art)

def single_inspection():
    """
    执行单机巡检流程。

    调用 input_db_info() 进行交互式连接信息输入，
    输入有效后调用 run_inspection() 执行巡检并生成报告。
    """
    print("\n=== " + _t("pg_cli_single_mode") + " ===")
    db_info = input_db_info()
    if not db_info:
        return
    run_inspection(db_info)

def batch_inspection():
    """
    执行批量巡检流程（从 Excel 文件导入数据库列表）。

    检查 Excel 模板文件是否存在，不存在时提示用户创建；
    读取模板后展示数据库列表并请用户确认，
    逐一调用 run_inspection() 完成每个数据库的巡检，
    最终汇总输出成功 / 失败统计。
    """
    print("\n=== " + _t("pg_cli_batch_mode") + " ===")
    excel_manager = ExcelTemplateManager()
    if not os.path.exists(excel_manager.template_file):
        print("\u274c " + _t("pg_cli_excel_not_exist"))
        create_template = input(_t("pg_cli_create_template_now")).strip().lower()
        if create_template in ['', 'y', 'yes']:
            excel_manager.create_template()
        return
    db_list = excel_manager.read_template()
    if not db_list:
        return
    print("\n\U0001f4cb " + _t("pg_cli_will_inspect_n").format(n=len(db_list)))
    for i, db in enumerate(db_list, 1):
        ssh_suffix = " " + _t("cli_ssh_suffix") if db.get("ssh_host") and (db.get("ssh_password") or db.get("ssh_key_file")) else ""
        print("  " + str(i) + ". " + db["name"] + " - " + db["ip"] + ":" + str(db["port"]) + ssh_suffix)
    confirm = input("\n" + _t("pg_cli_confirm_batch")).strip().lower()
    if confirm in ['', 'y', 'yes']:
        total_dbs = len(db_list)
        success_count = 0
        for i, db_info in enumerate(db_list, 1):
            print("\n[" + str(i) + "/" + str(total_dbs) + "] " + _t("pg_cli_start_inspect_n").format(name=db_info["name"]))
            if run_inspection(db_info):
                success_count += 1
        print("\n=== " + _t("pg_cli_batch_done") + " ===")
        print(_t("pg_cli_success_count").format(s=success_count, t=total_dbs))
        print(_t("pg_cli_report_dir"))

def create_excel_template():
    """
    创建批量巡检 Excel 配置模板。

    实例化 ExcelTemplateManager 并调用其 create_template() 方法，
    在当前目录生成 pg_batch_template.xlsx 文件。
    """
    print("\n=== " + _t("pg_cli_create_excel") + " ===")
    excel_manager = ExcelTemplateManager()
    excel_manager.create_template()

def create_word_template(inspector_name="Jack"):
    """
    在系统临时目录创建 Word 巡检报告模板文件。

    实例化 WordTemplateGenerator，调用 create_template() 生成包含
    Jinja2 占位符的 Word 模板文档，并保存到系统临时目录。

    :return: 成功时返回模板文件的完整路径（字符串）；失败时返回 None
    """
    try:
        temp_dir = tempfile.gettempdir()
        template_path = os.path.join(temp_dir, "pg_inspection_template.docx")
        generator = WordTemplateGenerator(inspector_name)
        doc = generator.create_template()
        doc.save(template_path)
        print("\u2705 " + _t("pg_cli_word_ok").format(path=template_path))
        return template_path
    except Exception as e:
        print("\u274c " + _t("pg_cli_word_fail").format(e=e))
        import traceback
        traceback.print_exc()
        return None

def run_inspection(db_info):
    """
    对单个数据库实例执行完整的巡检流程并生成 Word 报告。

    主要步骤：
    1. 从 db_info 中提取连接参数和 SSH 信息
    2. 调用 create_word_template() 生成临时 Word 模板文件
    3. 在 reports/ 目录下以时间戳命名输出文件
    4. 预先测试 PostgreSQL 连接并获取版本号
    5. 实例化 getData 执行巡检（checkdb）
    6. 将结果通过 saveDoc.contextsave() 渲染为最终报告
    7. 清理临时模板文件

    :param db_info: 包含数据库连接信息的字典，必须含 name、ip、port、user、password；
                    可选 ssh_host、ssh_port、ssh_user、ssh_password、ssh_key_file
    :return: 巡检并生成报告成功返回 True，任何环节失败返回 False
    """
    label_name = db_info['name']
    ip = db_info['ip']
    port = db_info['port']
    user = db_info['user']
    password = db_info['password']
    ssh_info = {}
    if 'ssh_host' in db_info and db_info['ssh_host']:
        ssh_info = {
            'ssh_host': db_info['ssh_host'],
            'ssh_port': db_info.get('ssh_port', 22),
            'ssh_user': db_info.get('ssh_user', 'root'),
            'ssh_password': db_info.get('ssh_password', ''),
            'ssh_key_file': db_info.get('ssh_key_file', '')
        }
    inspector_name = input(_t("pg_cli_inspector_prompt")).strip() or "Jack"
    ifile = create_word_template(inspector_name)
    if not ifile:
        return False
    dir_path = "reports"
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    file_name = _t("pg_cli_report_filename").format(name=label_name, ts=timestamp) + ".docx"
    ofile = os.path.join(dir_path, file_name)
    try:
        print("\U0001f50d " + _t("pg_cli_testing_connection").format(ip=ip, port=port))
        conn_test = psycopg2.connect(
            host=ip, port=int(port), user=user, password=password,
            dbname='postgres', client_encoding='UTF8', connect_timeout=10
        )
        cursor = conn_test.cursor()
        cursor.execute("SELECT version()")
        version = cursor.fetchone()[0]
        cursor.close()
        conn_test.close()
        print("\U0001f4ca " + _t("pg_cli_version").format(ver=version))
    except Exception as e:
        print("\u274c " + _t("pg_cli_conn_fail").format(e=e))
    data = getData(ip, port, user, password, database=db_info.get('database', 'postgres'), ssh_info=ssh_info, label=label_name)
    if data is None or data.conn_db2 is None:
        return False
    ret = data.checkdb('builtin')
    if not ret:
        return False
    ret.update({"co_name": [{'CO_NAME': label_name}]})
    ret.update({"port": [{'PORT': port}]})
    ret.update({"ip": [{'IP': ip}]})
    savedoc = saveDoc(context=ret, ofile=ofile, ifile=ifile, inspector_name=inspector_name)
    success = savedoc.contextsave()
    if success:
        print("\u2705 " + _t("pg_cli_report_ok").format(fname=file_name))
        try:
            if os.path.exists(ifile):
                os.remove(ifile)
        except:
            pass
        return True
    else:
        print("❌ " + _t("pg_cli_report_fail").format(name=label_name))
        return False

def main():
    """
    程序主入口函数。

    执行流程：
    1. 记录程序启动时间
    2. 打印横幅（print_banner）
    3. 进入主菜单循环：
       - 选项 1：执行单机巡检（single_inspection）
       - 选项 2：执行批量巡检（batch_inspection）
       - 选项 3：创建 Excel 模板（create_excel_template）
       - 选项 4：退出程序
    4. 每次操作完成后询问是否返回主菜单
    5. 程序退出前打印总运行耗时
    """
    start_time = time.time()

    # 支持从主入口通过 --template 直接生成 Excel 模板
    if len(sys.argv) > 1 and sys.argv[1] == '--template':
        create_excel_template()
        return

    print_banner()
    while True:
        choice = show_main_menu()
        if choice == '1':
            single_inspection()
        elif choice == '2':
            batch_inspection()
        elif choice == '3':
            create_excel_template()
        elif choice == '4':
            print("\n" + _t("pg_cli_thanks"))
            break
        if choice != '4':
            continue_choice = input("\n" + _t("pg_cli_back_menu")).strip().lower()
            if continue_choice in ['', 'y', 'yes']:
                continue
            else:
                print("\n" + _t("pg_cli_thanks"))
                break
    end_time = time.time()
    print("\n" + _t("pg_cli_total_time").format(t=end_time - start_time))

if __name__ == '__main__':
    infos = passArgu().get_argus()
    main()
